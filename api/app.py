import os
import re
import sys
import io
import base64
import math
import secrets
import subprocess
import threading
import queue as _queue_module
import time
import uuid
import shutil
import json
import logging
import sqlite3
import hashlib
import socket
import urllib.parse
try:
    import psycopg2
    import psycopg2.extras
except ImportError:
    psycopg2 = None
# Tuple of integrity-error types so except clauses work for both backends

try:
    import boto3
    from botocore.exceptions import ClientError as _BotocoreClientError
except ImportError:
    boto3 = None
    _BotocoreClientError = Exception
_IntegrityErrors: tuple = (sqlite3.IntegrityError,)
if psycopg2 is not None:
    _IntegrityErrors = (sqlite3.IntegrityError, psycopg2.IntegrityError)
import gzip
import zipfile
import ipaddress
import maxminddb
import asyncio
import mimetypes
import certifi
import yt_dlp
from pathlib import Path
from datetime import datetime, timedelta, timezone
from functools import wraps

import socketio as socketio_pkg
from fastapi import FastAPI, Request, Form, UploadFile, File, HTTPException
from fastapi.responses import JSONResponse, FileResponse, RedirectResponse, Response
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel
from starlette.middleware.sessions import SessionMiddleware
from starlette.middleware.cors import CORSMiddleware
from starlette.exceptions import HTTPException as StarletteHTTPException
from werkzeug.security import generate_password_hash, check_password_hash

# =========================================================
# PRODUCTION CONFIGURATION
# =========================================================

class Config:
    """Production configuration"""
    SECRET_KEY = os.environ.get("SECRET_KEY", os.urandom(24).hex())
    MAX_CONTENT_LENGTH = 100 * 1024 * 1024  # 100MB max upload
    DOWNLOAD_FOLDER = "downloads"
    TEMPLATES_FOLDER = "templates"
    STATIC_FOLDER = "static"
    MAX_DOWNLOADS_PER_IP = 10
    MAX_REVIEWS_PER_IP = 1
    MAX_CONCURRENT_DOWNLOADS = int(os.environ.get("MAX_CONCURRENT_DOWNLOADS", 5))
    MAX_QUEUE_SIZE = int(os.environ.get("MAX_QUEUE_SIZE", 50))
    DOWNLOAD_TIMEOUT = 3600  # 1 hour
    CLEANUP_INTERVAL = 60    # Run cleanup every 60 seconds
    FILE_RETENTION_MINUTES = 1  # Delete files older than 1 minute (~60 seconds) on each cleanup cycle
    SESSION_TYPE = 'filesystem'
    PERMANENT_SESSION_LIFETIME = timedelta(hours=1)
    # Admin authentication — set ADMIN_PASSWORD env var in production.
    # If the env var is not provided a cryptographically random secret is
    # generated so that every fresh deployment has a unique, strong password.
    ADMIN_PASSWORD = os.environ.get("ADMIN_PASSWORD") or secrets.token_hex(16)

# =========================================================
# PATH SETUP
# =========================================================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ROOT_DIR = BASE_DIR  # Assuming app.py is in root

# Handle case where app.py is in api/ subdirectory
if os.path.basename(BASE_DIR) == "api":
    ROOT_DIR = os.path.dirname(BASE_DIR)

TEMPLATES_DIR = os.path.join(ROOT_DIR, Config.TEMPLATES_FOLDER)
STATIC_DIR = os.path.join(ROOT_DIR, Config.STATIC_FOLDER)
# DOWNLOAD_DIR env var lets operators point downloads at any writable path
# (e.g. a mounted volume).  Falls back to /tmp/downloads when the default
# app-local path cannot be created so the server stays deployable on
# read-only or restricted file-systems (Railway, Render, Fly, etc.).
_configured_download_folder = os.environ.get(
    "DOWNLOAD_DIR", os.path.join(ROOT_DIR, Config.DOWNLOAD_FOLDER)
)
try:
    os.makedirs(_configured_download_folder, exist_ok=True)
    DOWNLOAD_FOLDER = _configured_download_folder
except OSError:
    DOWNLOAD_FOLDER = os.path.join("/tmp", "downloads")
    os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)


def _ensure_download_folder() -> bool:
    """Re-create DOWNLOAD_FOLDER if it was removed after startup.

    Returns True when the directory exists (or was just created), False when
    creation failed (e.g. read-only filesystem after the initial fallback).
    This prevents FileNotFoundError in os.listdir / file writes on fresh
    deployments or after container restarts that wipe the ephemeral fs.
    """
    if os.path.isdir(DOWNLOAD_FOLDER):
        return True
    try:
        os.makedirs(DOWNLOAD_FOLDER, exist_ok=True)
        return True
    except OSError as exc:
        logger.warning("Could not re-create download folder %r: %s", DOWNLOAD_FOLDER, exc)
        return False

DATA_DIR = os.path.join(ROOT_DIR, "data")
COOKIES_FILE = os.environ.get("COOKIES_FILE", os.path.join(DATA_DIR, "cookies.txt"))
# React frontend build output
FRONTEND_DIST = os.path.join(ROOT_DIR, "frontend_dist")

# Avatar upload directory (served as /static/avatars/)
AVATARS_DIR = os.path.join(STATIC_DIR, "avatars")

# Create remaining directories (non-fatal if they fail on restricted systems)
for _d in (TEMPLATES_DIR, STATIC_DIR, DATA_DIR, AVATARS_DIR):
    try:
        os.makedirs(_d, exist_ok=True)
    except OSError:
        pass

# =========================================================
# LOGGING SETUP
# =========================================================

logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.StreamHandler(sys.stdout),
        *(
            [logging.FileHandler(os.path.join(ROOT_DIR, 'app.log'))]
            if os.access(ROOT_DIR, os.W_OK)
            else []
        ),
    ]
)
logger = logging.getLogger(__name__)

# =========================================================
# FASTAPI APP INITIALIZATION
# =========================================================

_allowed_origins = os.environ.get("ALLOWED_ORIGINS", "*").split(",")

# Event loop reference for thread-safe Socket.IO emits
_loop: asyncio.AbstractEventLoop | None = None

from contextlib import asynccontextmanager

@asynccontextmanager
async def lifespan(application):
    global _loop
    _loop = asyncio.get_running_loop()
    # Ensure all database tables exist before the app starts serving requests.
    init_db()
    # Initialize the local GeoIP database in a background thread so the
    # server stays responsive while the (potentially large) download runs.
    threading.Thread(target=_init_geoip_db, daemon=True).start()
    yield

fastapi_app = FastAPI(lifespan=lifespan)

# Session middleware (replaces Flask's built-in session)
fastapi_app.add_middleware(SessionMiddleware, secret_key=Config.SECRET_KEY)

# CORS middleware (replaces Flask-CORS)
fastapi_app.add_middleware(
    CORSMiddleware,
    allow_origins=_allowed_origins,
    allow_methods=["GET", "POST", "DELETE", "OPTIONS"],
    allow_headers=["Content-Type"],
)

# Jinja2 template rendering (still used for legacy /admin/login fallback)
templates = Jinja2Templates(directory=TEMPLATES_DIR)

# Mount static files
from starlette.staticfiles import StaticFiles
if os.path.isdir(STATIC_DIR):
    fastapi_app.mount("/static", StaticFiles(directory=STATIC_DIR), name="static")

# Mount React frontend build assets (CSS/JS chunks)
if os.path.isdir(FRONTEND_DIST):
    _frontend_assets = os.path.join(FRONTEND_DIST, "assets")
    if os.path.isdir(_frontend_assets):
        fastapi_app.mount("/assets", StaticFiles(directory=_frontend_assets), name="frontend_assets")

# Socket.IO with ASGI mode
sio = socketio_pkg.AsyncServer(
    async_mode="asgi",
    cors_allowed_origins=_allowed_origins,
    logger=True if os.environ.get("FLASK_DEBUG") else False,
    engineio_logger=True if os.environ.get("FLASK_DEBUG") else False,
    ping_timeout=60,
    ping_interval=25,
    max_http_buffer_size=int(1e8),
)

# The top-level ASGI app: Socket.IO wrapping FastAPI
app = socketio_pkg.ASGIApp(sio, other_asgi_app=fastapi_app)


def emit_from_thread(event, data=None, room=None):
    """Thread-safe wrapper to emit Socket.IO events from background threads."""
    if _loop is None:
        return
    try:
        asyncio.run_coroutine_threadsafe(
            sio.emit(event, data, room=room),
            _loop,
        )
    except Exception as e:
        logger.error(f"emit_from_thread error: {e}")

# =========================================================
# GLOBAL STATE WITH THREAD SAFETY
# =========================================================

from threading import Lock, RLock
import urllib.request

downloads_lock = RLock()  # Reentrant lock for nested access
downloads = {}            # download_id -> metadata
active_threads = {}       # download_id -> thread
ip_download_count = {}    # ip -> count

# ── Download queue infrastructure ──────────────────────────────────────────
# Pending downloads wait in this queue until a concurrency slot opens.
_download_queue: _queue_module.Queue = _queue_module.Queue(maxsize=Config.MAX_QUEUE_SIZE)
# BoundedSemaphore limits how many downloads run simultaneously.
_active_semaphore = threading.BoundedSemaphore(Config.MAX_CONCURRENT_DOWNLOADS)

visitors_lock = Lock()
visitors = []             # List of visitor dicts tracked on page visits
ip_country_cache = {}     # ip -> {"country": str, "code": str}

reviews_lock = Lock()
reviews: list = []        # List of review dicts submitted by visitors

# Paths for persistent storage
DATA_DIR = os.path.join(ROOT_DIR, "data")
DB_PATH = os.path.join(DATA_DIR, "admin.db")
os.makedirs(DATA_DIR, exist_ok=True)

# PostgreSQL support: when DATABASE_URL is set, use PostgreSQL instead of SQLite
DATABASE_URL = os.environ.get("DATABASE_URL")
USE_POSTGRES = bool(DATABASE_URL and psycopg2 is not None)

# =========================================================
# S3-COMPATIBLE OBJECT STORAGE
# Set all five BUCKET_* variables to enable cloud storage.
# When configured, completed downloads are uploaded to the
# bucket and served via presigned URLs; files are also deleted
# from the bucket on removal/cleanup.
# =========================================================

BUCKET_NAME             = os.environ.get("BUCKET_NAME", "")
BUCKET_REGION           = os.environ.get("BUCKET_REGION", "")
BUCKET_ENDPOINT         = os.environ.get("BUCKET_ENDPOINT", "")
BUCKET_ACCESS_KEY_ID    = os.environ.get("BUCKET_ACCESS_KEY_ID", "")
BUCKET_SECRET_ACCESS_KEY = os.environ.get("BUCKET_SECRET_ACCESS_KEY", "")

_S3_ENABLED = bool(
    boto3 is not None
    and BUCKET_NAME
    and BUCKET_ACCESS_KEY_ID
    and BUCKET_SECRET_ACCESS_KEY
)

_s3_client_instance = None
_s3_client_lock = threading.Lock()


def _get_s3_client():
    """Return a cached boto3 S3 client, or *None* when S3 is not configured."""
    global _s3_client_instance
    if not _S3_ENABLED:
        return None
    if _s3_client_instance is not None:
        return _s3_client_instance
    with _s3_client_lock:
        if _s3_client_instance is None:
            kwargs = {
                "aws_access_key_id":     BUCKET_ACCESS_KEY_ID,
                "aws_secret_access_key": BUCKET_SECRET_ACCESS_KEY,
            }
            if BUCKET_REGION:
                kwargs["region_name"] = BUCKET_REGION
            if BUCKET_ENDPOINT:
                kwargs["endpoint_url"] = BUCKET_ENDPOINT
            _s3_client_instance = boto3.client("s3", **kwargs)
    return _s3_client_instance


def _s3_upload_file(local_path: str, key: str) -> bool:
    """Upload *local_path* to the configured bucket under *key*.

    Returns ``True`` on success, ``False`` when S3 is not configured or the
    upload fails (errors are logged but never re-raised so callers always get
    a usable result).
    """
    client = _get_s3_client()
    if client is None:
        return False
    try:
        client.upload_file(local_path, BUCKET_NAME, key)
        logger.info("S3 upload: %s → s3://%s/%s", local_path, BUCKET_NAME, key)
        return True
    except Exception as exc:
        logger.error("S3 upload failed for %s: %s", key, exc)
        return False


def _s3_delete_file(key: str) -> bool:
    """Delete *key* from the configured bucket.

    Returns ``True`` on success, ``False`` otherwise.
    """
    client = _get_s3_client()
    if client is None:
        return False
    try:
        client.delete_object(Bucket=BUCKET_NAME, Key=key)
        logger.info("S3 delete: s3://%s/%s", BUCKET_NAME, key)
        return True
    except Exception as exc:
        logger.error("S3 delete failed for %s: %s", key, exc)
        return False


def _s3_presigned_url(key: str, expires_in: int = 3600) -> str | None:
    """Generate a presigned download URL for *key* (default expiry 1 hour).

    Returns the URL string on success, or ``None`` when S3 is not configured or
    URL generation fails.
    """
    client = _get_s3_client()
    if client is None:
        return None
    try:
        url = client.generate_presigned_url(
            "get_object",
            Params={"Bucket": BUCKET_NAME, "Key": key},
            ExpiresIn=expires_in,
        )
        return url
    except Exception as exc:
        logger.error("S3 presigned URL failed for %s: %s", key, exc)
        return None


def _s3_object_exists(key: str) -> bool:
    """Return *True* when *key* exists in the configured bucket."""
    client = _get_s3_client()
    if client is None:
        return False
    try:
        client.head_object(Bucket=BUCKET_NAME, Key=key)
        return True
    except _BotocoreClientError:
        return False
    except Exception as exc:
        logger.error("S3 head_object failed for %s: %s", key, exc)
        return False


def _s3_upload_bytes(data: bytes, key: str, content_type: str = "application/octet-stream") -> bool:
    """Upload raw *data* bytes to the configured bucket under *key*.

    Returns ``True`` on success, ``False`` when S3 is not configured or the
    upload fails (errors are logged but never re-raised).
    """
    client = _get_s3_client()
    if client is None:
        return False
    try:
        client.put_object(Body=io.BytesIO(data), Bucket=BUCKET_NAME, Key=key, ContentType=content_type)
        logger.info("S3 upload bytes: s3://%s/%s (%d bytes)", BUCKET_NAME, key, len(data))
        return True
    except Exception as exc:
        logger.error("S3 upload bytes failed for %s: %s", key, exc)
        return False


def _bucket_write_json(folder: str, type_prefix: str, record_id: str, data: dict) -> bool:
    """Write *data* as a JSON file to *folder* in the configured bucket.

    The key is ``{folder}/{type_prefix}_{timestamp}_{record_id}.json``.
    This follows the bucket file-naming convention:
        {type}_{timestamp}_{uuid}.json
    Returns ``True`` on success, ``False`` when S3 is not configured or the
    write fails (errors are logged but never re-raised).
    """
    ts = datetime.now(timezone.utc).strftime("%Y%m%d_%H%M%S")
    key = f"{folder.rstrip('/')}/{type_prefix}_{ts}_{record_id}.json"
    payload = json.dumps(data, default=str).encode("utf-8")
    return _s3_upload_bytes(payload, key, "application/json")


# Paths whose visits are tracked for analytics (main site + admin page)
TRACKED_VISITOR_PATHS = {"/const", "/"}

# =========================================================
# ISO 3166-1 ALPHA-2 → COUNTRY NAME LOOKUP
# =========================================================

_ISO2_TO_NAME: dict[str, str] = {
    "AF": "Afghanistan", "AX": "Aland Islands", "AL": "Albania",
    "DZ": "Algeria", "AS": "American Samoa", "AD": "Andorra",
    "AO": "Angola", "AI": "Anguilla", "AQ": "Antarctica",
    "AG": "Antigua and Barbuda", "AR": "Argentina", "AM": "Armenia",
    "AW": "Aruba", "AU": "Australia", "AT": "Austria", "AZ": "Azerbaijan",
    "BS": "Bahamas", "BH": "Bahrain", "BD": "Bangladesh", "BB": "Barbados",
    "BY": "Belarus", "BE": "Belgium", "BZ": "Belize", "BJ": "Benin",
    "BM": "Bermuda", "BT": "Bhutan", "BO": "Bolivia",
    "BQ": "Bonaire, Sint Eustatius and Saba",
    "BA": "Bosnia and Herzegovina", "BW": "Botswana",
    "BV": "Bouvet Island", "BR": "Brazil",
    "IO": "British Indian Ocean Territory", "BN": "Brunei",
    "BG": "Bulgaria", "BF": "Burkina Faso", "BI": "Burundi",
    "CV": "Cabo Verde", "KH": "Cambodia", "CM": "Cameroon", "CA": "Canada",
    "KY": "Cayman Islands", "CF": "Central African Republic", "TD": "Chad",
    "CL": "Chile", "CN": "China", "CX": "Christmas Island",
    "CC": "Cocos (Keeling) Islands", "CO": "Colombia", "KM": "Comoros",
    "CG": "Congo", "CD": "DR Congo", "CK": "Cook Islands",
    "CR": "Costa Rica", "CI": "Ivory Coast", "HR": "Croatia", "CU": "Cuba",
    "CW": "Curacao", "CY": "Cyprus", "CZ": "Czech Republic",
    "DK": "Denmark", "DJ": "Djibouti", "DM": "Dominica",
    "DO": "Dominican Republic", "EC": "Ecuador", "EG": "Egypt",
    "SV": "El Salvador", "GQ": "Equatorial Guinea", "ER": "Eritrea",
    "EE": "Estonia", "SZ": "Eswatini", "ET": "Ethiopia",
    "FK": "Falkland Islands", "FO": "Faroe Islands", "FJ": "Fiji",
    "FI": "Finland", "FR": "France", "GF": "French Guiana",
    "PF": "French Polynesia", "TF": "French Southern Territories",
    "GA": "Gabon", "GM": "Gambia", "GE": "Georgia", "DE": "Germany",
    "GH": "Ghana", "GI": "Gibraltar", "GR": "Greece", "GL": "Greenland",
    "GD": "Grenada", "GP": "Guadeloupe", "GU": "Guam",
    "GT": "Guatemala", "GG": "Guernsey", "GN": "Guinea",
    "GW": "Guinea-Bissau", "GY": "Guyana", "HT": "Haiti",
    "HM": "Heard Island and McDonald Islands",
    "VA": "Holy See", "HN": "Honduras", "HK": "Hong Kong",
    "HU": "Hungary", "IS": "Iceland", "IN": "India", "ID": "Indonesia",
    "IR": "Iran", "IQ": "Iraq", "IE": "Ireland", "IM": "Isle of Man",
    "IL": "Israel", "IT": "Italy", "JM": "Jamaica", "JP": "Japan",
    "JE": "Jersey", "JO": "Jordan", "KZ": "Kazakhstan", "KE": "Kenya",
    "KI": "Kiribati", "KP": "North Korea", "KR": "South Korea",
    "KW": "Kuwait", "KG": "Kyrgyzstan", "LA": "Laos", "LV": "Latvia",
    "LB": "Lebanon", "LS": "Lesotho", "LR": "Liberia", "LY": "Libya",
    "LI": "Liechtenstein", "LT": "Lithuania", "LU": "Luxembourg",
    "MO": "Macao", "MG": "Madagascar", "MW": "Malawi", "MY": "Malaysia",
    "MV": "Maldives", "ML": "Mali", "MT": "Malta", "MH": "Marshall Islands",
    "MQ": "Martinique", "MR": "Mauritania", "MU": "Mauritius",
    "YT": "Mayotte", "MX": "Mexico", "FM": "Micronesia", "MD": "Moldova",
    "MC": "Monaco", "MN": "Mongolia", "ME": "Montenegro",
    "MS": "Montserrat", "MA": "Morocco", "MZ": "Mozambique",
    "MM": "Myanmar", "NA": "Namibia", "NR": "Nauru", "NP": "Nepal",
    "NL": "Netherlands", "NC": "New Caledonia", "NZ": "New Zealand",
    "NI": "Nicaragua", "NE": "Niger", "NG": "Nigeria", "NU": "Niue",
    "NF": "Norfolk Island", "MK": "North Macedonia",
    "MP": "Northern Mariana Islands", "NO": "Norway", "OM": "Oman",
    "PK": "Pakistan", "PW": "Palau", "PS": "Palestine", "PA": "Panama",
    "PG": "Papua New Guinea", "PY": "Paraguay", "PE": "Peru",
    "PH": "Philippines", "PN": "Pitcairn", "PL": "Poland",
    "PT": "Portugal", "PR": "Puerto Rico", "QA": "Qatar",
    "RE": "Reunion", "RO": "Romania", "RU": "Russia", "RW": "Rwanda",
    "BL": "Saint Barthelemy", "SH": "Saint Helena",
    "KN": "Saint Kitts and Nevis", "LC": "Saint Lucia",
    "MF": "Saint Martin", "PM": "Saint Pierre and Miquelon",
    "VC": "Saint Vincent and the Grenadines", "WS": "Samoa",
    "SM": "San Marino", "ST": "Sao Tome and Principe",
    "SA": "Saudi Arabia", "SN": "Senegal", "RS": "Serbia",
    "SC": "Seychelles", "SL": "Sierra Leone", "SG": "Singapore",
    "SX": "Sint Maarten", "SK": "Slovakia", "SI": "Slovenia",
    "SB": "Solomon Islands", "SO": "Somalia", "ZA": "South Africa",
    "GS": "South Georgia and the South Sandwich Islands",
    "SS": "South Sudan", "ES": "Spain", "LK": "Sri Lanka", "SD": "Sudan",
    "SR": "Suriname", "SJ": "Svalbard and Jan Mayen", "SE": "Sweden",
    "CH": "Switzerland", "SY": "Syria", "TW": "Taiwan",
    "TJ": "Tajikistan", "TZ": "Tanzania", "TH": "Thailand",
    "TL": "Timor-Leste", "TG": "Togo", "TK": "Tokelau", "TO": "Tonga",
    "TT": "Trinidad and Tobago", "TN": "Tunisia", "TR": "Turkey",
    "TM": "Turkmenistan", "TC": "Turks and Caicos Islands",
    "TV": "Tuvalu", "UG": "Uganda", "UA": "Ukraine",
    "AE": "United Arab Emirates", "GB": "United Kingdom",
    "US": "United States", "UM": "United States Minor Outlying Islands",
    "UY": "Uruguay", "UZ": "Uzbekistan", "VU": "Vanuatu",
    "VE": "Venezuela", "VN": "Vietnam",
    "VG": "British Virgin Islands", "VI": "U.S. Virgin Islands",
    "WF": "Wallis and Futuna", "EH": "Western Sahara",
    "YE": "Yemen", "ZM": "Zambia", "ZW": "Zimbabwe",
}

# =========================================================
# LOCAL GEOIP DATABASE (DB-IP City Lite – free, MMDB format)
# =========================================================

GEOIP_DB_PATH = os.path.join(DATA_DIR, "dbip-city-lite.mmdb")
GEOIP_DB_REFRESH_SECONDS = 35 * 86400  # Re-download after ~35 days (database is updated monthly)
_geoip_reader: maxminddb.Reader | None = None

# Optional ipapi.com (ipstack) API key – enables an extra geolocation fallback.
# Sign up at https://ipapi.com/signup/free for a free key, then set the env var.
IPAPI_ACCESS_KEY = os.environ.get("IPAPI_ACCESS_KEY", "")


def _init_geoip_db():
    """Download the free DB-IP City Lite MMDB database if absent or stale, then open it."""
    global _geoip_reader

    need_download = not os.path.exists(GEOIP_DB_PATH)
    if not need_download:
        age = time.time() - os.path.getmtime(GEOIP_DB_PATH)
        if age > GEOIP_DB_REFRESH_SECONDS:
            need_download = True

    if need_download:
        now = datetime.utcnow()
        url = (
            f"https://download.db-ip.com/free/"
            f"dbip-city-lite-{now.year}-{now.month:02d}.mmdb.gz"
        )
        gz_path = GEOIP_DB_PATH + ".gz"
        tmp_path = GEOIP_DB_PATH + ".tmp"
        try:
            req = urllib.request.Request(url, headers={"User-Agent": "Mozilla/5.0"})
            with urllib.request.urlopen(req, timeout=60) as resp:
                with open(gz_path, "wb") as f:
                    shutil.copyfileobj(resp, f)
            with gzip.open(gz_path, "rb") as gz_in, open(tmp_path, "wb") as out:
                shutil.copyfileobj(gz_in, out)
            os.replace(tmp_path, GEOIP_DB_PATH)
            logger.info("GeoIP database downloaded successfully")
        except Exception as e:
            logger.warning(f"Failed to download GeoIP database: {e}")
        finally:
            # Clean up temp files
            for p in (gz_path, tmp_path):
                try:
                    os.remove(p)
                except OSError:
                    pass

    if os.path.exists(GEOIP_DB_PATH):
        try:
            _geoip_reader = maxminddb.open_database(GEOIP_DB_PATH)
            logger.info("GeoIP database loaded successfully")
        except Exception as e:
            logger.warning(f"Failed to open GeoIP database: {e}")
            _geoip_reader = None


# =========================================================
# SSL CERTIFICATE FIX
# =========================================================

os.environ['SSL_CERT_FILE'] = certifi.where()
os.environ['REQUESTS_CA_BUNDLE'] = certifi.where()

# =========================================================
# DATABASE PERSISTENCE  (PostgreSQL when DATABASE_URL is set, else SQLite)
# =========================================================

_db_lock = threading.Lock()


def _get_db():
    """Open a short-lived database connection for the calling thread.

    Callers hold _db_lock and immediately close the connection after use,
    so each connection is used by exactly one thread.
    """
    if USE_POSTGRES:
        conn = psycopg2.connect(DATABASE_URL)
        return conn
    else:
        conn = sqlite3.connect(DB_PATH)
        conn.row_factory = sqlite3.Row
        conn.execute("PRAGMA journal_mode=WAL")
        conn.execute("PRAGMA foreign_keys=ON")
        return conn


def _execute(conn, sql, params=None):
    """Execute *sql* on *conn* and return a cursor-like object.

    Handles the dialect differences between SQLite (``?`` placeholders,
    ``conn.execute``) and PostgreSQL (``%s`` placeholders, cursor with
    ``RealDictCursor``).

    Note: The ``?`` → ``%s`` replacement is safe here because all SQL in this
    application uses ``?`` exclusively as parameter placeholders (no question
    marks appear inside string literals or comments).
    """
    if USE_POSTGRES:
        sql = sql.replace("?", "%s")
        cur = conn.cursor(cursor_factory=psycopg2.extras.RealDictCursor)
        cur.execute(sql, params)
        return cur
    else:
        return conn.execute(sql, params or ())


def init_db():
    """Create tables if they don't exist."""
    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS downloads (
                        id TEXT PRIMARY KEY,
                        data TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS visitors (
                        id SERIAL PRIMARY KEY,
                        data TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS admin_users (
                        id SERIAL PRIMARY KEY,
                        username TEXT UNIQUE NOT NULL,
                        password_hash TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS reviews (
                        id SERIAL PRIMARY KEY,
                        data TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS app_users (
                        id SERIAL PRIMARY KEY,
                        user_id TEXT UNIQUE NOT NULL,
                        name TEXT NOT NULL,
                        email TEXT UNIQUE NOT NULL,
                        password_hash TEXT NOT NULL,
                        role TEXT NOT NULL DEFAULT 'passenger',
                        location_lat REAL,
                        location_lng REAL,
                        location_name TEXT,
                        avatar_url TEXT,
                        bio TEXT,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS rides (
                        id SERIAL PRIMARY KEY,
                        ride_id TEXT UNIQUE NOT NULL,
                        user_id TEXT NOT NULL,
                        driver_name TEXT NOT NULL,
                        origin TEXT NOT NULL,
                        destination TEXT NOT NULL,
                        origin_lat REAL,
                        origin_lng REAL,
                        dest_lat REAL,
                        dest_lng REAL,
                        fare REAL,
                        departure TEXT NOT NULL,
                        seats INTEGER NOT NULL DEFAULT 1,
                        notes TEXT,
                        status TEXT NOT NULL DEFAULT 'open',
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS driver_applications (
                        id SERIAL PRIMARY KEY,
                        app_id TEXT UNIQUE NOT NULL,
                        user_id TEXT UNIQUE NOT NULL,
                        vehicle_make TEXT NOT NULL,
                        vehicle_model TEXT NOT NULL,
                        vehicle_year INTEGER NOT NULL,
                        vehicle_color TEXT NOT NULL DEFAULT '',
                        license_plate TEXT NOT NULL,
                        status TEXT NOT NULL DEFAULT 'pending',
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS notifications (
                        id SERIAL PRIMARY KEY,
                        notif_id TEXT UNIQUE NOT NULL,
                        user_id TEXT NOT NULL,
                        type TEXT NOT NULL,
                        title TEXT NOT NULL,
                        body TEXT NOT NULL,
                        read INTEGER NOT NULL DEFAULT 0,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS ride_chat_messages (
                        id SERIAL PRIMARY KEY,
                        msg_id TEXT UNIQUE NOT NULL,
                        ride_id TEXT NOT NULL,
                        sender_name TEXT NOT NULL,
                        sender_role TEXT NOT NULL DEFAULT 'passenger',
                        text TEXT,
                        media_type TEXT,
                        media_data TEXT,
                        lat REAL,
                        lng REAL,
                        ts REAL NOT NULL,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS real_estate_agents (
                        id SERIAL PRIMARY KEY,
                        agent_id TEXT UNIQUE NOT NULL,
                        user_id TEXT,
                        name TEXT NOT NULL,
                        avatar TEXT,
                        bio TEXT,
                        email TEXT,
                        phone TEXT,
                        availability_status TEXT NOT NULL DEFAULT 'available',
                        lat REAL,
                        lng REAL,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS agent_reviews (
                        id SERIAL PRIMARY KEY,
                        review_id TEXT UNIQUE NOT NULL,
                        agent_id TEXT NOT NULL,
                        reviewer_user_id TEXT NOT NULL,
                        reviewer_name TEXT NOT NULL,
                        rating INTEGER NOT NULL,
                        text TEXT,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS agent_likes (
                        id SERIAL PRIMARY KEY,
                        agent_id TEXT NOT NULL,
                        user_id TEXT NOT NULL,
                        created_at TEXT NOT NULL,
                        UNIQUE(agent_id, user_id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS agent_chat_messages (
                        id SERIAL PRIMARY KEY,
                        msg_id TEXT UNIQUE NOT NULL,
                        agent_id TEXT NOT NULL,
                        user_id TEXT NOT NULL,
                        sender_role TEXT NOT NULL DEFAULT 'user',
                        text TEXT NOT NULL,
                        ts REAL NOT NULL,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS dm_conversations (
                        id SERIAL PRIMARY KEY,
                        conv_id TEXT UNIQUE NOT NULL,
                        user1_id TEXT NOT NULL,
                        user2_id TEXT NOT NULL,
                        unread_u1 INTEGER NOT NULL DEFAULT 0,
                        unread_u2 INTEGER NOT NULL DEFAULT 0,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS dm_messages (
                        id SERIAL PRIMARY KEY,
                        msg_id TEXT UNIQUE NOT NULL,
                        conv_id TEXT NOT NULL,
                        sender_id TEXT NOT NULL,
                        content TEXT NOT NULL,
                        status TEXT NOT NULL DEFAULT 'sent',
                        reply_to_id TEXT,
                        ts REAL NOT NULL,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS properties (
                        id SERIAL PRIMARY KEY,
                        property_id TEXT UNIQUE NOT NULL,
                        title TEXT NOT NULL,
                        description TEXT,
                        price REAL NOT NULL DEFAULT 0,
                        address TEXT,
                        lat REAL,
                        lng REAL,
                        images_json TEXT NOT NULL DEFAULT '[]',
                        status TEXT NOT NULL DEFAULT 'active',
                        owner_user_id TEXT,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS property_agents (
                        id SERIAL PRIMARY KEY,
                        property_id TEXT NOT NULL,
                        agent_id TEXT NOT NULL,
                        position INTEGER NOT NULL DEFAULT 0,
                        UNIQUE(property_id, agent_id)
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS property_conversations (
                        id SERIAL PRIMARY KEY,
                        conv_id TEXT UNIQUE NOT NULL,
                        property_id TEXT NOT NULL,
                        user_id TEXT NOT NULL,
                        agent_id TEXT NOT NULL,
                        unread_user INTEGER NOT NULL DEFAULT 0,
                        unread_agent INTEGER NOT NULL DEFAULT 0,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS property_messages (
                        id SERIAL PRIMARY KEY,
                        msg_id TEXT UNIQUE NOT NULL,
                        conv_id TEXT NOT NULL,
                        sender_id TEXT NOT NULL,
                        sender_role TEXT NOT NULL DEFAULT 'user',
                        content TEXT NOT NULL,
                        ts REAL NOT NULL,
                        created_at TEXT NOT NULL
                    )
                """)
                cur.execute("""
                    CREATE TABLE IF NOT EXISTS agent_applications (
                        id SERIAL PRIMARY KEY,
                        app_id TEXT UNIQUE NOT NULL,
                        user_id TEXT UNIQUE NOT NULL,
                        full_name TEXT NOT NULL,
                        email TEXT NOT NULL,
                        phone TEXT NOT NULL DEFAULT '',
                        agency_name TEXT NOT NULL DEFAULT '',
                        license_number TEXT NOT NULL,
                        status TEXT NOT NULL DEFAULT 'pending',
                        created_at TEXT NOT NULL
                    )
                """)
                conn.commit()
                # Migrations: add new columns to existing tables if needed
                for col, coldef in [("avatar_url", "TEXT"), ("bio", "TEXT"), ("public_key", "TEXT"), ("can_post_properties", "INTEGER DEFAULT 0")]:
                    try:
                        cur.execute(f"ALTER TABLE app_users ADD COLUMN {col} {coldef}")
                        conn.commit()
                    except Exception:
                        conn.rollback()
                        pass  # column already exists
                for col, coldef in [("dest_lat", "REAL"), ("dest_lng", "REAL"), ("fare", "REAL"), ("ride_type", "TEXT DEFAULT 'airport'")]:
                    try:
                        cur.execute(f"ALTER TABLE rides ADD COLUMN {col} {coldef}")
                        conn.commit()
                    except Exception:
                        conn.rollback()
                        pass  # column already exists
                for col, coldef in [("subscription_type", "TEXT DEFAULT 'monthly'")]:
                    try:
                        cur.execute(f"ALTER TABLE driver_applications ADD COLUMN {col} {coldef}")
                        conn.commit()
                    except Exception:
                        conn.rollback()
                        pass  # column already exists
            else:
                conn.executescript("""
                    CREATE TABLE IF NOT EXISTS downloads (
                        id TEXT PRIMARY KEY,
                        data TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS visitors (
                        rowid INTEGER PRIMARY KEY AUTOINCREMENT,
                        data TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS admin_users (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        username TEXT UNIQUE NOT NULL,
                        password_hash TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS reviews (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        data TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS app_users (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        user_id TEXT UNIQUE NOT NULL,
                        name TEXT NOT NULL,
                        email TEXT UNIQUE NOT NULL,
                        password_hash TEXT NOT NULL,
                        role TEXT NOT NULL DEFAULT 'passenger',
                        location_lat REAL,
                        location_lng REAL,
                        location_name TEXT,
                        avatar_url TEXT,
                        bio TEXT,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS rides (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        ride_id TEXT UNIQUE NOT NULL,
                        user_id TEXT NOT NULL,
                        driver_name TEXT NOT NULL,
                        origin TEXT NOT NULL,
                        destination TEXT NOT NULL,
                        origin_lat REAL,
                        origin_lng REAL,
                        dest_lat REAL,
                        dest_lng REAL,
                        fare REAL,
                        departure TEXT NOT NULL,
                        seats INTEGER NOT NULL DEFAULT 1,
                        notes TEXT,
                        status TEXT NOT NULL DEFAULT 'open',
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS driver_applications (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        app_id TEXT UNIQUE NOT NULL,
                        user_id TEXT UNIQUE NOT NULL,
                        vehicle_make TEXT NOT NULL,
                        vehicle_model TEXT NOT NULL,
                        vehicle_year INTEGER NOT NULL,
                        vehicle_color TEXT NOT NULL DEFAULT '',
                        license_plate TEXT NOT NULL,
                        status TEXT NOT NULL DEFAULT 'pending',
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS notifications (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        notif_id TEXT UNIQUE NOT NULL,
                        user_id TEXT NOT NULL,
                        type TEXT NOT NULL,
                        title TEXT NOT NULL,
                        body TEXT NOT NULL,
                        read INTEGER NOT NULL DEFAULT 0,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS ride_chat_messages (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        msg_id TEXT UNIQUE NOT NULL,
                        ride_id TEXT NOT NULL,
                        sender_name TEXT NOT NULL,
                        sender_role TEXT NOT NULL DEFAULT 'passenger',
                        text TEXT,
                        media_type TEXT,
                        media_data TEXT,
                        lat REAL,
                        lng REAL,
                        ts REAL NOT NULL,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS real_estate_agents (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        agent_id TEXT UNIQUE NOT NULL,
                        user_id TEXT,
                        name TEXT NOT NULL,
                        avatar TEXT,
                        bio TEXT,
                        email TEXT,
                        phone TEXT,
                        availability_status TEXT NOT NULL DEFAULT 'available',
                        lat REAL,
                        lng REAL,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS agent_reviews (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        review_id TEXT UNIQUE NOT NULL,
                        agent_id TEXT NOT NULL,
                        reviewer_user_id TEXT NOT NULL,
                        reviewer_name TEXT NOT NULL,
                        rating INTEGER NOT NULL,
                        text TEXT,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS agent_likes (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        agent_id TEXT NOT NULL,
                        user_id TEXT NOT NULL,
                        created_at TEXT NOT NULL,
                        UNIQUE(agent_id, user_id)
                    );
                    CREATE TABLE IF NOT EXISTS agent_chat_messages (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        msg_id TEXT UNIQUE NOT NULL,
                        agent_id TEXT NOT NULL,
                        user_id TEXT NOT NULL,
                        sender_role TEXT NOT NULL DEFAULT 'user',
                        text TEXT NOT NULL,
                        ts REAL NOT NULL,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS dm_conversations (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        conv_id TEXT UNIQUE NOT NULL,
                        user1_id TEXT NOT NULL,
                        user2_id TEXT NOT NULL,
                        unread_u1 INTEGER NOT NULL DEFAULT 0,
                        unread_u2 INTEGER NOT NULL DEFAULT 0,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS dm_messages (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        msg_id TEXT UNIQUE NOT NULL,
                        conv_id TEXT NOT NULL,
                        sender_id TEXT NOT NULL,
                        content TEXT NOT NULL,
                        status TEXT NOT NULL DEFAULT 'sent',
                        reply_to_id TEXT,
                        ts REAL NOT NULL,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS properties (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        property_id TEXT UNIQUE NOT NULL,
                        title TEXT NOT NULL,
                        description TEXT,
                        price REAL NOT NULL DEFAULT 0,
                        address TEXT,
                        lat REAL,
                        lng REAL,
                        images_json TEXT NOT NULL DEFAULT '[]',
                        status TEXT NOT NULL DEFAULT 'active',
                        owner_user_id TEXT,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS property_agents (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        property_id TEXT NOT NULL,
                        agent_id TEXT NOT NULL,
                        position INTEGER NOT NULL DEFAULT 0,
                        UNIQUE(property_id, agent_id)
                    );
                    CREATE TABLE IF NOT EXISTS property_conversations (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        conv_id TEXT UNIQUE NOT NULL,
                        property_id TEXT NOT NULL,
                        user_id TEXT NOT NULL,
                        agent_id TEXT NOT NULL,
                        unread_user INTEGER NOT NULL DEFAULT 0,
                        unread_agent INTEGER NOT NULL DEFAULT 0,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS property_messages (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        msg_id TEXT UNIQUE NOT NULL,
                        conv_id TEXT NOT NULL,
                        sender_id TEXT NOT NULL,
                        sender_role TEXT NOT NULL DEFAULT 'user',
                        content TEXT NOT NULL,
                        ts REAL NOT NULL,
                        created_at TEXT NOT NULL
                    );
                    CREATE TABLE IF NOT EXISTS agent_applications (
                        id INTEGER PRIMARY KEY AUTOINCREMENT,
                        app_id TEXT UNIQUE NOT NULL,
                        user_id TEXT UNIQUE NOT NULL,
                        full_name TEXT NOT NULL,
                        email TEXT NOT NULL,
                        phone TEXT NOT NULL DEFAULT '',
                        agency_name TEXT NOT NULL DEFAULT '',
                        license_number TEXT NOT NULL,
                        status TEXT NOT NULL DEFAULT 'pending',
                        created_at TEXT NOT NULL
                    );
                """)
                # SQLite migrations: add new columns to existing tables if needed
                for col, coldef in [("avatar_url", "TEXT"), ("bio", "TEXT"), ("public_key", "TEXT"), ("can_post_properties", "INTEGER DEFAULT 0")]:
                    try:
                        conn.execute(f"ALTER TABLE app_users ADD COLUMN {col} {coldef}")
                    except Exception:
                        pass  # column already exists
                for col, coldef in [("dest_lat", "REAL"), ("dest_lng", "REAL"), ("fare", "REAL"), ("ride_type", "TEXT DEFAULT 'airport'")]:
                    try:
                        conn.execute(f"ALTER TABLE rides ADD COLUMN {col} {coldef}")
                    except Exception:
                        pass  # column already exists
                for col, coldef in [("subscription_type", "TEXT DEFAULT 'monthly'")]:
                    try:
                        conn.execute(f"ALTER TABLE driver_applications ADD COLUMN {col} {coldef}")
                    except Exception:
                        pass  # column already exists
            conn.commit()
        finally:
            conn.close()


def admin_user_exists() -> bool:
    """Return True if at least one admin account is registered."""
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(conn, "SELECT COUNT(*) AS cnt FROM admin_users")
            row = cur.fetchone()
            return (row["cnt"] if USE_POSTGRES else row[0]) > 0
        finally:
            conn.close()


def register_admin_user(username: str, password: str) -> tuple[bool, str]:
    """Register admin. Returns (success, error_message).

    Passwords are stored using werkzeug's PBKDF2-HMAC-SHA256 implementation
    which includes a per-user salt and iteration count.
    """
    if admin_user_exists():
        return False, "User already exists"
    ph = generate_password_hash(password)
    with _db_lock:
        conn = _get_db()
        try:
            _execute(
                conn,
                "INSERT INTO admin_users (username, password_hash) VALUES (?, ?)",
                (username, ph),
            )
            conn.commit()
            return True, ""
        except _IntegrityErrors:
            conn.rollback()
            return False, "User already exists"
        finally:
            conn.close()


def verify_admin_user(username: str, password: str) -> bool:
    """Check username+password against the stored hash."""
    with _db_lock:
        conn = _get_db()
        try:
            row = _execute(
                conn,
                "SELECT password_hash FROM admin_users WHERE username=?",
                (username,),
            ).fetchone()
        finally:
            conn.close()
    if row is None:
        return False
    return check_password_hash(row["password_hash"], password)


def load_persistence():
    """Load downloads and visitors from the database on startup."""
    global downloads, visitors
    # Migrate legacy JSON files into the database (one-time migration)
    _migrate_json_to_db()

    with _db_lock:
        conn = _get_db()
        try:
            rows = _execute(conn, "SELECT id, data FROM downloads").fetchall()
            saved_dl = {r["id"]: json.loads(r["data"]) for r in rows}
            visitor_order_col = "id" if USE_POSTGRES else "rowid"
            rows_v = _execute(conn, f"SELECT data FROM visitors ORDER BY {visitor_order_col}").fetchall()
            saved_v = [json.loads(r["data"]) for r in rows_v]
        except Exception as exc:
            # Catch missing-table errors from both backends so the app can
            # still start when init_db() hasn't run yet or failed silently.
            _is_missing_table = isinstance(exc, sqlite3.OperationalError) or (
                psycopg2 is not None
                and isinstance(exc, psycopg2.errors.UndefinedTable)
            )
            if _is_missing_table:
                logger.warning(
                    "downloads/visitors table does not exist yet — skipping "
                    "load_persistence(). Tables will be created by init_db()."
                )
                conn.close()
                return
            conn.close()
            raise
        else:
            conn.close()

    with downloads_lock:
        downloads.update(saved_dl)
    with visitors_lock:
        visitors.extend(saved_v)
    logger.info(f"Loaded {len(saved_dl)} download records and {len(saved_v)} visitor records from {'PostgreSQL' if USE_POSTGRES else 'SQLite'}")

    # Load reviews
    _load_reviews_from_db()
    with reviews_lock:
        logger.info(f"Loaded {len(reviews)} review records from {'PostgreSQL' if USE_POSTGRES else 'SQLite'}")


def _migrate_json_to_db():
    """One-time migration: import legacy JSON persistence files into the database."""
    legacy_dl = os.path.join(DATA_DIR, "downloads.json")
    legacy_v = os.path.join(DATA_DIR, "visitors.json")
    conn = _get_db()
    try:
        if os.path.exists(legacy_dl):
            try:
                with open(legacy_dl) as fh:
                    data = json.load(fh)
                for did, d in data.items():
                    if USE_POSTGRES:
                        cur = conn.cursor()
                        cur.execute(
                            "INSERT INTO downloads (id, data) VALUES (%s, %s) ON CONFLICT (id) DO NOTHING",
                            (did, json.dumps(d, default=str)),
                        )
                    else:
                        conn.execute(
                            "INSERT OR IGNORE INTO downloads (id, data) VALUES (?, ?)",
                            (did, json.dumps(d, default=str)),
                        )
                conn.commit()
                os.rename(legacy_dl, legacy_dl + ".migrated")
                logger.info(f"Migrated {len(data)} download records from JSON to {'PostgreSQL' if USE_POSTGRES else 'SQLite'}")
            except Exception as exc:
                logger.error(f"JSON→DB migration failed for downloads: {exc}")
        if os.path.exists(legacy_v):
            try:
                with open(legacy_v) as fh:
                    data = json.load(fh)
                for v in data:
                    _execute(
                        conn,
                        "INSERT INTO visitors (data) VALUES (?)",
                        (json.dumps(v, default=str),),
                    )
                conn.commit()
                os.rename(legacy_v, legacy_v + ".migrated")
                logger.info(f"Migrated {len(data)} visitor records from JSON to {'PostgreSQL' if USE_POSTGRES else 'SQLite'}")
            except Exception as exc:
                logger.error(f"JSON→DB migration failed for visitors: {exc}")
    finally:
        conn.close()


def save_downloads_to_disk():
    """Persist current downloads dict to the database (upsert all records)."""
    try:
        with downloads_lock:
            data = dict(downloads)
        with _db_lock:
            conn = _get_db()
            try:
                for did, d in data.items():
                    if USE_POSTGRES:
                        cur = conn.cursor()
                        cur.execute(
                            "INSERT INTO downloads (id, data) VALUES (%s, %s) "
                            "ON CONFLICT (id) DO UPDATE SET data = EXCLUDED.data",
                            (did, json.dumps(d, default=str)),
                        )
                    else:
                        conn.execute(
                            "INSERT OR REPLACE INTO downloads (id, data) VALUES (?, ?)",
                            (did, json.dumps(d, default=str)),
                        )
                conn.commit()
            finally:
                conn.close()
    except Exception as exc:
        logger.error(f"Could not persist downloads: {exc}")


def save_visitors_to_disk():
    """Persist visitor list to the database (replace all rows)."""
    try:
        with visitors_lock:
            data = list(visitors)
        with _db_lock:
            conn = _get_db()
            try:
                _execute(conn, "DELETE FROM visitors")
                for v in data:
                    _execute(
                        conn,
                        "INSERT INTO visitors (data) VALUES (?)",
                        (json.dumps(v, default=str),),
                    )
                conn.commit()
            finally:
                conn.close()
    except Exception as exc:
        logger.error(f"Could not persist visitors: {exc}")


# Shared timer for debounced visitor saves (avoids thread-per-visit churn)
_visitor_save_timer: threading.Timer | None = None
_visitor_save_timer_lock = threading.Lock()


def _schedule_visitor_save(delay: float = 5.0):
    """Schedule a visitor save in `delay` seconds, cancelling any pending save."""
    global _visitor_save_timer
    with _visitor_save_timer_lock:
        if _visitor_save_timer is not None:
            _visitor_save_timer.cancel()
        _visitor_save_timer = threading.Timer(delay, save_visitors_to_disk)
        _visitor_save_timer.daemon = True
        _visitor_save_timer.start()


def _parse_browser(ua: str) -> str:
    """Best-effort browser detection from User-Agent string."""
    ua_lower = ua.lower()
    if "edg/" in ua_lower or "edghtml" in ua_lower:
        return "Edge"
    if "chrome" in ua_lower:
        return "Chrome"
    if "firefox" in ua_lower:
        return "Firefox"
    if "safari" in ua_lower:
        return "Safari"
    if "msie" in ua_lower or "trident" in ua_lower:
        return "IE"
    return "Other"


def _parse_device(ua: str) -> str:
    """Best-effort device type detection from User-Agent string."""
    ua_lower = ua.lower()
    if "tablet" in ua_lower or "ipad" in ua_lower or "kindle" in ua_lower or "silk" in ua_lower:
        return "Tablet"
    if "android" in ua_lower and "mobi" not in ua_lower:
        return "Tablet"
    if "mobi" in ua_lower or "iphone" in ua_lower or "ipod" in ua_lower:
        return "Mobile"
    if ua_lower:
        return "Desktop"
    return "Other"


def _is_private_ip(ip: str) -> bool:
    """Return True if *ip* is a loopback, private, link-local or reserved address."""
    try:
        addr = ipaddress.ip_address(ip)
        return addr.is_private or addr.is_loopback or addr.is_link_local or addr.is_reserved
    except ValueError:
        return False


# Ordered list of headers that may carry the real client IP when the
# application sits behind a CDN, load-balancer, or reverse proxy.
# Priority (highest first):
#   1. CDN-specific single-IP headers  (Cloudflare, Akamai, Fastly, Fly.io)
#   2. Standard proxy headers          (X-Real-IP, X-Forwarded-For)
#   3. Less-common proxy headers       (X-Client-IP, X-Cluster-Client-IP)
_REAL_IP_HEADERS: tuple[str, ...] = (
    "CF-Connecting-IP",      # Cloudflare – always a single, verified IP
    "True-Client-IP",        # Akamai / Cloudflare Enterprise
    "Fastly-Client-IP",      # Fastly CDN
    "Fly-Client-IP",         # Fly.io
    "X-Real-IP",             # nginx / Apache mod_remoteip
    "X-Forwarded-For",       # Standard multi-proxy chain header
    "X-Client-IP",           # HAProxy / some load-balancers
    "X-Cluster-Client-IP",   # Rackspace Cloud Load Balancers
    "Forwarded",             # RFC 7239 – "for=<ip>" syntax
)

# CDN/proxy headers that carry the visitor's country ISO-2 code directly.
# Checked in priority order in _get_country_from_headers().
_CDN_GEO_HEADERS: tuple[str, ...] = (
    "CF-IPCountry",              # Cloudflare
    "CloudFront-Viewer-Country", # AWS CloudFront
    "X-Vercel-IP-Country",       # Vercel
    "X-Appengine-Country",       # Google App Engine
    "Fastly-Geo-Country",        # Fastly CDN
    "X-Country-Code",            # generic reverse-proxy
    "X-GeoIP-Country",           # generic reverse-proxy
    "X-Forwarded-Country",       # generic reverse-proxy
)


def _get_real_ip(request) -> str:
    """Return the best-guess real client IP from a FastAPI/Starlette request.

    When the application runs behind a reverse proxy or CDN the TCP peer
    visible at ``request.client.host`` is the proxy, not the actual visitor.
    This helper inspects a prioritised list of well-known forwarding headers
    to recover the original client address.

    The returned IP is always validated with :func:`ipaddress.ip_address` so
    that header injection cannot sneak in garbage values.  Private / loopback
    addresses in the ``X-Forwarded-For`` chain are skipped so that an internal
    hop does not obscure the public client IP.
    """
    for header in _REAL_IP_HEADERS:
        raw = request.headers.get(header, "").strip()
        if not raw:
            continue
        # RFC 7239 "Forwarded" header uses "for=<token>" syntax
        if header.lower() == "forwarded":
            for part in raw.split(","):
                for field in part.split(";"):
                    field = field.strip()
                    if field.lower().startswith("for="):
                        candidate = field[4:].strip().strip('"')
                        # RFC 7239 tokens may be "[IPv6]:port" or "IPv4:port"
                        if candidate.startswith("["):
                            # IPv6 in brackets – extract content, drop zone id
                            candidate = candidate.split("]")[0].lstrip("[").split("%")[0]
                        elif candidate.count(":") == 1:
                            # IPv4:port – keep only the IP part
                            candidate = candidate.split(":")[0]
                        # For bare IPv6 (no brackets) just drop any zone id
                        candidate = candidate.split("%")[0]
                        try:
                            addr = ipaddress.ip_address(candidate)
                            if not (addr.is_private or addr.is_loopback or addr.is_link_local):
                                return str(addr)
                        except ValueError:
                            continue
            continue
        # X-Forwarded-For may contain a comma-separated chain; take the first
        # non-private public IP (leftmost = original client in most setups).
        for candidate in raw.split(","):
            candidate = candidate.strip().split("%")[0]  # strip IPv6 zone id
            # Handle "ip:port" notation – keep only the IP part
            if candidate.startswith("["):
                candidate = candidate.split("]")[0].lstrip("[")
            elif candidate.count(":") == 1:  # IPv4:port
                candidate = candidate.split(":")[0]
            try:
                addr = ipaddress.ip_address(candidate)
                if not (addr.is_private or addr.is_loopback or addr.is_link_local):
                    return str(addr)
            except ValueError:
                continue
    # Fallback: direct TCP peer (works for direct connections / local dev)
    return request.client.host if request.client else "unknown"


# ── Accept-Language → country heuristic ──────────────────────────────────
# Maps common primary language codes to their most likely country ISO-2 code.
# Used only as a last-resort fallback when all geo-IP services fail.
_LANG_TO_COUNTRY: dict[str, str] = {
    "en": "US", "zh": "CN", "hi": "IN", "es": "ES", "fr": "FR",
    "ar": "SA", "bn": "BD", "pt": "BR", "ru": "RU", "ja": "JP",
    "de": "DE", "ko": "KR", "vi": "VN", "it": "IT", "tr": "TR",
    "pl": "PL", "uk": "UA", "nl": "NL", "th": "TH", "id": "ID",
    "sv": "SE", "da": "DK", "fi": "FI", "nb": "NO", "no": "NO",
    "cs": "CZ", "ro": "RO", "hu": "HU", "el": "GR", "he": "IL",
    "ms": "MY", "tl": "PH", "sw": "KE", "fa": "IR",
    "am": "ET", "az": "AZ", "be": "BY", "bg": "BG", "ca": "ES",
    "cy": "GB", "eu": "ES", "gl": "ES", "hr": "HR", "hy": "AM",
    "is": "IS", "ka": "GE", "km": "KH", "lo": "LA", "lt": "LT",
    "lv": "LV", "mk": "MK", "mn": "MN", "my": "MM", "ne": "NP",
    "si": "LK", "sk": "SK", "sl": "SI", "sq": "AL", "sr": "RS",
    "ta": "IN", "te": "IN", "ur": "PK", "uz": "UZ", "kk": "KZ",
    "ky": "KG", "tk": "TM", "tg": "TJ", "ps": "AF", "pa": "IN",
    "gu": "IN", "mr": "IN", "ml": "IN", "kn": "IN", "or": "IN",
    "so": "SO", "ha": "NG", "yo": "NG", "ig": "NG",
}

# Maps common IANA timezone prefixes/names to their most likely country ISO-2
# code. Used as a secondary heuristic before Accept-Language fallback.
_TZ_TO_COUNTRY: dict[str, str] = {
    # United States
    "America/New_York": "US", "America/Chicago": "US", "America/Denver": "US",
    "America/Los_Angeles": "US", "America/Anchorage": "US", "Pacific/Honolulu": "US",
    "America/Phoenix": "US", "America/Detroit": "US", "America/Indiana/Indianapolis": "US",
    "America/Indiana/Knox": "US", "America/Indiana/Marengo": "US",
    "America/Indiana/Petersburg": "US", "America/Indiana/Tell_City": "US",
    "America/Indiana/Vevay": "US", "America/Indiana/Vincennes": "US",
    "America/Indiana/Winamac": "US", "America/Kentucky/Louisville": "US",
    "America/Kentucky/Monticello": "US", "America/North_Dakota/Beulah": "US",
    "America/North_Dakota/Center": "US", "America/North_Dakota/New_Salem": "US",
    "America/Adak": "US", "America/Boise": "US", "America/Juneau": "US",
    "America/Metlakatla": "US", "America/Nome": "US", "America/Sitka": "US",
    "America/Yakutat": "US",
    # Canada
    "America/Toronto": "CA", "America/Vancouver": "CA", "America/Edmonton": "CA",
    "America/Winnipeg": "CA", "America/Halifax": "CA", "America/St_Johns": "CA",
    "America/Moncton": "CA", "America/Glace_Bay": "CA", "America/Goose_Bay": "CA",
    "America/Iqaluit": "CA", "America/Nipigon": "CA", "America/Pangnirtung": "CA",
    "America/Rainy_River": "CA", "America/Rankin_Inlet": "CA",
    "America/Regina": "CA", "America/Resolute": "CA", "America/Swift_Current": "CA",
    "America/Thunder_Bay": "CA", "America/Whitehorse": "CA", "America/Yellowknife": "CA",
    "America/Dawson": "CA", "America/Dawson_Creek": "CA", "America/Fort_Nelson": "CA",
    "America/Cambridge_Bay": "CA", "America/Creston": "CA",
    # Mexico
    "America/Mexico_City": "MX", "America/Cancun": "MX", "America/Tijuana": "MX",
    "America/Hermosillo": "MX", "America/Chihuahua": "MX", "America/Matamoros": "MX",
    "America/Mazatlan": "MX", "America/Merida": "MX", "America/Monterrey": "MX",
    "America/Ojinaga": "MX", "America/Santa_Isabel": "MX",
    # South America
    "America/Sao_Paulo": "BR", "America/Fortaleza": "BR", "America/Manaus": "BR",
    "America/Belem": "BR", "America/Boa_Vista": "BR", "America/Campo_Grande": "BR",
    "America/Cuiaba": "BR", "America/Eirunepe": "BR", "America/Maceio": "BR",
    "America/Noronha": "BR", "America/Porto_Velho": "BR", "America/Recife": "BR",
    "America/Rio_Branco": "BR", "America/Santarem": "BR",
    "America/Argentina/Buenos_Aires": "AR", "America/Argentina/Catamarca": "AR",
    "America/Argentina/Cordoba": "AR", "America/Argentina/Jujuy": "AR",
    "America/Argentina/La_Rioja": "AR", "America/Argentina/Mendoza": "AR",
    "America/Argentina/Rio_Gallegos": "AR", "America/Argentina/Salta": "AR",
    "America/Argentina/San_Juan": "AR", "America/Argentina/San_Luis": "AR",
    "America/Argentina/Tucuman": "AR", "America/Argentina/Ushuaia": "AR",
    "America/Bogota": "CO", "America/Lima": "PE", "America/Santiago": "CL",
    "America/Caracas": "VE", "America/Guayaquil": "EC", "America/La_Paz": "BO",
    "America/Asuncion": "PY", "America/Montevideo": "UY",
    "America/Cayenne": "GF", "America/Guyana": "GY", "America/Paramaribo": "SR",
    # Central America & Caribbean
    "America/Panama": "PA", "America/Costa_Rica": "CR", "America/Guatemala": "GT",
    "America/Havana": "CU", "America/Jamaica": "JM", "America/Port-au-Prince": "HT",
    "America/Santo_Domingo": "DO", "America/Tegucigalpa": "HN",
    "America/Managua": "NI", "America/El_Salvador": "SV", "America/Belize": "BZ",
    "America/Nassau": "BS", "America/Barbados": "BB", "America/Martinique": "MQ",
    "America/Port_of_Spain": "TT", "America/Aruba": "AW", "America/Curacao": "CW",
    "America/Puerto_Rico": "PR",
    # Europe
    "Europe/London": "GB", "Europe/Paris": "FR", "Europe/Berlin": "DE",
    "Europe/Madrid": "ES", "Europe/Rome": "IT", "Europe/Amsterdam": "NL",
    "Europe/Brussels": "BE", "Europe/Zurich": "CH", "Europe/Vienna": "AT",
    "Europe/Stockholm": "SE", "Europe/Oslo": "NO", "Europe/Copenhagen": "DK",
    "Europe/Helsinki": "FI", "Europe/Warsaw": "PL", "Europe/Prague": "CZ",
    "Europe/Budapest": "HU", "Europe/Bucharest": "RO", "Europe/Sofia": "BG",
    "Europe/Athens": "GR", "Europe/Istanbul": "TR", "Europe/Moscow": "RU",
    "Europe/Kiev": "UA", "Europe/Kyiv": "UA", "Europe/Lisbon": "PT",
    "Europe/Dublin": "IE", "Europe/Belgrade": "RS", "Europe/Zagreb": "HR",
    "Europe/Bratislava": "SK", "Europe/Ljubljana": "SI", "Europe/Tallinn": "EE",
    "Europe/Riga": "LV", "Europe/Vilnius": "LT", "Europe/Minsk": "BY",
    "Europe/Luxembourg": "LU", "Europe/Monaco": "MC", "Europe/Andorra": "AD",
    "Europe/Malta": "MT", "Europe/Nicosia": "CY", "Europe/Tirane": "AL",
    "Europe/Skopje": "MK", "Europe/Sarajevo": "BA", "Europe/Podgorica": "ME",
    "Europe/Kaliningrad": "RU", "Europe/Samara": "RU", "Europe/Saratov": "RU",
    "Europe/Ulyanovsk": "RU", "Europe/Volgograd": "RU",
    "Atlantic/Reykjavik": "IS", "Atlantic/Faroe": "FO",
    # Asia
    "Asia/Tokyo": "JP", "Asia/Seoul": "KR", "Asia/Shanghai": "CN",
    "Asia/Hong_Kong": "HK", "Asia/Taipei": "TW", "Asia/Singapore": "SG",
    "Asia/Kolkata": "IN", "Asia/Calcutta": "IN", "Asia/Karachi": "PK",
    "Asia/Dhaka": "BD", "Asia/Bangkok": "TH", "Asia/Ho_Chi_Minh": "VN",
    "Asia/Jakarta": "ID", "Asia/Manila": "PH", "Asia/Kuala_Lumpur": "MY",
    "Asia/Dubai": "AE", "Asia/Riyadh": "SA", "Asia/Tehran": "IR",
    "Asia/Baghdad": "IQ", "Asia/Jerusalem": "IL", "Asia/Beirut": "LB",
    "Asia/Colombo": "LK", "Asia/Kathmandu": "NP", "Asia/Yangon": "MM",
    "Asia/Almaty": "KZ", "Asia/Tashkent": "UZ", "Asia/Bishkek": "KG",
    "Asia/Kabul": "AF", "Asia/Yerevan": "AM", "Asia/Baku": "AZ",
    "Asia/Tbilisi": "GE", "Asia/Ashgabat": "TM", "Asia/Dushanbe": "TJ",
    "Asia/Katmandu": "NP", "Asia/Rangoon": "MM", "Asia/Saigon": "VN",
    "Asia/Phnom_Penh": "KH", "Asia/Vientiane": "LA", "Asia/Brunei": "BN",
    "Asia/Makassar": "ID", "Asia/Jayapura": "ID", "Asia/Pontianak": "ID",
    "Asia/Kuching": "MY", "Asia/Macau": "MO", "Asia/Ulaanbaatar": "MN",
    "Asia/Choibalsan": "MN", "Asia/Hovd": "MN", "Asia/Aden": "YE",
    "Asia/Kuwait": "KW", "Asia/Bahrain": "BH", "Asia/Qatar": "QA",
    "Asia/Muscat": "OM", "Asia/Nicosia": "CY", "Asia/Amman": "JO",
    "Asia/Damascus": "SY", "Asia/Gaza": "PS", "Asia/Hebron": "PS",
    "Asia/Pyongyang": "KP", "Asia/Urumqi": "CN", "Asia/Chongqing": "CN",
    "Asia/Harbin": "CN", "Asia/Kashgar": "CN",
    "Asia/Novosibirsk": "RU", "Asia/Omsk": "RU", "Asia/Krasnoyarsk": "RU",
    "Asia/Irkutsk": "RU", "Asia/Yakutsk": "RU", "Asia/Vladivostok": "RU",
    "Asia/Magadan": "RU", "Asia/Kamchatka": "RU", "Asia/Anadyr": "RU",
    "Asia/Sakhalin": "RU", "Asia/Srednekolymsk": "RU", "Asia/Chita": "RU",
    "Asia/Khandyga": "RU", "Asia/Ust-Nera": "RU", "Asia/Yekaterinburg": "RU",
    # Africa
    "Africa/Cairo": "EG", "Africa/Lagos": "NG", "Africa/Nairobi": "KE",
    "Africa/Johannesburg": "ZA", "Africa/Casablanca": "MA", "Africa/Algiers": "DZ",
    "Africa/Tunis": "TN", "Africa/Accra": "GH", "Africa/Addis_Ababa": "ET",
    "Africa/Dar_es_Salaam": "TZ", "Africa/Kampala": "UG", "Africa/Khartoum": "SD",
    "Africa/Abidjan": "CI", "Africa/Dakar": "SN", "Africa/Maputo": "MZ",
    "Africa/Lusaka": "ZM", "Africa/Harare": "ZW", "Africa/Tripoli": "LY",
    "Africa/Luanda": "AO", "Africa/Douala": "CM", "Africa/Brazzaville": "CG",
    "Africa/Kinshasa": "CD", "Africa/Libreville": "GA", "Africa/Bangui": "CF",
    "Africa/Ndjamena": "TD", "Africa/Malabo": "GQ", "Africa/Windhoek": "NA",
    "Africa/Gaborone": "BW", "Africa/Maseru": "LS", "Africa/Mbabane": "SZ",
    "Africa/Bujumbura": "BI", "Africa/Kigali": "RW", "Africa/Djibouti": "DJ",
    "Africa/Mogadishu": "SO", "Africa/Asmara": "ER", "Africa/Juba": "SS",
    "Africa/Bamako": "ML", "Africa/Conakry": "GN", "Africa/Bissau": "GW",
    "Africa/Freetown": "SL", "Africa/Monrovia": "LR", "Africa/Ouagadougou": "BF",
    "Africa/Lome": "TG", "Africa/Porto-Novo": "BJ", "Africa/Niamey": "NE",
    "Africa/Nouakchott": "MR", "Africa/El_Aaiun": "EH", "Africa/Ceuta": "ES",
    # Oceania / Pacific
    "Australia/Sydney": "AU", "Australia/Melbourne": "AU", "Australia/Brisbane": "AU",
    "Australia/Perth": "AU", "Australia/Adelaide": "AU", "Australia/Darwin": "AU",
    "Australia/Hobart": "AU", "Australia/Lord_Howe": "AU",
    "Australia/Broken_Hill": "AU", "Australia/Currie": "AU",
    "Australia/Eucla": "AU", "Australia/Lindeman": "AU",
    "Pacific/Auckland": "NZ", "Pacific/Chatham": "NZ",
    "Pacific/Fiji": "FJ", "Pacific/Guam": "GU", "Pacific/Port_Moresby": "PG",
    "Pacific/Bougainville": "PG", "Pacific/Noumea": "NC",
    "Pacific/Guadalcanal": "SB", "Pacific/Efate": "VU",
    "Pacific/Tarawa": "KI", "Pacific/Enderbury": "KI", "Pacific/Kiritimati": "KI",
    "Pacific/Funafuti": "TV", "Pacific/Majuro": "MH", "Pacific/Kwajalein": "MH",
    "Pacific/Nauru": "NR", "Pacific/Palau": "PW", "Pacific/Yap": "FM",
    "Pacific/Pohnpei": "FM", "Pacific/Kosrae": "FM", "Pacific/Chuuk": "FM",
    "Pacific/Honolulu": "US", "Pacific/Pago_Pago": "AS", "Pacific/Midway": "UM",
    "Pacific/Tongatapu": "TO", "Pacific/Apia": "WS", "Pacific/Niue": "NU",
    "Pacific/Rarotonga": "CK", "Pacific/Tahiti": "PF", "Pacific/Gambier": "PF",
    "Pacific/Marquesas": "PF", "Pacific/Pitcairn": "PN",
    "Pacific/Easter": "CL", "Pacific/Galapagos": "EC",
    "Indian/Maldives": "MV", "Indian/Mauritius": "MU", "Indian/Reunion": "RE",
    "Indian/Mayotte": "YT", "Indian/Comoro": "KM", "Indian/Antananarivo": "MG",
    "Indian/Mahe": "SC", "Indian/Kerguelen": "TF", "Indian/Cocos": "CC",
    "Indian/Christmas": "CX",
}


def _country_from_accept_language(accept_lang: str) -> tuple[str, str]:
    """Best-effort country guess from an Accept-Language header value.

    Returns (country_name, iso2_code) or ("", "") when no guess can be made.
    Handles formats like "en-US,en;q=0.9" or "en,en-GB;q=0.9".

    Accuracy improvement: scans ALL language tags (not just the first) for an
    explicit region subtag (e.g. "en-GB") before falling back to the primary
    language code.  This fixes cases where a browser sends the bare language
    first (e.g. "en") followed by the region-qualified tag (e.g. "en-GB"),
    which previously caused UK visitors to be reported as United States.
    """
    if not accept_lang:
        return "", ""
    # Extract all tags, stripped of quality weights
    tags = [t.split(";")[0].strip() for t in accept_lang.split(",")]
    # First pass: scan every tag for an explicit two-letter region subtag
    for tag in tags:
        parts = tag.replace("_", "-").split("-")
        if len(parts) >= 2:
            region = parts[1].upper()
            if len(region) == 2 and region.isalpha():
                name = _ISO2_TO_NAME.get(region)
                if name:
                    return name, region
    # Second pass: fall back to primary language code of the first tag
    first_parts = tags[0].replace("_", "-").split("-") if tags else []
    lang = first_parts[0].lower() if first_parts else ""
    code = _LANG_TO_COUNTRY.get(lang, "")
    if code:
        return _ISO2_TO_NAME.get(code, code), code
    return "", ""


def _lookup_country_async(ip: str, accept_language: str = "", client_tz: str = ""):
    """Resolve an IP to its country (and city/region when available).

    Lookup order:
    1. Local GeoIP database (DB-IP City Lite MMDB – most accurate, offline)
    2. ip-api.com / ipinfo.io / ipwhois.app / ipapi.co / api.country.is
       / reallyfreegeoip.org / freeipapi.com / ipapi.com (ipstack)
    3. Timezone heuristic via worldtimeapi.org
    4. Browser-reported IANA timezone (client_tz) – more precise than Accept-Language
    5. Accept-Language header heuristic (last resort)

    Args:
        ip: The IP address to geo-locate.
        accept_language: Optional Accept-Language header value used as a
            last-resort heuristic when all geo-IP services fail.
        client_tz: Optional IANA timezone string reported by the browser
            (``Intl.DateTimeFormat().resolvedOptions().timeZone``).  Used as a
            high-confidence fallback when all network-based lookups fail.
    """
    if ip in ip_country_cache:
        return

    # Private / loopback addresses cannot be geo-located
    if _is_private_ip(ip):
        ip_country_cache[ip] = {"country": "Local", "code": "", "city": "", "region": ""}
        with visitors_lock:
            for v in visitors[-200:]:
                if v.get("ip") == ip and not v.get("country"):
                    v["country"] = "Local"
                    v["country_code"] = ""
                    v["city"] = ""
                    v["region"] = ""
        with downloads_lock:
            for d in downloads.values():
                if d.get("ip") == ip and not d.get("country"):
                    d["country"] = "Local"
                    d["country_code"] = ""
                    d["city"] = ""
                    d["region"] = ""
        _schedule_visitor_save()
        threading.Thread(target=save_downloads_to_disk, daemon=True).start()
        return

    country, code, city, region = "Unknown", "", "", ""

    # --- Primary: Local GeoIP database (accurate, offline, no rate limits) ---
    if _geoip_reader is not None:
        try:
            result = _geoip_reader.get(ip)
            if result:
                c = result.get("country", {})
                iso = (c.get("iso_code") or "").upper()
                if iso and len(iso) == 2:
                    code = iso
                    # Prefer our canonical dictionary name for consistency
                    country = _ISO2_TO_NAME.get(code) or c.get("names", {}).get("en") or code
                    city_data = result.get("city", {})
                    city = city_data.get("names", {}).get("en", "")
                    subdivisions = result.get("subdivisions", [])
                    if subdivisions:
                        region = subdivisions[0].get("names", {}).get("en", "")
        except Exception:
            pass

    # --- Service 1: ip-api.com (free, no key) ---
    if country == "Unknown":
        try:
            with urllib.request.urlopen(
                f"https://ip-api.com/json/{ip}?fields=status,country,countryCode,city,regionName", timeout=5
            ) as resp:
                data = json.loads(resp.read())
            if data.get("status") == "success" and data.get("countryCode"):
                code = data["countryCode"].upper()
                # Normalise name through our dictionary for consistency
                country = _ISO2_TO_NAME.get(code, data.get("country", code))
                city = data.get("city", "")
                region = data.get("regionName", "")
        except Exception:
            pass

    # --- Service 2: ipinfo.io (fallback) ---
    if country == "Unknown":
        try:
            with urllib.request.urlopen(
                f"https://ipinfo.io/{ip}/json", timeout=5
            ) as resp:
                data = json.loads(resp.read())
            if data.get("country") and not data.get("bogon"):
                code = data["country"].upper()  # ipinfo returns 2-letter code in "country"
                country = _ISO2_TO_NAME.get(code, code)
                city = data.get("city", "")
                region = data.get("region", "")
        except Exception:
            pass

    # --- Service 3: ipwhois.app (second fallback) ---
    if country == "Unknown":
        try:
            with urllib.request.urlopen(
                f"https://ipwhois.app/json/{ip}?objects=country,country_code,city,region", timeout=5
            ) as resp:
                data = json.loads(resp.read())
            if data.get("country_code"):
                code = data["country_code"].upper()
                country = _ISO2_TO_NAME.get(code, data.get("country", code))
                city = data.get("city", "")
                region = data.get("region", "")
        except Exception:
            pass

    # --- Service 4: ipapi.co (third fallback) ---
    if country == "Unknown":
        try:
            req = urllib.request.Request(
                f"https://ipapi.co/{ip}/json/",
                headers={"User-Agent": "Mozilla/5.0"},
            )
            with urllib.request.urlopen(req, timeout=5) as resp:
                data = json.loads(resp.read())
            if data.get("country_code") and not data.get("error"):
                code = data["country_code"].upper()
                country = _ISO2_TO_NAME.get(code, data.get("country_name", code))
                city = data.get("city", "")
                region = data.get("region", "")
        except Exception:
            pass

    # --- Service 5: api.country.is (lightweight, returns ISO code only) ---
    if country == "Unknown":
        try:
            with urllib.request.urlopen(
                f"https://api.country.is/{ip}", timeout=5
            ) as resp:
                data = json.loads(resp.read())
            iso = (data.get("country") or "").upper()
            if iso and len(iso) == 2:
                code = iso
                country = _ISO2_TO_NAME.get(code, code)
        except Exception:
            pass

    # --- Service 6: reallyfreegeoip.org (free, no key required) ---
    if country == "Unknown":
        try:
            with urllib.request.urlopen(
                f"https://reallyfreegeoip.org/json/{ip}", timeout=5
            ) as resp:
                data = json.loads(resp.read())
            iso = (data.get("country_code") or "").upper()
            if iso and len(iso) == 2:
                code = iso
                country = _ISO2_TO_NAME.get(code, data.get("country_name", code))
                city = data.get("city", "") or city
                region = data.get("region_name", "") or region
        except Exception:
            pass

    # --- Service 7: freeipapi.com (free, returns country + timezone) ---
    if country == "Unknown":
        try:
            with urllib.request.urlopen(
                f"https://freeipapi.com/api/json/{ip}", timeout=5
            ) as resp:
                data = json.loads(resp.read())
            iso = (data.get("countryCode") or "").upper()
            if iso and len(iso) == 2:
                code = iso
                country = _ISO2_TO_NAME.get(code, data.get("countryName", code))
                city = data.get("cityName", "") or city
                region = data.get("regionName", "") or region
        except Exception:
            pass

    # --- Service 8: ipapi.com / ipstack (requires IPAPI_ACCESS_KEY env var) ---
    if country == "Unknown" and IPAPI_ACCESS_KEY:
        try:
            req = urllib.request.Request(
                f"https://api.ipapi.com/{ip}?access_key={IPAPI_ACCESS_KEY}&output=json",
                headers={"User-Agent": "Mozilla/5.0"},
            )
            with urllib.request.urlopen(req, timeout=5) as resp:
                data = json.loads(resp.read())
            iso = (data.get("country_code") or "").upper()
            if iso and len(iso) == 2:
                code = iso
                country = _ISO2_TO_NAME.get(code, data.get("country_name", code))
                city = data.get("city", "") or city
                region = data.get("region_name", "") or region
        except Exception:
            pass

    # --- Timezone-based heuristic (maps IANA timezone to likely country) ---
    if country == "Unknown" and not code:
        try:
            # Try fetching timezone from worldtimeapi.org and map to a country
            with urllib.request.urlopen(
                f"https://worldtimeapi.org/api/ip/{ip}", timeout=5
            ) as resp:
                data = json.loads(resp.read())
            tz = data.get("timezone", "")
            if tz:
                tz_code = _TZ_TO_COUNTRY.get(tz, "")
                if tz_code and tz_code in _ISO2_TO_NAME:
                    code = tz_code
                    country = _ISO2_TO_NAME[code]
        except Exception:
            pass

    # --- Last-resort: Accept-Language header heuristic ---
    # The browser language region subtag (e.g. en-US, en-GB) reflects locale
    # preference, not geographic location, so it is imprecise.  However it is
    # better than leaving the country as "Unknown" when every geo-IP service
    # has failed.
    if country == "Unknown" and not code and client_tz:
        tz_code = _TZ_TO_COUNTRY.get(client_tz, "")
        if tz_code and tz_code in _ISO2_TO_NAME:
            code = tz_code
            country = _ISO2_TO_NAME[code]

    if country == "Unknown" and not code and accept_language:
        al_country, al_code = _country_from_accept_language(accept_language)
        if al_country and al_code:
            code = al_code
            country = al_country

    # Final normalisation: validate the ISO code and ensure the country name
    # matches our canonical dictionary so every record uses the exact same
    # spelling regardless of which lookup service resolved it.
    if code:
        code = code.upper()
        if len(code) == 2 and code in _ISO2_TO_NAME:
            country = _ISO2_TO_NAME[code]
        elif len(code) != 2:
            code = ""  # discard invalid codes

    ip_country_cache[ip] = {"country": country, "code": code, "city": city, "region": region}
    # Back-fill any visitor records that are waiting for this IP's country
    with visitors_lock:
        for v in visitors[-200:]:
            if v.get("ip") == ip and not v.get("country"):
                v["country"] = country
                v["country_code"] = code
                v["city"] = city
                v["region"] = region
    # Back-fill any download records that are waiting for this IP's country
    with downloads_lock:
        for d in downloads.values():
            if d.get("ip") == ip and not d.get("country"):
                d["country"] = country
                d["country_code"] = code
                d["city"] = city
                d["region"] = region
    _schedule_visitor_save()
    # Persist updated download records so country is not lost on restart
    threading.Thread(target=save_downloads_to_disk, daemon=True).start()


def _get_country_from_headers(req) -> tuple[str, str]:
    """Extract country from CDN/proxy headers before hitting external APIs.

    Supported headers (in priority order):
    - CF-IPCountry            (Cloudflare)
    - CloudFront-Viewer-Country (AWS CloudFront)
    - X-Vercel-IP-Country     (Vercel)
    - X-Appengine-Country     (Google App Engine)
    - Fastly-Geo-Country      (Fastly CDN)
    - X-Country-Code          (generic)
    - X-GeoIP-Country         (generic)
    - X-Forwarded-Country     (generic)

    Returns (country_name, iso2_code) or ("", "") when not found.
    """
    for header in _CDN_GEO_HEADERS:
        code = req.headers.get(header, "").strip().upper()
        if len(code) == 2 and code.isalpha() and code not in ("XX", "T1"):
            return _ISO2_TO_NAME.get(code, code), code
    return "", ""


# =========================================================
# RATE LIMITING DECORATOR
# =========================================================

def rate_limit(max_per_ip=Config.MAX_DOWNLOADS_PER_IP):
    """Rate limiting decorator for FastAPI endpoints.

    The decorated function must accept ``request: Request`` as its first
    positional or keyword argument so that the IP address can be extracted.
    """
    def decorator(f):
        @wraps(f)
        async def wrapped(*args, request: Request, **kwargs):
            ip = _get_real_ip(request)
            with downloads_lock:
                count = ip_download_count.get(ip, 0)
                if count >= max_per_ip:
                    return JSONResponse(
                        {"error": f"Rate limit exceeded. Maximum {max_per_ip} concurrent downloads per IP."},
                        status_code=429,
                    )
                ip_download_count[ip] = count + 1

            try:
                return await f(*args, request=request, **kwargs)
            finally:
                with downloads_lock:
                    ip_download_count[ip] = ip_download_count.get(ip, 1) - 1
                    if ip_download_count[ip] <= 0:
                        ip_download_count.pop(ip, None)
        return wrapped
    return decorator

# =========================================================
# ADMIN AUTHENTICATION DECORATOR
# =========================================================

def admin_required(f):
    """Decorator that requires admin login for the wrapped view.

    The decorated function must accept ``request: Request`` so that the
    session and accept headers can be inspected.
    """
    @wraps(f)
    async def wrapped(*args, request: Request, **kwargs):
        if not request.session.get("admin_logged_in"):
            if request.method != "GET":
                return JSONResponse({"error": "Authentication required"}, status_code=401)
            accept = request.headers.get("accept", "text/html")
            if "application/json" in accept and "text/html" not in accept:
                return JSONResponse({"error": "Authentication required"}, status_code=401)
            login_url = "/admin/login?next=" + str(request.url)
            return RedirectResponse(url=login_url, status_code=302)
        return await f(*args, request=request, **kwargs)
    return wrapped


def check_yt_dlp():
    """Check if yt-dlp is installed and accessible"""
    try:
        logger.info(f"OK: yt-dlp version: {yt_dlp.version.__version__}")
        return True
    except Exception as e:
        logger.error(f"ERROR: Error checking yt-dlp: {e}")
        return False

def check_ffmpeg():
    """Check if ffmpeg is installed"""
    ffmpeg_path = shutil.which('ffmpeg')
    if ffmpeg_path:
        logger.info(f"OK: ffmpeg found at: {ffmpeg_path}")
        return True
    else:
        logger.warning("WARNING: ffmpeg not found - some formats may not work")
        return False


def check_youtube_connectivity() -> dict:
    """Probe YouTube with yt-dlp to detect active bot-detection blocks.

    Returns a dict with keys:
      * ``reachable`` (bool) – True if the probe succeeded.
      * ``bot_detected`` (bool) – True if a bot/auth-detection error fired.
      * ``message`` (str) – Human-readable summary.

    The probe delegates to :func:`_get_cookieless_extractor_args` for its
    extractor args, using ``android_vr``, ``web_embedded``, and ``tv`` player
    clients.  ``android_vr`` requires no JS runtime and no PO Token, making it
    the most reliable client for an unauthenticated server environment.
    ``web_embedded`` and ``tv`` are included as higher-quality fallbacks when a
    JS runtime is available.  ``mweb`` is intentionally omitted: as of
    yt-dlp 2026.3.x it requires a GVS PO Token, so all its formats would be
    skipped and it would produce a spurious user-visible warning.
    """
    # A short, well-known public-domain video used solely as a reachability probe.
    _PROBE_URL = "https://www.youtube.com/watch?v=aqz-KE-bpKQ"

    ydl_opts = {
        "quiet": True,
        "no_warnings": True,
        "skip_download": True,
        "extract_flat": False,
        "extractor_args": _get_cookieless_extractor_args(),
    }
    try:
        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            ydl.extract_info(_PROBE_URL, download=False)
        logger.info("YouTube connectivity check: reachable")
        return {"reachable": True, "bot_detected": False, "message": "YouTube is reachable"}
    except Exception as exc:
        error_msg = str(exc)
        if _is_auth_error(error_msg):
            logger.warning(f"YouTube connectivity check: bot-detection active — {error_msg}")
            return {
                "reachable": False,
                "bot_detected": True,
                "message": "YouTube bot-detection is active. Downloads may fail.",
            }
        logger.error(f"YouTube connectivity check: unexpected error — {error_msg}")
        return {
            "reachable": False,
            "bot_detected": False,
            "message": f"YouTube unreachable: {error_msg}",
        }

# =========================================================
# HELPER FUNCTIONS
# =========================================================

def safe_filename(name: str) -> str:
    """Create a safe filename"""
    # Remove invalid characters
    name = "".join(c for c in name if c.isalnum() or c in (" ", "-", "_", ".", "(", ")"))
    # Limit length
    if len(name) > 100:
        name = name[:100]
    return name.strip()

def format_size(size_bytes):
    """Format file size"""
    if size_bytes == 0:
        return "0 B"
    size_names = ["B", "KB", "MB", "GB", "TB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    return f"{size_bytes:.2f} {size_names[i]}"

def get_ssl_env() -> dict:
    """Return environment with SSL certificate vars"""
    env = os.environ.copy()
    env["SSL_CERT_FILE"] = certifi.where()
    env["REQUESTS_CA_BUNDLE"] = certifi.where()
    return env

def _get_yt_extractor_args() -> dict:
    """Build YouTube extractor args letting yt-dlp pick the best player clients.

    ⚠️  DO NOT REMOVE ``"default"`` from ``player_client``.
    ``"default"`` lets yt-dlp automatically switch between:
      • unauthenticated defaults: ``android_vr``, ``web_safari``
      • authenticated defaults:   ``tv_downgraded``, ``web_safari``
    Removing it causes YouTube authentication errors (investigated in PR #78).

    ``web_embedded`` and ``tv`` are added alongside ``"default"`` because
    yt-dlp 2026.3.13 removed the ``web`` client from its unauthenticated
    defaults (``web`` and ``web_safari`` now require PO tokens for HTTPS/DASH
    streams).  ``web_embedded`` and ``tv`` have no PO-token requirement, support
    cookies, and work with or without a JS runtime — providing reliable fallback
    clients for unauthenticated sessions where ``web_safari`` would otherwise
    fail without a POT provider.

    ``mweb`` is intentionally excluded: as of yt-dlp 2026.3.x it requires a
    GVS PO Token for all streaming protocols.  Without a POT provider every
    ``mweb`` format is skipped and yt-dlp emits a spurious user-visible warning,
    making ``mweb`` useless in a server environment with no POT provider.

    See https://github.com/yt-dlp/yt-dlp/wiki/Extractors#youtube for details.
    """
    # ⚠️ DO NOT REMOVE "default" — see docstring above and PR #78
    # web_embedded + tv: no POT required, SUPPORTS_COOKIES=True — reliable fallbacks
    # mweb omitted: requires GVS PO Token in yt-dlp 2026.3.x, all formats skipped without POT provider
    args: dict = {"player_client": ["default", "web_embedded", "tv"]}
    return {"youtube": args}


def _get_cookieless_extractor_args() -> dict:
    """Build YouTube extractor args using only clients that work without authentication.

    ``android_vr`` is the most reliable POT-free, JS-free client — it requires
    neither a PO Token nor a JavaScript runtime and can always fetch publicly
    available videos without cookies.

    ``web_embedded`` and ``tv`` are added as higher-quality fallbacks: both
    require no PO Token and support cookies, so they provide better format
    selection when a JS runtime (Node.js / Deno) is present.

    ``mweb`` is intentionally excluded: as of yt-dlp 2026.3.x it requires a
    GVS PO Token for all streaming protocols.  Without a POT provider every
    ``mweb`` format is skipped, and yt-dlp emits a user-visible warning,
    making ``mweb`` worse than useless in a cookieless fallback path.
    """
    return {"youtube": {"player_client": ["android_vr", "web_embedded", "tv"]}}


def _get_cookie_opts() -> dict:
    """Return yt-dlp ``cookiefile`` option when a valid cookies file exists.

    The cookies file path is controlled by the ``COOKIES_FILE`` env-var and
    defaults to ``data/cookies.txt``.  If the file doesn't exist an empty dict
    is returned so callers can simply unpack it into their ``ydl_opts``.
    """
    if os.path.isfile(COOKIES_FILE):
        return {"cookiefile": COOKIES_FILE}
    return {}

# Bot-detection and login/auth patterns yt-dlp may emit
_AUTH_PATTERNS = (
    "sign in to confirm",
    "confirm you're not a bot",
    "login required",
    "this video requires login",
    "please sign in",
    "sign in to view",
    "session has been invalidated",
    "required to log in",
    "use --cookies",
    "use --cookies-from-browser",
    "requires authentication",
    "not authenticated",
    # Captcha challenge — YouTube bot-detection via CAPTCHA (yt-dlp ≥ 2026.3.x)
    "captcha challenge",
    "requiring a captcha",
)

# DRM protection patterns emitted by yt-dlp
_DRM_PATTERNS = (
    "drm protected",
    "this video is drm",
    "drm-protected",
    "widevine",
    "playready",
    "fairplay",
    "is drm",
)

# HTTP 403 Forbidden patterns — CDN-level access denial during video data download
_HTTP_FORBIDDEN_PATTERNS = (
    "http error 403",
    "403: forbidden",
    "403 forbidden",
    "unable to download video data",
)


def _is_auth_error(error_msg: str) -> bool:
    """Return ``True`` if *error_msg* matches a known authentication/bot-detection pattern."""
    lower = error_msg.lower()
    if any(p in lower for p in _AUTH_PATTERNS):
        return True
    # Also catch cookie-specific failures (expired / invalid / missing cookies)
    if "cookie" in lower and any(w in lower for w in ("invalid", "expired", "missing", "rejected")):
        return True
    return False


def _is_drm_error(error_msg: str) -> bool:
    """Return ``True`` if *error_msg* indicates DRM-protected content.

    DRM (Digital Rights Management) errors cannot be resolved by switching
    player clients or using alternative authentication.  When detected, the
    downloader should try alternative tools (gallery-dl, you-get, streamlink)
    which may be able to fetch DRM-free versions of the same content.
    """
    lower = error_msg.lower()
    return any(p in lower for p in _DRM_PATTERNS)


def _is_http_forbidden_error(error_msg: str) -> bool:
    """Return ``True`` if *error_msg* indicates an HTTP 403 Forbidden error.

    YouTube CDN servers sometimes return HTTP 403 when downloading video data
    even though extraction succeeded.  This can happen when:
    - The CDN stream URL's embedded token has expired between extraction and download
    - The server IP is temporarily rate-limited or geo-blocked at the CDN level
    - The selected player client produces stream URLs that YouTube CDN refuses

    When detected, the downloader retries with alternative player clients and,
    if still failing, falls back to alternative download tools.
    """
    lower = error_msg.lower()
    return any(p in lower for p in _HTTP_FORBIDDEN_PATTERNS)


def _is_auth_or_forbidden_error(error_msg: str) -> bool:
    """Return ``True`` when *error_msg* is a bot-detection, auth, or HTTP 403 error.

    These errors will not resolve on retry, so exponential backoff should be
    skipped and the cookieless / alternative-tool fallback path reached
    immediately.
    """
    return _is_auth_error(error_msg) or _is_http_forbidden_error(error_msg)


# ── Alternative tool fallback ─────────────────────────────────────────────────

# Media file extensions that alternative tools may produce
_ALT_MEDIA_EXTS: tuple[str, ...] = (
    ".mp4", ".webm", ".mkv", ".avi", ".mov", ".ts", ".3gp",
    ".mp3", ".m4a", ".ogg", ".wav", ".opus", ".flac", ".aac",
    ".flv", ".wmv", ".m4v",
)

# Alternative download tools tried in order.  Each entry is a list of tokens
# that will be passed to subprocess.Popen; ``{url}`` and ``{out}`` are
# replaced at call time with the URL and output directory respectively.
# Tools absent from PATH are silently skipped.
_ALTERNATIVE_TOOL_COMMANDS: list[list[str]] = [
    ["gallery-dl", "--dest", "{out}", "{url}"],
    ["you-get", "--output-dir", "{out}", "{url}"],
    ["streamlink", "--output", "{out}/stream.ts", "{url}", "best"],
]

# Polling interval (seconds) between cancellation checks during alternative tool runs
_CANCELLATION_CHECK_INTERVAL_SECONDS = 0.5
_GENTLE_FAILURE_MESSAGE = (
    "This video could not be downloaded. "
    "It may be DRM-protected, region-restricted, or temporarily unavailable. "
    "Please try again later or try a different video."
)

# ── ssyoutube / savefrom.net fallback ────────────────────────────────────────

# savefrom.net API (powers ssyoutube.com) used as a last-resort fallback when
# all yt-dlp strategies and alternative CLI tools have failed for YouTube URLs.
_SSYOUTUBE_API_URL = "https://worker.sf-tools.com/savefrom.php"

# Ordered list of video quality labels to try (highest quality first)
_SSYOUTUBE_QUALITY_PREFERENCE = ("1080", "720", "480", "360", "240", "144")

# Regex patterns for extracting a YouTube video ID from common URL forms
_YOUTUBE_URL_RE = re.compile(
    r"(?:https?://)?(?:www\.)?(?:youtube\.com|youtu\.be|m\.youtube\.com)"
    r"/(?:watch\?v=|embed/|shorts/|v/)?([A-Za-z0-9_-]{11})"
)


def _is_youtube_url(url: str) -> bool:
    """Return ``True`` when *url* is a YouTube URL."""
    return bool(_YOUTUBE_URL_RE.search(url))


def _extract_youtube_video_id(url: str) -> str | None:
    """Return the 11-character YouTube video ID embedded in *url*, or ``None``."""
    match = _YOUTUBE_URL_RE.search(url)
    return match.group(1) if match else None


def _try_ssyoutube_download(
    url: str,
    output_dir: str,
    download_id: str | None = None,
) -> str | None:
    """Download a YouTube video via the ssyoutube / savefrom.net API.

    This function replicates the logic used by ``ssyoutube.com`` to bypass
    YouTube download restrictions: it sends the video URL to the savefrom.net
    worker API, parses the returned JSON for the best available MP4 download
    link, and streams the file to *output_dir*.

    Returns the absolute path of the downloaded file on success, or ``None``
    if the URL is not a YouTube URL, the API returns no usable links, or any
    network / IO error occurs.  All failures are logged at INFO level so they
    are visible in the server log without being alarming to operators.

    The optional *download_id* is checked for cancellation between the API
    request and each streaming chunk so that a queued cancel is honoured
    promptly.
    """
    if not _is_youtube_url(url):
        return None

    video_id = _extract_youtube_video_id(url)
    if not video_id:
        return None

    logger.info("Trying ssyoutube fallback for video ID %s", video_id)

    # ── Step 1: call the savefrom.net API to obtain download links ──────────
    try:
        post_data = urllib.parse.urlencode({"sf_url": url}).encode()
        req = urllib.request.Request(
            _SSYOUTUBE_API_URL,
            data=post_data,
            headers={
                "User-Agent": _CHROME_UA,
                "Referer": "https://ssyoutube.com/",
                "Origin": "https://ssyoutube.com",
                "Content-Type": "application/x-www-form-urlencoded",
            },
            method="POST",
        )
        ctx = __import__("ssl").create_default_context(cafile=certifi.where())
        with urllib.request.urlopen(req, timeout=30, context=ctx) as resp:
            raw = resp.read().decode(errors="replace")
        data: dict = json.loads(raw)
    except Exception as exc:
        logger.info("ssyoutube API request failed for %s: %s", video_id, exc)
        return None

    # ── Step 2: pick the best quality download URL ──────────────────────────
    # The API returns a dict keyed by quality label, each value a list of
    # format dicts: {"url": "...", "ext": "mp4"|"webm", "no_audio": bool, ...}
    url_map: dict = data.get("url", {})
    if not isinstance(url_map, dict) or not url_map:
        logger.info("ssyoutube API returned no download URLs for %s", video_id)
        return None

    download_url: str | None = None
    chosen_ext = "mp4"
    for quality in _SSYOUTUBE_QUALITY_PREFERENCE:
        formats = url_map.get(quality, [])
        if not isinstance(formats, list):
            continue
        for fmt in formats:
            if not isinstance(fmt, dict):
                continue
            fmt_url = fmt.get("url", "")
            fmt_ext = fmt.get("ext", "mp4")
            # Prefer MP4 with audio; skip audio-only or video-only streams
            if fmt_url and not fmt.get("no_audio") and fmt_ext in ("mp4", "webm"):
                download_url = fmt_url
                chosen_ext = fmt_ext
                break
        if download_url:
            break

    if not download_url:
        logger.info("ssyoutube API returned no playable format for %s", video_id)
        return None

    # ── Step 3: check for cancellation before downloading ───────────────────
    if download_id is not None:
        with downloads_lock:
            if downloads.get(download_id, {}).get("status") == "cancelled":
                return None

    # ── Step 4: stream the file to disk ──────────────────────────────────────
    output_path = os.path.join(output_dir, f"{video_id}.{chosen_ext}")
    try:
        dl_req = urllib.request.Request(
            download_url,
            headers={"User-Agent": _CHROME_UA},
        )
        ctx = __import__("ssl").create_default_context(cafile=certifi.where())
        with urllib.request.urlopen(dl_req, timeout=120, context=ctx) as resp, \
                open(output_path, "wb") as out_file:
            while True:
                chunk = resp.read(65536)
                if not chunk:
                    break
                # Honour cancellation between chunks
                if download_id is not None:
                    with downloads_lock:
                        if downloads.get(download_id, {}).get("status") == "cancelled":
                            out_file.close()
                            try:
                                os.remove(output_path)
                            except OSError:
                                pass
                            return None
                out_file.write(chunk)
    except Exception as exc:
        logger.info("ssyoutube file download failed for %s: %s", video_id, exc)
        try:
            os.remove(output_path)
        except OSError:
            pass
        return None

    if os.path.getsize(output_path) == 0:
        logger.info("ssyoutube produced empty file for %s — discarding", video_id)
        try:
            os.remove(output_path)
        except OSError:
            pass
        return None

    logger.info("ssyoutube fallback succeeded for %s → %s", video_id, output_path)
    return output_path


# ── URL Validation & Input Sanitization ──────────────────────────────────────

# Schemes that are never valid for media downloads
_BLOCKED_URL_SCHEMES = {"javascript", "data", "vbscript", "file", "about", "blob"}

# Patterns that indicate script injection or other dangerous input.
# The event-handler pattern requires whitespace/quote before "on" to avoid
# matching legitimate query parameters like "version=" or "connection=".
_URL_INJECTION_PATTERNS = re.compile(
    r"<\s*script|javascript\s*:|(?:[\s\"']|^)on\w+\s*=|<\s*iframe|<\s*img|"
    r"eval\s*\(|expression\s*\(",
    re.IGNORECASE,
)


def _validate_url(url: str) -> str | None:
    """Validate *url* for safety and basic reachability prerequisites.

    Returns ``None`` when the URL is acceptable, or a short human-readable
    error string describing the problem.
    """
    if not url or not url.strip():
        return "URL is required"

    # Reject obvious script-injection payloads before parsing
    if _URL_INJECTION_PATTERNS.search(url):
        return "Invalid URL: script injection detected"

    try:
        parsed = urllib.parse.urlparse(url.strip())
    except Exception:
        return "Invalid URL"

    scheme = (parsed.scheme or "").lower()
    if scheme in _BLOCKED_URL_SCHEMES:
        return f"Invalid URL: scheme '{scheme}' is not allowed"
    if scheme not in ("http", "https"):
        return "Invalid URL: only http and https URLs are supported"

    hostname = parsed.hostname or ""
    if not hostname:
        return "Invalid URL: missing hostname"

    # Guard against bare IP-like or empty hostnames that are not real domains
    if hostname in ("localhost", "localhost.localdomain") or hostname.startswith("127.") or hostname.startswith("0."):
        return "Invalid URL: local addresses are not supported"

    return None


# ── URL Deduplication ─────────────────────────────────────────────────────────

# Maps url_hash (hex-digest) → {"download_id": str, "filename": str | None}
# Protected by downloads_lock.
_download_url_cache: dict[str, dict] = {}

# How long (seconds) a completed download stays in the deduplication cache
_URL_CACHE_TTL_SECONDS = 1800  # 30 minutes


def _url_hash(url: str) -> str:
    """Return the SHA-256 hex digest of the normalised *url*."""
    normalised = url.strip().lower()
    return hashlib.sha256(normalised.encode("utf-8")).hexdigest()


def _get_cached_download(url: str) -> dict | None:
    """Return the cached download record for *url* if it is still valid.

    Returns a dict with at least ``download_id`` and ``filename`` keys, or
    ``None`` when no valid cache entry exists.  Must be called while
    *downloads_lock* is held.
    """
    h = _url_hash(url)
    entry = _download_url_cache.get(h)
    if entry is None:
        return None
    # Invalidate if the file no longer exists on disk
    filename = entry.get("filename")
    if filename:
        filepath = os.path.join(DOWNLOAD_FOLDER, filename)
        if not os.path.isfile(filepath):
            del _download_url_cache[h]
            return None
    # Invalidate if the TTL has expired
    cached_at = entry.get("cached_at", 0)
    if time.time() - cached_at > _URL_CACHE_TTL_SECONDS:
        del _download_url_cache[h]
        return None
    return entry


def _cache_completed_download(url: str, download_id: str, filename: str | None) -> None:
    """Store a completed download in the deduplication cache.

    Must be called while *downloads_lock* is held.
    """
    h = _url_hash(url)
    _download_url_cache[h] = {
        "download_id": download_id,
        "filename": filename,
        "cached_at": time.time(),
    }


# ── Exponential Backoff Retry ─────────────────────────────────────────────────

# Delays in seconds for each retry attempt (1 s → 3 s → 7 s)
_RETRY_DELAYS: tuple[float, ...] = (1.0, 3.0, 7.0)


def _with_exponential_backoff(
    fn,
    max_retries: int = 3,
    delays: tuple[float, ...] = _RETRY_DELAYS,
    retriable_exc: tuple = (Exception,),
    is_retriable=None,
):
    """Call *fn()* with up to *max_retries* retries using exponential backoff.

    *delays* specifies the wait time (in seconds) before each successive retry.
    Only exceptions that are instances of *retriable_exc* are retried; any
    other exception propagates immediately.

    The optional *is_retriable* callable receives the caught exception and
    returns ``True`` when the error should be retried.  Returning ``False``
    causes the exception to propagate immediately without waiting, which is
    useful for errors (e.g. bot-detection / auth) that will not resolve on
    retry.  When *is_retriable* is ``None`` all matching exceptions are retried.

    Returns the value returned by *fn()* on success, or re-raises the last
    exception after all retries are exhausted.
    """
    last_exc: Exception = RuntimeError("_with_exponential_backoff: no attempts made")
    for attempt in range(max_retries + 1):
        try:
            return fn()
        except retriable_exc as exc:
            if is_retriable is not None and not is_retriable(exc):
                raise  # non-retriable error — propagate immediately
            last_exc = exc
            if attempt < max_retries:
                delay = delays[attempt] if attempt < len(delays) else delays[-1]
                logger.info(
                    "Attempt %d/%d failed (%s) — retrying in %.1fs",
                    attempt + 1, max_retries + 1, type(exc).__name__, delay,
                )
                time.sleep(delay)
    raise last_exc


# ── Circuit Breaker ───────────────────────────────────────────────────────────

# Number of consecutive failures before the circuit opens
_CIRCUIT_BREAKER_THRESHOLD = 5
# How long (seconds) the circuit stays open before allowing a probe
_CIRCUIT_BREAKER_COOLDOWN = 300  # 5 minutes


class _CircuitBreaker:
    """Simple thread-safe circuit breaker for the yt-dlp extractor.

    States
    ------
    closed  Normal operation — requests pass through.
    open    Too many consecutive failures — requests are rejected immediately
            and the caller falls back to alternative tools.
    half-open  Cooldown has elapsed — one probe attempt is allowed to test if
               the extractor has recovered.
    """

    def __init__(
        self,
        threshold: int = _CIRCUIT_BREAKER_THRESHOLD,
        cooldown: float = _CIRCUIT_BREAKER_COOLDOWN,
    ) -> None:
        self._threshold = threshold
        self._cooldown = cooldown
        self._failures = 0
        self._open_since: float | None = None
        self._lock = threading.Lock()

    def is_open(self) -> bool:
        """Return ``True`` when the circuit is open (extractor should be skipped)."""
        with self._lock:
            if self._open_since is None:
                return False
            if time.time() - self._open_since >= self._cooldown:
                # Transition to half-open: allow one probe
                return False
            return True

    def record_failure(self) -> None:
        """Record one extractor failure.  Opens the circuit after *threshold* failures."""
        with self._lock:
            self._failures += 1
            if self._failures >= self._threshold and self._open_since is None:
                self._open_since = time.time()
                logger.warning(
                    "Circuit breaker opened after %d consecutive failures — "
                    "yt-dlp extractor temporarily disabled for %ds",
                    self._failures,
                    int(self._cooldown),
                )

    def record_success(self) -> None:
        """Record one successful extraction.  Closes the circuit and resets the counter."""
        with self._lock:
            if self._failures > 0 or self._open_since is not None:
                logger.info(
                    "Circuit breaker reset after successful extraction "
                    "(was at %d failures)",
                    self._failures,
                )
            self._failures = 0
            self._open_since = None

    @property
    def failure_count(self) -> int:
        with self._lock:
            return self._failures


# Global circuit breaker instance for the primary yt-dlp extractor
_extractor_circuit_breaker = _CircuitBreaker()


def _find_media_file(directory: str) -> str | None:
    """Return the path of the first media file found in *directory*, or None."""
    try:
        for entry in os.scandir(directory):
            if entry.is_file() and os.path.splitext(entry.name)[1].lower() in _ALT_MEDIA_EXTS:
                return entry.path
    except OSError:
        pass
    return None


def _try_alternative_tools_download(
    url: str,
    output_dir: str,
    download_id: str | None = None,
) -> str | None:
    """Try alternative download tools for URLs that yt-dlp cannot handle.

    Attempts gallery-dl, you-get, and streamlink (whichever are installed and
    found in PATH) in order, up to a maximum of 3 tools.  Returns the path of
    the downloaded file on success, or ``None`` if all tools fail.

    The optional *download_id* is used to check for cancellation between tool
    attempts so that a queued cancel request is honoured promptly.
    """
    tried = 0
    for cmd_template in _ALTERNATIVE_TOOL_COMMANDS:
        if tried >= 3:
            break
        tool_name = cmd_template[0]
        if not shutil.which(tool_name):
            logger.debug("Alternative tool %r not found in PATH — skipping", tool_name)
            continue
        cmd = [
            tok.replace("{url}", url).replace("{out}", output_dir)
            for tok in cmd_template
        ]
        tried += 1
        logger.info("Trying alternative tool %r for %s", tool_name, url)
        try:
            proc = subprocess.Popen(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                env=get_ssl_env(),
            )
            # Poll while the process runs so cancellation is detected promptly.
            while proc.poll() is None:
                if download_id is not None:
                    with downloads_lock:
                        if downloads.get(download_id, {}).get("status") == "cancelled":
                            proc.terminate()
                            logger.info("Alternative tool cancelled for %s", download_id)
                            return None
                time.sleep(_CANCELLATION_CHECK_INTERVAL_SECONDS)
            if proc.returncode == 0:
                media_path = _find_media_file(output_dir)
                if media_path:
                    logger.info(
                        "Alternative tool %r succeeded for %s → %s",
                        tool_name, url, media_path,
                    )
                    return media_path
                logger.warning(
                    "Alternative tool %r exited 0 but no media file found for %s",
                    tool_name, url,
                )
            else:
                stderr = proc.stderr.read().decode(errors="replace")
                logger.info(
                    "Alternative tool %r failed (rc=%d) for %s: %s",
                    tool_name, proc.returncode, url, stderr[:200],
                )
        except Exception as exc:
            logger.warning("Error running alternative tool %r for %s: %s", tool_name, url, exc)
    return None


def _friendly_cookie_error(error_msg: str) -> str:
    """Return a user-friendly message when YouTube bot-detection triggers.

    Detects the ``Sign in to confirm you're not a bot`` error emitted by
    yt-dlp (and related authentication / login-required errors) and replaces
    them with a plain, actionable message that does not expose admin-panel
    instructions to regular users.
    """
    lower = error_msg.lower()

    if _is_drm_error(error_msg):
        return _GENTLE_FAILURE_MESSAGE

    if _is_http_forbidden_error(error_msg):
        return (
            "This video cannot be downloaded right now — the server received "
            "HTTP 403 Forbidden from YouTube's CDN. "
            "Please try again in a few minutes, or try a different video."
        )

    if _is_auth_error(error_msg):
        return (
            "This video cannot be downloaded right now. "
            "Please try again in a few minutes, or try a different video."
        )

    # Rate-limited by YouTube — hide technical details from regular users
    if "rate-limited by youtube" in lower:
        return (
            "This video cannot be downloaded right now — YouTube has "
            "temporarily rate-limited this server. Please try again in an hour."
        )

    # Private / age-restricted videos
    if "private video" in lower:
        return (
            "This video is private and cannot be downloaded."
        )
    if "age" in lower and ("restricted" in lower or "gate" in lower):
        return (
            "This video is age-restricted and cannot be downloaded "
            "without an authenticated session."
        )

    return error_msg


def format_speed(bytes_per_sec) -> str:
    """Format bytes/s to a human-readable speed string"""
    if bytes_per_sec is None or bytes_per_sec <= 0:
        return ""
    for unit in ("B", "KiB", "MiB", "GiB"):
        if bytes_per_sec < 1024.0:
            return f"{bytes_per_sec:.2f}{unit}/s"
        bytes_per_sec /= 1024.0
    return f"{bytes_per_sec:.2f}TiB/s"

def format_eta(seconds) -> str:
    """Format seconds to mm:ss ETA string"""
    if seconds is None:
        return ""
    m, s = divmod(int(seconds), 60)
    return f"{m}:{s:02d}"

_CHROME_UA = (
    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
    "AppleWebKit/537.36 (KHTML, like Gecko) "
    "Chrome/134.0.0.0 Safari/537.36"
)


def normalize_format_spec(fmt: str) -> str:
    """Ensure the format spec has a fallback so that 'format not available'
    errors are avoided.  Any spec that already contains '/' already has its
    own fallback chain and is returned unchanged.  Otherwise '/best' is
    appended so yt-dlp can always find something to download.
    """
    fmt = (fmt or "best").strip() or "best"
    # Already contains a fallback separator — leave untouched
    if "/" in fmt:
        return fmt
    # Append a generic fallback so yt-dlp can always find something to download
    return f"{fmt}/best"


# Output extension allowlists used by the public download endpoints.
_VALID_OUTPUT_EXTS: set[str] = {"mp4", "webm", "mkv", "avi", "mov", "ts", "3gp"}
_AUDIO_OUTPUT_EXTS: set[str] = {"mp3", "m4a", "ogg", "wav", "opus", "flac", "aac", "weba"}


def _normalize_output_ext(ext: str | None, *, allow_audio: bool = True) -> str:
    """Return a supported output extension, defaulting invalid values to mp4."""
    output_ext = ext.strip().lower() if ext else "mp4"
    allowed_exts = set(_VALID_OUTPUT_EXTS)
    if allow_audio:
        allowed_exts.update(_AUDIO_OUTPUT_EXTS)
    return output_ext if output_ext in allowed_exts else "mp4"



# Formats that must be produced via post-processing conversion rather than
# direct container muxing.  AVI does not support the VP9/Opus/AAC codec
# combinations commonly returned by yt-dlp, so using merge_output_format=avi
# causes "ERROR: Postprocessing: Conversion failed!".  Instead we let yt-dlp
# download in its preferred container (e.g. mp4/mkv) and then invoke
# FFmpegVideoConvertor to remux/transcode into the requested format.
_CONVERT_OUTPUT_EXTS: set[str] = {"avi"}


def _apply_output_ext(ydl_opts: dict, output_template: str, output_ext: str) -> str:
    """Apply container or audio-extraction options and return the final template."""
    if output_ext in _AUDIO_OUTPUT_EXTS:
        ydl_opts["postprocessors"] = [{
            "key": "FFmpegExtractAudio",
            "preferredcodec": output_ext,
            "preferredquality": "192",
        }]
        output_template = os.path.splitext(output_template)[0] + ".%(ext)s"
    elif output_ext in _CONVERT_OUTPUT_EXTS:
        # Use FFmpegVideoConvertor so that yt-dlp downloads in its native best
        # format first and then converts, avoiding merge failures caused by
        # codec incompatibility with the AVI container.
        ydl_opts["postprocessors"] = [{
            "key": "FFmpegVideoConvertor",
            "preferedformat": output_ext,
        }]
        output_template = os.path.splitext(output_template)[0] + ".%(ext)s"
    elif output_ext in _VALID_OUTPUT_EXTS:
        ydl_opts["merge_output_format"] = output_ext
    ydl_opts["outtmpl"] = output_template
    return output_template


def _build_info_dict(info: dict) -> dict:
    """Convert a yt-dlp info dict to the shape returned by ``get_video_info``."""
    return {
        "title": info.get("title", "Unknown"),
        "duration": info.get("duration", 0),
        "uploader": info.get("uploader", "Unknown"),
        "thumbnail": info.get("thumbnail", ""),
        "formats": [
            {
                "format_id": f.get("format_id"),
                "ext": f.get("ext"),
                "resolution": f.get("resolution", "N/A"),
                "filesize": f.get("filesize", 0),
                "format_note": f.get("format_note", "")
            }
            for f in info.get("formats", [])
            if f.get("vcodec") != "none"
        ],
        "audio_formats": [
            {
                "format_id": f.get("format_id"),
                "ext": f.get("ext"),
                "abr": f.get("abr", 0),
                "filesize": f.get("filesize", 0),
                "format_note": f.get("format_note", "")
            }
            for f in info.get("formats", [])
            if f.get("acodec") != "none" and f.get("vcodec") == "none"
        ]
    }


def get_video_info(url: str) -> dict:
    """Get video information without downloading, using the yt-dlp Python API.

    If the initial attempt fails with an authentication/bot-detection error and
    no cookies file is present, a second attempt is made using only the
    ``web_embedded`` and ``tv`` player clients which require no PO tokens or
    authentication and can fetch publicly available videos without cookies.
    """
    _base_opts = {
        "quiet": True,
        "no_warnings": True,
        "noplaylist": True,
        "http_headers": {"User-Agent": _CHROME_UA},
        "extractor_retries": 5,
        "retries": 5,
        "fragment_retries": 5,
        "file_access_retries": 3,
        "sleep_requests": 1,
        "sleep_interval": 5,
        "max_sleep_interval": 10,
        "geo_bypass": True,
        # ⚠️ DO NOT REMOVE — Node.js fallback for JS challenge solving (PR #78)
        "js_runtimes": {"deno": {}, "node": {}},
    }
    try:
        ydl_opts = {
            **_base_opts,
            "extractor_args": _get_yt_extractor_args(),
            **_get_cookie_opts(),
        }

        with yt_dlp.YoutubeDL(ydl_opts) as ydl:
            info = ydl.extract_info(url, download=False)

        return _build_info_dict(info)
    except yt_dlp.utils.DownloadError as e:
        # When bot-detection or HTTP 403 fires and there are no cookies, retry
        # with only the POT-free clients (web_embedded + tv) that work without auth.
        if (_is_auth_error(str(e)) or _is_http_forbidden_error(str(e))) and not os.path.isfile(COOKIES_FILE):
            logger.info("Auth/403 error without cookies — retrying with cookieless clients")
            try:
                ydl_opts_retry = {
                    **_base_opts,
                    "extractor_args": _get_cookieless_extractor_args(),
                }
                with yt_dlp.YoutubeDL(ydl_opts_retry) as ydl:
                    info = ydl.extract_info(url, download=False)
                return _build_info_dict(info)
            except yt_dlp.utils.DownloadError as retry_err:
                logger.info("Cookieless retry also failed: %s", retry_err)
            except Exception as retry_err:
                logger.warning("Unexpected error during cookieless retry: %s", retry_err)
        error_msg = _friendly_cookie_error(str(e))
        return {"error": error_msg}
    except Exception as e:
        return {"error": _friendly_cookie_error(str(e))}

# =========================================================
# DOWNLOAD QUEUE DISPATCHER
# =========================================================

def _queue_dispatcher():
    """Daemon thread that drains _download_queue and dispatches workers
    as concurrency slots become available via _active_semaphore.

    Each item in the queue is a 5-tuple:
        (download_id, url, output_template, format_spec, output_ext)

    The dispatcher blocks on _active_semaphore.acquire() until a running
    download finishes and releases its slot.  Cancelled-while-queued items
    are skipped without consuming a slot.
    """
    while True:
        item = _download_queue.get()  # blocks until an item is available
        download_id, url, output_template, format_spec, output_ext = item

        # Wait for a free concurrency slot
        _active_semaphore.acquire()

        # If the download was cancelled while waiting in the queue, skip it
        with downloads_lock:
            if downloads.get(download_id, {}).get("status") == "cancelled":
                _active_semaphore.release()
                _download_queue.task_done()
                continue
            downloads[download_id]["status"] = "starting"

        # Notify the frontend that this download is now starting
        emit_from_thread("status_update", {"id": download_id, "status": "starting"}, room=download_id)

        def _worker_with_release(did, u, ot, fs, oe):
            try:
                download_worker(did, u, ot, fs, oe)
            finally:
                _active_semaphore.release()
                _download_queue.task_done()

        t = threading.Thread(
            target=_worker_with_release,
            args=(download_id, url, output_template, format_spec, output_ext),
            daemon=True,
        )
        t.start()
        with downloads_lock:
            active_threads[download_id] = t

# Start the dispatcher as a daemon thread so it runs for the lifetime of the process
threading.Thread(target=_queue_dispatcher, daemon=True, name="download-queue-dispatcher").start()

# =========================================================
# DOWNLOAD WORKER
# =========================================================

def download_worker(download_id, url, output_template, format_spec, output_ext=None):
    """Background thread for downloading using the yt-dlp Python API"""

    # Ensure the download directory exists before writing any files.
    # On new deployments or after container restarts the folder may be absent.
    _ensure_download_folder()

    # ── Phase 1: notify the frontend that the worker has started ──────────────
    emit_from_thread("started", {"id": download_id}, room=download_id)

    # ── Phase 2: fetch video info to resolve the real title / filename ─────────
    with downloads_lock:
        downloads[download_id]["status"] = "fetching_info"
    emit_from_thread(
        "status_update",
        {"id": download_id, "status": "fetching_info", "message": "Fetching video info…"},
        room=download_id,
    )

    info = get_video_info(url)

    # Check if the user cancelled the download while info was being fetched.
    with downloads_lock:
        if downloads.get(download_id, {}).get("status") == "cancelled":
            emit_from_thread("cancelled", {"id": download_id}, room=download_id)
            threading.Thread(target=save_downloads_to_disk, daemon=True).start()
            return

    if info and "error" not in info:
        real_title = info.get("title", "")
        if real_title:
            safe_real_title = safe_filename(real_title)
            output_template = os.path.join(DOWNLOAD_FOLDER, f"{safe_real_title}.%(ext)s")
            with downloads_lock:
                downloads[download_id]["title"] = real_title
                downloads[download_id]["safe_title"] = safe_real_title
                downloads[download_id]["output_template"] = output_template
            emit_from_thread(
                "title_update",
                {"id": download_id, "title": real_title},
                room=download_id,
            )
    elif info and "error" in info:
        logger.warning(f"Info error for {url}: {info['error']}")
        with downloads_lock:
            downloads[download_id]["info_error"] = info["error"]
        emit_from_thread(
            "warning",
            {"id": download_id, "message": info["error"]},
            room=download_id,
        )

    def progress_hook(d):
        if d["status"] != "downloading":
            return

        downloaded = d.get("downloaded_bytes") or 0
        total = d.get("total_bytes") or d.get("total_bytes_estimate") or 0
        percent = (100.0 * downloaded / total) if total else 0
        speed = format_speed(d.get("speed"))
        eta = format_eta(d.get("eta"))
        size = format_size(total) if total else ""

        with downloads_lock:
            # Check if download was cancelled (e.g. by page refresh or admin)
            if downloads.get(download_id, {}).get("status") == "cancelled":
                raise yt_dlp.utils.DownloadCancelled("Download cancelled")
            downloads[download_id].update({
                "percent": percent,
                "speed": speed,
                "eta": eta,
                "size": size,
            })

        try:
            emit_from_thread(
                "progress",
                {
                    "id": download_id,
                    "line": "",
                    "percent": percent,
                    "speed": speed,
                    "eta": eta,
                    "size": size,
                },
                room=download_id,
            )
        except Exception as e:
            logger.error(f"Socket emit error: {e}")

    ydl_opts = {
        "format": normalize_format_spec(format_spec),
        "outtmpl": output_template,
        "noplaylist": True,
        "extractor_args": _get_yt_extractor_args(),
        "http_headers": {"User-Agent": _CHROME_UA},
        "extractor_retries": 5,
        "retries": 5,
        "fragment_retries": 5,
        "file_access_retries": 3,
        "sleep_requests": 1,
        "sleep_interval": 5,
        "max_sleep_interval": 10,
        "geo_bypass": True,
        # ⚠️ DO NOT REMOVE — Node.js fallback for JS challenge solving (PR #78)
        "js_runtimes": {"deno": {}, "node": {}},
        "progress_hooks": [progress_hook],
        "quiet": True,
        "no_warnings": True,
        **_get_cookie_opts(),
    }
    output_template = _apply_output_ext(ydl_opts, output_template, output_ext)

    # Add ffmpeg if available
    ffmpeg_path = shutil.which('ffmpeg')
    if ffmpeg_path:
        ydl_opts["ffmpeg_location"] = ffmpeg_path

    def _do_download(opts: dict) -> None:
        """Run yt-dlp download with the given options."""
        with yt_dlp.YoutubeDL(opts) as ydl:
            ydl.download([url])

    def _finalize_completed() -> None:
        """Update state and emit events after a successful download."""
        with downloads_lock:
            downloads[download_id].update({
                "status": "completed",
                "end_time": time.time(),
                "percent": 100
            })

            # Find downloaded file
            base_name = os.path.splitext(os.path.basename(output_template))[0]
            try:
                folder_listing = os.listdir(DOWNLOAD_FOLDER)
            except FileNotFoundError:
                logger.warning(
                    "Downloads folder %r not found while finalising download %s; filename will be unset.",
                    DOWNLOAD_FOLDER, download_id,
                )
                folder_listing = []
            for file in folder_listing:
                if file.startswith(base_name):
                    file_path = os.path.join(DOWNLOAD_FOLDER, file)
                    downloads[download_id].update({
                        "filename": file,
                        "file_size": os.path.getsize(file_path),
                        "file_size_hr": format_size(os.path.getsize(file_path))
                    })
                    # Upload to S3 if configured
                    if _S3_ENABLED:
                        threading.Thread(
                            target=_s3_upload_file,
                            args=(file_path, file),
                            daemon=True,
                        ).start()
                    break

            # Record in the URL deduplication cache
            _cache_completed_download(
                url,
                download_id,
                downloads[download_id].get("filename"),
            )

        emit_from_thread(
            "progress",
            {"id": download_id, "line": "", "percent": 100,
             "speed": "", "eta": "", "size": downloads[download_id].get("size", "")},
            room=download_id,
        )
        emit_from_thread("completed", {
            "id": download_id,
            "filename": downloads[download_id].get("filename"),
            "title": downloads[download_id].get("title")
        }, room=download_id)
        emit_from_thread("files_updated")
        threading.Thread(target=save_downloads_to_disk, daemon=True).start()

    def _ssyoutube_fallback() -> bool:
        """Try the ssyoutube / savefrom.net API and finalize on success.

        Emits a status update, calls :func:`_try_ssyoutube_download`, and on
        success updates the download record and emits completion events.

        Returns ``True`` when the download was handled (success or cancellation),
        ``False`` when the ssyoutube attempt produced no file and the caller
        should proceed to report an error.
        """
        if not _is_youtube_url(url):
            return False
        emit_from_thread(
            "status_update",
            {
                "id": download_id,
                "status": "downloading",
                "message": "Trying ssyoutube fallback…",
            },
            room=download_id,
        )
        ssyt_path = _try_ssyoutube_download(url, DOWNLOAD_FOLDER, download_id)
        if ssyt_path:
            with downloads_lock:
                ssyt_filename = os.path.basename(ssyt_path)
                downloads[download_id].update({
                    "status": "completed",
                    "end_time": time.time(),
                    "percent": 100,
                    "filename": ssyt_filename,
                    "file_size": os.path.getsize(ssyt_path),
                    "file_size_hr": format_size(os.path.getsize(ssyt_path)),
                })
                _cache_completed_download(url, download_id, ssyt_filename)
            if _S3_ENABLED:
                threading.Thread(
                    target=_s3_upload_file,
                    args=(ssyt_path, ssyt_filename),
                    daemon=True,
                ).start()
            emit_from_thread(
                "progress",
                {"id": download_id, "line": "", "percent": 100,
                 "speed": "", "eta": "", "size": downloads[download_id].get("size", "")},
                room=download_id,
            )
            emit_from_thread("completed", {
                "id": download_id,
                "filename": ssyt_filename,
                "title": downloads[download_id].get("title"),
            }, room=download_id)
            emit_from_thread("files_updated")
            threading.Thread(target=save_downloads_to_disk, daemon=True).start()
            return True
        with downloads_lock:
            if downloads.get(download_id, {}).get("status") == "cancelled":
                emit_from_thread("cancelled", {"id": download_id}, room=download_id)
                threading.Thread(target=save_downloads_to_disk, daemon=True).start()
                return True
        return False

    with downloads_lock:
        downloads[download_id]["status"] = "downloading"
        downloads[download_id]["start_time"] = time.time()

    # Check circuit breaker before attempting the primary yt-dlp download.
    # If the circuit is open (too many recent consecutive failures) we skip
    # directly to alternative tools so users get faster feedback.
    if _extractor_circuit_breaker.is_open():
        logger.warning(
            "Circuit breaker is open for %s — skipping yt-dlp, trying alternative tools",
            download_id,
        )
        emit_from_thread(
            "status_update",
            {
                "id": download_id,
                "status": "downloading",
                "message": "Primary extractor temporarily unavailable — trying fallback…",
            },
            room=download_id,
        )
        alt_path = _try_alternative_tools_download(url, DOWNLOAD_FOLDER, download_id)
        if alt_path:
            with downloads_lock:
                alt_filename = os.path.basename(alt_path)
                downloads[download_id].update({
                    "status": "completed",
                    "end_time": time.time(),
                    "percent": 100,
                    "filename": alt_filename,
                    "file_size": os.path.getsize(alt_path),
                    "file_size_hr": format_size(os.path.getsize(alt_path)),
                })
                _cache_completed_download(url, download_id, alt_filename)
            if _S3_ENABLED:
                threading.Thread(
                    target=_s3_upload_file,
                    args=(alt_path, alt_filename),
                    daemon=True,
                ).start()
            emit_from_thread(
                "progress",
                {"id": download_id, "line": "", "percent": 100,
                 "speed": "", "eta": "", "size": downloads[download_id].get("size", "")},
                room=download_id,
            )
            emit_from_thread("completed", {
                "id": download_id,
                "filename": alt_filename,
                "title": downloads[download_id].get("title"),
            }, room=download_id)
            emit_from_thread("files_updated")
            threading.Thread(target=save_downloads_to_disk, daemon=True).start()
        else:
            # Alternative tools all failed; try ssyoutube for YouTube URLs.
            if _ssyoutube_fallback():
                return
            with downloads_lock:
                downloads[download_id].update({
                    "status": "failed",
                    "error": _GENTLE_FAILURE_MESSAGE,
                    "end_time": time.time(),
                })
            emit_from_thread("failed", {"id": download_id, "error": _GENTLE_FAILURE_MESSAGE}, room=download_id)
            threading.Thread(target=save_downloads_to_disk, daemon=True).start()
        return

    try:
        # Wrap the primary yt-dlp download with exponential backoff retry.
        # DownloadCancelled must NOT be retried (user intent), so we only retry
        # on generic DownloadError and unexpected exceptions.
        # Auth/bot-detection errors are also not retried — they will not resolve
        # on retry and we want to fall through to the cookieless/alternative-tool
        # path as quickly as possible.
        def _primary_download():
            _do_download(ydl_opts)

        _with_exponential_backoff(
            _primary_download,
            max_retries=3,
            delays=_RETRY_DELAYS,
            retriable_exc=(yt_dlp.utils.DownloadError, OSError),
            is_retriable=lambda e: not _is_auth_or_forbidden_error(str(e)),
        )
        _extractor_circuit_breaker.record_success()
        _finalize_completed()

    except yt_dlp.utils.DownloadCancelled:
        logger.info(f"Download cancelled via hook: {download_id}")
        with downloads_lock:
            downloads[download_id].update({
                "status": "cancelled",
                "end_time": time.time(),
            })
        emit_from_thread("cancelled", {"id": download_id}, room=download_id)
        threading.Thread(target=save_downloads_to_disk, daemon=True).start()

    except yt_dlp.utils.DownloadError as e:
        # Record failure for circuit breaker tracking
        _extractor_circuit_breaker.record_failure()

        # When bot-detection or HTTP 403 fires and no cookies are present, retry using only
        # the POT-free clients (web_embedded + tv + mweb) that work without authentication.
        final_error: Exception = e
        err_str_initial = str(e)
        if (_is_auth_error(err_str_initial) or _is_http_forbidden_error(err_str_initial)) and not os.path.isfile(COOKIES_FILE):
            logger.info(
                f"Auth/403 error without cookies for {download_id} — "
                "retrying with cookieless clients"
            )
            emit_from_thread(
                "status_update",
                {
                    "id": download_id,
                    "status": "downloading",
                    "message": "Retrying without authentication…",
                },
                room=download_id,
            )
            ydl_opts_retry = {
                **ydl_opts,
                "extractor_args": _get_cookieless_extractor_args(),
            }
            # Defensive: ensure no cookiefile leaks into the cookieless retry
            ydl_opts_retry.pop("cookiefile", None)
            try:
                _do_download(ydl_opts_retry)
                _extractor_circuit_breaker.record_success()
                _finalize_completed()
                return
            except yt_dlp.utils.DownloadCancelled:
                logger.info(f"Download cancelled via hook (cookieless retry): {download_id}")
                with downloads_lock:
                    downloads[download_id].update({
                        "status": "cancelled",
                        "end_time": time.time(),
                    })
                emit_from_thread("cancelled", {"id": download_id}, room=download_id)
                threading.Thread(target=save_downloads_to_disk, daemon=True).start()
                return
            except Exception as retry_err:
                logger.info("Cookieless retry also failed for %s: %s", download_id, retry_err)
                final_error = retry_err

        # For DRM errors, HTTP 403, or after all yt-dlp strategies are exhausted,
        # try alternative tools: gallery-dl, you-get, streamlink (up to 3 tools).
        err_str = str(final_error)
        if _is_drm_error(err_str) or _is_auth_error(err_str) or _is_http_forbidden_error(err_str):
            emit_from_thread(
                "status_update",
                {
                    "id": download_id,
                    "status": "downloading",
                    "message": "Trying alternative download tools…",
                },
                room=download_id,
            )
            alt_path = _try_alternative_tools_download(url, DOWNLOAD_FOLDER, download_id)
            if alt_path:
                # Adopt the file produced by the alternative tool
                with downloads_lock:
                    alt_filename = os.path.basename(alt_path)
                    downloads[download_id].update({
                        "status": "completed",
                        "end_time": time.time(),
                        "percent": 100,
                        "filename": alt_filename,
                        "file_size": os.path.getsize(alt_path),
                        "file_size_hr": format_size(os.path.getsize(alt_path)),
                    })
                    _cache_completed_download(url, download_id, alt_filename)
                if _S3_ENABLED:
                    threading.Thread(
                        target=_s3_upload_file,
                        args=(alt_path, alt_filename),
                        daemon=True,
                    ).start()
                emit_from_thread(
                    "progress",
                    {"id": download_id, "line": "", "percent": 100,
                     "speed": "", "eta": "", "size": downloads[download_id].get("size", "")},
                    room=download_id,
                )
                emit_from_thread("completed", {
                    "id": download_id,
                    "filename": alt_filename,
                    "title": downloads[download_id].get("title"),
                }, room=download_id)
                emit_from_thread("files_updated")
                threading.Thread(target=save_downloads_to_disk, daemon=True).start()
                return
            # Check if cancelled during alternative tool attempts
            with downloads_lock:
                if downloads.get(download_id, {}).get("status") == "cancelled":
                    emit_from_thread("cancelled", {"id": download_id}, room=download_id)
                    threading.Thread(target=save_downloads_to_disk, daemon=True).start()
                    return

            # Last resort: try the ssyoutube / savefrom.net API for YouTube URLs.
            if _ssyoutube_fallback():
                return

        error_msg = _friendly_cookie_error(err_str)
        with downloads_lock:
            downloads[download_id].update({
                "status": "failed",
                "error": error_msg,
                "end_time": time.time(),
            })
        emit_from_thread("failed", {
            "id": download_id,
            "error": error_msg
        }, room=download_id)
        threading.Thread(target=save_downloads_to_disk, daemon=True).start()

    except Exception as e:
        _extractor_circuit_breaker.record_failure()
        logger.error(f"Download worker error: {e}")
        error_msg = _friendly_cookie_error(str(e))
        with downloads_lock:
            downloads[download_id].update({
                "status": "failed",
                "error": error_msg,
                "end_time": time.time(),
            })
        emit_from_thread("failed", {
            "id": download_id,
            "error": error_msg
        }, room=download_id)
        threading.Thread(target=save_downloads_to_disk, daemon=True).start()


    finally:
        # Cleanup
        with downloads_lock:
            if download_id in active_threads:
                del active_threads[download_id]



# =========================================================
# REACT SPA HELPERS
# =========================================================

def _react_index():
    """Return the React build index.html, falling back to a basic HTML stub."""
    idx = os.path.join(FRONTEND_DIST, "index.html")
    if os.path.isfile(idx):
        return FileResponse(idx, media_type="text/html")
    # Fallback stub while the React build is absent (dev/CI)
    return Response(
        '<html><body><h1>YOT Downloader</h1>'
        '<p>React frontend build not found. Run <code>npm run build</code> inside the '
        '<code>frontend/</code> directory.</p></body></html>',
        media_type="text/html",
        status_code=200,
    )


# =========================================================
# ROUTES
# =========================================================

@fastapi_app.get("/")
async def index(request: Request):
    """Serve the React SPA (main page)"""
    return _react_index()

@fastapi_app.get("/ads.txt")
async def ads_txt():
    """Serve ads.txt for Google AdSense verification"""
    ads_txt_path = os.path.join(ROOT_DIR, "ads.txt")
    if os.path.exists(ads_txt_path):
        return FileResponse(ads_txt_path, media_type="text/plain")
    logger.warning("ads.txt file not found at %s", ads_txt_path)
    return JSONResponse({"error": "ads.txt not found"}, status_code=404)

@fastapi_app.get("/yotweek.png")
async def yotweek_icon():
    """Serve the yotweek brand icon."""
    icon_path = os.path.join(ROOT_DIR, "yotweek.png")
    if os.path.exists(icon_path):
        return FileResponse(icon_path, media_type="image/png")
    return JSONResponse({"error": "Icon not found"}, status_code=404)

@fastapi_app.get("/health")
async def health():
    """Health check endpoint.

    Verifies:
    - yt-dlp is importable and returns a version string
    - Sufficient disk space is available in the downloads folder
    - Basic network connectivity (DNS resolution)
    - Circuit breaker state

    Returns HTTP 200 when all checks pass, or HTTP 503 when a critical
    subsystem is degraded.
    """
    checks: dict[str, object] = {}
    overall_healthy = True

    # --- yt-dlp check ---
    try:
        ytdlp_version = yt_dlp.version.__version__
        checks["yt_dlp"] = {"ok": True, "version": ytdlp_version}
    except Exception as exc:
        checks["yt_dlp"] = {"ok": False, "error": str(exc)}
        overall_healthy = False

    # --- Disk space check (warn below 500 MB, fail below 100 MB) ---
    try:
        usage = shutil.disk_usage(DOWNLOAD_FOLDER)
        free_mb = usage.free // (1024 * 1024)
        checks["disk"] = {"ok": free_mb >= 100, "free_mb": free_mb}
        if free_mb < 100:
            overall_healthy = False
    except Exception as exc:
        checks["disk"] = {"ok": False, "error": str(exc)}
        overall_healthy = False

    # --- Network connectivity check (DNS probe with 2-second timeout) ---
    # Tries multiple well-known DNS servers so corporate / restricted networks
    # that block 8.8.8.8 can still succeed via an alternative probe.
    _DNS_PROBES = [("8.8.8.8", 53), ("1.1.1.1", 53), ("9.9.9.9", 53)]
    network_ok = False
    for _host, _port in _DNS_PROBES:
        try:
            _sock = socket.create_connection((_host, _port), timeout=2)
            _sock.close()
            network_ok = True
            break
        except Exception:
            pass
    checks["network"] = {"ok": network_ok}
    # Network failure is reported but does not mark the service unhealthy
    # because it may be a transient DNS hiccup or environment restriction.

    # --- Circuit breaker state ---
    cb_open = _extractor_circuit_breaker.is_open()
    checks["circuit_breaker"] = {
        "ok": not cb_open,
        "open": cb_open,
        "failures": _extractor_circuit_breaker.failure_count,
    }
    if cb_open:
        overall_healthy = False

    status_code = 200 if overall_healthy else 503
    return JSONResponse(
        {
            "status": "healthy" if overall_healthy else "degraded",
            "timestamp": datetime.now().isoformat(),
            "version": "1.0.0",
            "checks": checks,
        },
        status_code=status_code,
    )


@fastapi_app.get("/api/youtube_status")
async def youtube_status():
    """Probe YouTube reachability and report whether bot-detection is active.

    Runs ``check_youtube_connectivity()`` in a thread so the async event loop
    is not blocked during the network call.  The result is returned as JSON:

    .. code-block:: json

        {
            "reachable":     true,
            "bot_detected":  false,
            "message":       "YouTube is reachable",
            "timestamp":     "2026-01-01T00:00:00"
        }

    The endpoint is intentionally unauthenticated so monitoring tools can poll
    it without credentials.
    """
    loop = asyncio.get_event_loop()
    result = await loop.run_in_executor(None, check_youtube_connectivity)
    result["timestamp"] = datetime.now().isoformat()
    status_code = 200 if result["reachable"] else 503
    return JSONResponse(result, status_code=status_code)


@fastapi_app.post("/api/client_hints")
async def client_hints(request: Request):
    """Accept client-side hints (browser timezone, language) for improved geo-detection.

    The browser reports its local ``Intl.DateTimeFormat().resolvedOptions().timeZone``
    and ``navigator.language`` via a small JS snippet.  These signals are used to
    refine the cached country for the client IP when all server-side geo-IP
    services have already been tried but the country is still unknown.

    The endpoint is intentionally lightweight – it returns immediately and does
    all cache updates synchronously (the data is already in-process).
    """
    ip = _get_real_ip(request)
    if not ip or ip == "unknown":
        return JSONResponse({"ok": False}, status_code=400)

    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"ok": False}, status_code=400)

    tz = (body.get("timezone") or "").strip()
    lang = (body.get("language") or "").strip()

    # Only refine – never overwrite a successfully resolved country
    cached = ip_country_cache.get(ip, {})
    if cached.get("country") and cached["country"] not in ("Unknown", ""):
        return JSONResponse({"ok": True, "country": cached["country"]})

    country, code = "", ""

    # 1. Try IANA timezone from the browser (more precise than IP heuristic)
    if tz and not code:
        tz_code = _TZ_TO_COUNTRY.get(tz, "")
        if tz_code and tz_code in _ISO2_TO_NAME:
            code = tz_code
            country = _ISO2_TO_NAME[code]

    # 2. Try Accept-Language-style hint from navigator.language
    if not code and lang:
        country, code = _country_from_accept_language(lang)

    if country and code:
        ip_country_cache[ip] = {
            "country": country, "code": code, "city": "", "region": ""
        }
        # Back-fill pending visitor / download records (last 200 entries is a
        # reasonable window that covers recent unresolved visitors without
        # iterating the entire in-memory list)
        with visitors_lock:
            for v in visitors[-200:]:
                if v.get("ip") == ip and not v.get("country"):
                    v["country"] = country
                    v["country_code"] = code
        with downloads_lock:
            for d in downloads.values():
                if d.get("ip") == ip and not d.get("country"):
                    d["country"] = country
                    d["country_code"] = code
        _schedule_visitor_save()

    return JSONResponse({"ok": True, "country": country or "Unknown"})

@fastapi_app.post("/video_info")
@rate_limit()
async def video_info_endpoint(request: Request, url: str = Form(None)):
    """Fetch video metadata (title, thumbnail, duration, available formats) without downloading."""
    if not url or not url.strip():
        return JSONResponse({"error": "URL is required"}, status_code=400)
    url = url.strip()
    info = get_video_info(url)
    if not info or "error" in info:
        return JSONResponse({"error": (info or {}).get("error", "Failed to fetch video info")}, status_code=422)
    return JSONResponse(info)


@fastapi_app.post("/start_download")
@rate_limit()
async def start_download(request: Request, url: str = Form(None), format: str = Form("bv*+ba/b"), ext: str = Form("mp4"), session_id: str = Form(None)):
    """Start a download with better error feedback"""
    format_spec = format
    output_ext  = _normalize_output_ext(ext, allow_audio=True)

    if not url:
        return JSONResponse({"error": "URL is required"}, status_code=400)

    # Validate and sanitize the URL (blocks scripts, invalid schemes, etc.)
    url_error = _validate_url(url)
    if url_error:
        return JSONResponse({"error": url_error}, status_code=400)

    # Reject only if the queue is already full
    if _download_queue.full():
        return JSONResponse({
            "error": "Waiting for available slot…"
        }, status_code=429)

    # Deduplication: return the existing download if this URL was recently completed
    with downloads_lock:
        cached = _get_cached_download(url)
        if cached is not None:
            cached_id = cached.get("download_id", "")
            existing = downloads.get(cached_id, {})
            if existing.get("status") == "completed" and existing.get("filename"):
                logger.info("Returning cached download for %s → %s", url, cached_id)
                return JSONResponse({
                    "download_id": cached_id,
                    "title": existing.get("title", ""),
                    "status": "completed",
                    "filename": existing.get("filename"),
                    "cached": True,
                })

    download_id = str(uuid.uuid4())

    # Use a placeholder title; the worker will resolve the real title via
    # get_video_info() in the background and emit a title_update event.
    title = f"video_{download_id[:8]}"
    safe_title = safe_filename(title)
    output_template = os.path.join(DOWNLOAD_FOLDER, f"{safe_title}.%(ext)s")

    # Store download info
    ip = _get_real_ip(request)
    # Try CDN/proxy headers first for country
    hdr_country, hdr_code = _get_country_from_headers(request)
    if hdr_country and ip not in ip_country_cache:
        ip_country_cache[ip] = {"country": hdr_country, "code": hdr_code, "city": "", "region": ""}
    elif ip not in ip_country_cache and _is_private_ip(ip):
        ip_country_cache[ip] = {"country": "Local", "code": "", "city": "", "region": ""}
    cached_geo = ip_country_cache.get(ip, {})
    with downloads_lock:
        downloads[download_id] = {
            "id": download_id,
            "url": url,
            "title": title,
            "safe_title": safe_title,
            "status": "queued",
            "percent": 0,
            "output_template": output_template,
            "format": format_spec,
            "type": "video",
            "created_at": time.time(),
            "filename": None,
            "ip": ip,
            "country": cached_geo.get("country", ""),
            "country_code": cached_geo.get("code", ""),
            "city": cached_geo.get("city", ""),
            "region": cached_geo.get("region", ""),
            "info_error": None,
            "owner_session": session_id or "",
        }
    # Resolve the requester's country in background if not already cached
    if ip not in ip_country_cache:
        accept_lang = request.headers.get("accept-language", "")
        threading.Thread(
            target=_lookup_country_async, args=(ip, accept_lang), daemon=True
        ).start()

    # Enqueue the download; the dispatcher will start it when a slot is free.
    _download_queue.put((download_id, url, output_template, format_spec, output_ext))

    threading.Thread(target=save_downloads_to_disk, daemon=True).start()

    return JSONResponse({
        "download_id": download_id,
        "title": title,
        "status": "queued",
    })

@fastapi_app.get("/status/{download_id}")
async def get_status(download_id: str):
    """Get download status"""
    with downloads_lock:
        download = downloads.get(download_id, {})
        safe_download = {
            "id": download.get("id"),
            "title": download.get("title"),
            "status": download.get("status"),
            "percent": download.get("percent"),
            "speed": download.get("speed"),
            "eta": download.get("eta"),
            "size": download.get("size"),
            "filename": download.get("filename"),
            "file_size_hr": download.get("file_size_hr"),
            "error": download.get("error")
        }
    return JSONResponse(safe_download)

@fastapi_app.get("/files")
async def list_files(request: Request, session_id: str = None):
    """List downloaded files.

    When *session_id* is provided (non-admin users), only files that were
    downloaded in the same session are returned.  Admin users see every file.
    """
    is_admin = request.session.get("admin_logged_in", False)

    # Build a mapping from filename → owner_session from in-memory download records
    with downloads_lock:
        filename_to_session: dict[str, str] = {
            d["filename"]: d.get("owner_session", "")
            for d in downloads.values()
            if d.get("filename")
        }

    files = []
    try:
        _ensure_download_folder()
        for name in os.listdir(DOWNLOAD_FOLDER):
            path = os.path.join(DOWNLOAD_FOLDER, name)
            if not os.path.isfile(path):
                continue
            owner_sess = filename_to_session.get(name, "")
            # Non-admin users only see files belonging to their session
            if not is_admin and session_id:
                if owner_sess != session_id:
                    continue
            elif not is_admin and not session_id:
                # No session_id provided by a non-admin: show nothing
                continue
            stat = os.stat(path)
            files.append({
                "name": name,
                "size": stat.st_size,
                "size_hr": format_size(stat.st_size),
                "modified": stat.st_mtime,
                "modified_str": datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M"),
                "url": str(request.url_for('download_file', filename=name)),
                "owner_session": owner_sess,
            })
        files.sort(key=lambda f: f["modified"], reverse=True)
    except FileNotFoundError:
        # Downloads directory was removed after startup (e.g. container restart).
        # Return an empty list rather than a 500 so the UI stays functional.
        logger.warning("Downloads folder %r not found while listing files; returning empty list.", DOWNLOAD_FOLDER)
    except Exception as e:
        logger.error(f"Error listing files: {e}")
        return JSONResponse({"error": "Failed to list files"}, status_code=500)

    return JSONResponse(files)


@fastapi_app.delete("/session/{session_id}")
async def delete_session_files(session_id: str):
    """Delete all files associated with *session_id*.

    Called by the frontend on page load to clean up the previous session's
    downloaded files, ensuring each page refresh starts fresh.
    """
    if not session_id or len(session_id) > 128:
        return JSONResponse({"deleted": []})

    deleted = []
    with downloads_lock:
        owned_files = [
            d["filename"]
            for d in downloads.values()
            if d.get("owner_session") == session_id and d.get("filename")
        ]

    for filename in owned_files:
        filepath = os.path.join(DOWNLOAD_FOLDER, filename)
        try:
            if _S3_ENABLED:
                _s3_delete_file(filename)
            local_deleted = (
                os.path.abspath(filepath).startswith(os.path.abspath(DOWNLOAD_FOLDER))
                and os.path.isfile(filepath)
            )
            if local_deleted:
                os.remove(filepath)
            if local_deleted or _S3_ENABLED:
                deleted.append(filename)
        except Exception as e:
            logger.warning(f"Could not delete session file {filename}: {e}")

    return JSONResponse({"deleted": deleted})

# MIME types that Python's mimetypes module maps incorrectly or leaves absent.
# Explicit overrides are applied for both the /downloads/ (attachment) and
# /stream/ (inline) endpoints so browsers receive an accurate Content-Type.
# The correct MIME type is especially important on iOS Safari, which uses the
# Content-Type to decide whether to attempt inline playback or hand the file
# off to the Files app / a share sheet.
_MIME_OVERRIDES = {
    # Video
    ".avi":  "video/x-msvideo",  # Python maps .avi → None or video/avi (non-standard)
    ".mkv":  "video/x-matroska", # Python has no mapping for .mkv
    ".mov":  "video/quicktime",  # Python maps .mov → video/quicktime (correct, kept explicit)
    ".wmv":  "video/x-ms-wmv",   # Python has no reliable mapping
    ".flv":  "video/x-flv",      # Python has no reliable mapping
    ".ts":   "video/mp2t",       # Python maps .ts → text/vnd.trolltech.linguist
    ".3gp":  "video/3gpp",       # Python maps .3gp → audio/3gpp (wrong for video)
    ".3g2":  "video/3gpp2",      # Python maps .3g2 → audio/3gpp2 (wrong for video)
    # Audio
    ".weba": "audio/webm",       # Python has no mapping for .weba
    ".opus": "audio/opus",       # Python maps .opus → audio/ogg (imprecise)
}

# Extensions that iOS Safari cannot play natively (no built-in codec support).
# Used by the /downloads/ endpoint to add an X-iOS-Unsupported header so the
# frontend can display a user-friendly warning.
_IOS_UNSUPPORTED_EXTS = {".avi", ".mkv", ".wmv", ".flv"}

@fastapi_app.get("/downloads/{filename:path}", name="download_file")
async def download_file(filename: str):
    """Serve a downloaded file as an attachment with the correct MIME type.

    Content-Disposition: attachment is set explicitly so that iOS Safari
    triggers a download / share-sheet prompt instead of trying to play the
    file inline.  The proper MIME type (e.g. video/x-msvideo for .avi) is
    used rather than the generic application/octet-stream so the OS can
    associate the file with an appropriate app on the device.

    For formats that iOS cannot play natively (AVI, MKV, WMV, FLV) an
    ``X-iOS-Unsupported: true`` response header is included so the frontend
    can surface a warning to the user.
    """
    try:
        # When S3 is configured, redirect to a presigned URL if the object exists
        if _S3_ENABLED:
            presigned = _s3_presigned_url(filename)
            if presigned:
                return RedirectResponse(presigned)
            # Fall through to local file if the S3 object is not found yet
        filepath = os.path.join(DOWNLOAD_FOLDER, filename)
        if not os.path.abspath(filepath).startswith(os.path.abspath(DOWNLOAD_FOLDER)):
            return JSONResponse({"error": "Invalid filename"}, status_code=400)
        if not os.path.isfile(filepath):
            return JSONResponse({"error": "File not found"}, status_code=404)
        ext = os.path.splitext(filename)[1].lower()
        media_type = _MIME_OVERRIDES.get(ext) or mimetypes.guess_type(filename)[0] or "application/octet-stream"
        # Use a safe ASCII filename in the Content-Disposition header; for
        # filenames with non-ASCII characters also include a UTF-8 encoded
        # filename* parameter so modern browsers (including Mobile Safari)
        # use the correct name.  Fall back to "download" when stripping
        # non-ASCII characters would leave an empty string.
        safe_ascii = filename.encode("ascii", "ignore").decode().strip() or "download"
        encoded    = urllib.parse.quote(filename, safe="")
        disposition = f'attachment; filename="{safe_ascii}"; filename*=UTF-8\'\'{encoded}'
        headers = {"Content-Disposition": disposition}
        if ext in _IOS_UNSUPPORTED_EXTS:
            headers["X-iOS-Unsupported"] = "true"
        return FileResponse(filepath, media_type=media_type, headers=headers)
    except Exception as e:
        logger.error(f"Download error: {e}")
        return JSONResponse({"error": "File not found"}, status_code=404)

@fastapi_app.get("/stream/{filename:path}")
async def stream_file(filename: str):
    """Serve a downloaded file inline for in-browser preview.

    Differs from /downloads/ in that the file is served without the
    Content-Disposition: attachment header, so browsers (including iOS Safari)
    can play it directly in a <video>/<audio> element.

    Explicit MIME type overrides are applied for extensions that Python's
    mimetypes module maps incorrectly (e.g. .ts → text/vnd.trolltech.linguist),
    which would otherwise prevent the browser from decoding the audio track.
    """
    try:
        # When S3 is configured, redirect to a presigned URL for inline playback
        if _S3_ENABLED:
            presigned = _s3_presigned_url(filename)
            if presigned:
                return RedirectResponse(presigned)
            # Fall through to local file if the S3 object is not found yet
        filepath = os.path.join(DOWNLOAD_FOLDER, filename)
        if not os.path.realpath(filepath).startswith(os.path.realpath(DOWNLOAD_FOLDER)):
            return JSONResponse({"error": "Invalid filename"}, status_code=400)
        if not os.path.isfile(filepath):
            return JSONResponse({"error": "File not found"}, status_code=404)
        ext = os.path.splitext(filename)[1].lower()
        media_type = _MIME_OVERRIDES.get(ext) or mimetypes.guess_type(filename)[0] or "application/octet-stream"
        return FileResponse(
            filepath,
            media_type=media_type,
            content_disposition_type="inline",
        )
    except Exception as e:
        logger.error(f"Stream error: {e}")
        return JSONResponse({"error": "File not found"}, status_code=404)


@fastapi_app.delete("/delete/{filename:path}")
async def delete_file(filename: str):
    """Delete a downloaded file"""
    try:
        filepath = os.path.join(DOWNLOAD_FOLDER, filename)

        # Security check
        if not os.path.abspath(filepath).startswith(os.path.abspath(DOWNLOAD_FOLDER)):
            return JSONResponse({"error": "Invalid filename"}, status_code=400)

        # Delete from S3 if configured
        if _S3_ENABLED:
            _s3_delete_file(filename)

        local_exists = os.path.exists(filepath) and os.path.isfile(filepath)
        if local_exists:
            os.remove(filepath)
            logger.info(f"Deleted file: {filename}")

        if local_exists or _S3_ENABLED:
            emit_from_thread("files_updated")
            return JSONResponse({"success": True})
        else:
            return JSONResponse({"error": "File not found"}, status_code=404)
    except Exception as e:
        logger.error(f"Delete error: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)

@fastapi_app.get("/stats")
async def get_stats():
    """Get download statistics including all-time totals"""
    try:
        files = []
        total_size = 0
        try:
            _ensure_download_folder()
            listing = os.listdir(DOWNLOAD_FOLDER)
        except FileNotFoundError:
            listing = []
            logger.warning("Downloads folder %r not found while computing stats; using zero counts.", DOWNLOAD_FOLDER)
        for name in listing:
            path = os.path.join(DOWNLOAD_FOLDER, name)
            if os.path.isfile(path):
                size = os.path.getsize(path)
                total_size += size
                files.append(name)

        with downloads_lock:
            active_count = sum(1 for d in downloads.values()
                              if d["status"] in ("starting", "fetching_info", "queued", "downloading"))

        persistent = _get_persistent_stats()

        return JSONResponse({
            "file_count": len(files),
            "total_size": total_size,
            "total_size_hr": format_size(total_size),
            "active_downloads": active_count,
            "max_concurrent": Config.MAX_CONCURRENT_DOWNLOADS,
            "total_downloads": persistent["total_downloads"],
            "total_visitors": persistent["total_site_visitors"],
        })
    except Exception as e:
        logger.error(f"Stats error: {e}")
        return JSONResponse({"error": str(e)}, status_code=500)

@fastapi_app.get("/active_downloads")
async def active_downloads_list():
    """Return count and details of active/queued downloads"""
    with downloads_lock:
        all_active = [
            d for d in downloads.values()
            if d["status"] in ("starting", "fetching_info", "queued", "downloading")
        ]

    # Assign queue positions to items still waiting (status == "queued"),
    # ordered by creation time so position 1 is the next to be dispatched.
    queued_sorted = sorted(
        [d for d in all_active if d["status"] == "queued"],
        key=lambda d: d.get("created_at", 0),
    )
    queue_pos = {d["id"]: i + 1 for i, d in enumerate(queued_sorted)}

    active = [
        {
            "id": d["id"],
            "title": d.get("title"),
            "status": d["status"],
            "percent": d.get("percent", 0),
            "speed": d.get("speed", ""),
            "eta": d.get("eta", ""),
            "size": d.get("size", ""),
            "queue_position": queue_pos.get(d["id"]),
        }
        for d in all_active
    ]
    return JSONResponse({
        "count": len(active),
        "downloads": active,
        "queue_size": _download_queue.qsize(),
        "max_concurrent": Config.MAX_CONCURRENT_DOWNLOADS,
    })

@fastapi_app.post("/cancel/{download_id}")
async def cancel_download(download_id: str):
    """Cancel an ongoing download"""
    with downloads_lock:
        if download_id in downloads:
            if downloads[download_id]["status"] in ("starting", "fetching_info", "queued", "downloading"):
                downloads[download_id]["status"] = "cancelled"
                downloads[download_id]["end_time"] = time.time()
                emit_from_thread("cancelled", {"id": download_id}, room=download_id)
                logger.info(f"Cancelled download: {download_id}")
                return JSONResponse({"success": True})

    return JSONResponse({"error": "Download not found"}, status_code=404)

@fastapi_app.post("/cancel_all")
async def cancel_all_downloads(request: Request):
    """Cancel all active/queued downloads for the requesting client IP.
    Called on page refresh so orphaned downloads are cleaned up."""
    ip = _get_real_ip(request)
    cancelled_ids = []
    with downloads_lock:
        for did, d in downloads.items():
            if d["status"] in ("starting", "fetching_info", "queued", "downloading") and d.get("ip") == ip:
                d["status"] = "cancelled"
                d["end_time"] = time.time()
                cancelled_ids.append(did)
    for did in cancelled_ids:
        emit_from_thread("cancelled", {"id": did}, room=did)
    if cancelled_ids:
        logger.info(f"Cancelled {len(cancelled_ids)} downloads for IP {ip}")
        threading.Thread(target=save_downloads_to_disk, daemon=True).start()
    return JSONResponse({"success": True, "cancelled": len(cancelled_ids)})


# ── Review endpoints ──────────────────────────────────────────────
def _load_reviews_from_db():
    """Load reviews from the database into memory on startup."""
    try:
        with _db_lock:
            conn = _get_db()
            try:
                order_col = "id" if USE_POSTGRES else "rowid"
                rows = _execute(conn, f"SELECT data FROM reviews ORDER BY {order_col} DESC").fetchall()
            finally:
                conn.close()
        loaded = []
        for row in rows:
            try:
                loaded.append(json.loads(row["data"]))
            except Exception:
                pass
        with reviews_lock:
            reviews.clear()
            reviews.extend(loaded)
    except Exception as e:
        logger.warning("Could not load reviews from DB: %s", e)


def _save_review_to_db(review_data: dict):
    """Persist a single review to the database."""
    try:
        with _db_lock:
            conn = _get_db()
            try:
                _execute(conn, "INSERT INTO reviews (data) VALUES (?)", (json.dumps(review_data),))
                conn.commit()
            finally:
                conn.close()
    except Exception as e:
        logger.warning("Could not save review to DB: %s", e)


@fastapi_app.get("/reviews")
async def get_reviews():
    """Return all reviews for the public home page (IP stripped for privacy)."""
    with reviews_lock:
        safe = [{k: v for k, v in r.items() if k != "ip"} for r in reviews]
    return JSONResponse(safe)


@fastapi_app.get("/reviews/can_submit")
async def can_submit_review(request: Request):
    """Return whether this IP is allowed to submit another review."""
    ip = _get_real_ip(request)
    with reviews_lock:
        count = sum(1 for r in reviews if r.get("ip") == ip)
    return JSONResponse({"can_submit": count < Config.MAX_REVIEWS_PER_IP})


@fastapi_app.post("/reviews")
async def submit_review(request: Request):
    """Submit a new review from a visitor (max 1 per IP)."""
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "Invalid JSON"}, status_code=400)

    name = (body.get("name") or "Anonymous").strip()[:50] or "Anonymous"
    comment = (body.get("comment") or "").strip()[:500]
    rating = body.get("rating")

    try:
        rating = int(rating)
        if rating < 1 or rating > 5:
            raise ValueError
    except (TypeError, ValueError):
        return JSONResponse({"error": "Rating must be 1-5"}, status_code=400)

    ip = _get_real_ip(request)
    with reviews_lock:
        ip_review_count = sum(1 for r in reviews if r.get("ip") == ip)
        if ip_review_count >= Config.MAX_REVIEWS_PER_IP:
            return JSONResponse(
                {"error": f"You have already submitted the maximum of {Config.MAX_REVIEWS_PER_IP} reviews."},
                status_code=429,
            )

    review = {
        "id": str(uuid.uuid4()),
        "name": name,
        "comment": comment,
        "rating": rating,
        "timestamp": time.time(),
        "ip": ip,
    }

    with reviews_lock:
        reviews.insert(0, review)

    _save_review_to_db(review)

    return JSONResponse({"success": True, "review": {k: v for k, v in review.items() if k != "ip"}})


@fastapi_app.get("/api/admin/reviews")
async def api_admin_reviews(request: Request):
    """Return all reviews including IP info (admin only)."""
    if not request.session.get("admin_user") and not request.session.get("admin_logged_in"):
        return JSONResponse({"error": "Admin login required."}, status_code=401)
    with reviews_lock:
        all_reviews = list(reviews)
    return JSONResponse({"reviews": all_reviews})


@fastapi_app.delete("/api/admin/reviews/{review_id}")
async def api_admin_delete_review(request: Request, review_id: str):
    """Delete (reject) a review by its ID (admin only)."""
    if not request.session.get("admin_user") and not request.session.get("admin_logged_in"):
        return JSONResponse({"error": "Admin login required."}, status_code=401)

    with reviews_lock:
        before = len(reviews)
        reviews[:] = [r for r in reviews if r.get("id") != review_id]
        removed = before - len(reviews)

    if removed == 0:
        return JSONResponse({"error": "Review not found."}, status_code=404)

    # Remove from database
    try:
        with _db_lock:
            conn = _get_db()
            try:
                if USE_POSTGRES:
                    cur = conn.cursor()
                    cur.execute(
                        "DELETE FROM reviews WHERE (data::jsonb)->>'id' = %s",
                        (review_id,),
                    )
                else:
                    conn.execute(
                        "DELETE FROM reviews WHERE json_extract(data, '$.id') = ?",
                        (review_id,),
                    )
                conn.commit()
            finally:
                conn.close()
    except Exception as e:
        logger.warning(f"Could not delete review from DB: {e}")

    return JSONResponse({"ok": True})


@fastapi_app.get("/const")
async def admin_page(request: Request):
    """Admin dashboard — served via React SPA (authentication checked client-side via /admin/auth_status)"""
    return _react_index()


@fastapi_app.get("/admin/login")
async def admin_login_get(request: Request):
    """Admin login — served via React SPA."""
    return _react_index()


@fastapi_app.post("/admin/login")
async def admin_login_post(request: Request):
    """Admin login (POST with form data) — legacy form-based endpoint kept for compatibility."""
    form = await request.form()
    username = (form.get("username") or "").strip()
    password = form.get("password") or ""
    error = None

    if not admin_user_exists():
        if secrets.compare_digest(password, Config.ADMIN_PASSWORD):
            request.session["admin_logged_in"] = True
            request.session["admin_username"] = username or "admin"
            next_url = request.query_params.get("next") or "/const"
            return RedirectResponse(url=next_url, status_code=302)
        error = "Incorrect password. Please try again."
    else:
        if verify_admin_user(username, password):
            request.session["admin_logged_in"] = True
            request.session["admin_username"] = username
            next_url = request.query_params.get("next") or "/const"
            return RedirectResponse(url=next_url, status_code=302)
        error = "Incorrect username or password. Please try again."

    return templates.TemplateResponse(
        "admin_login.html",
        {"request": request, "error": error, "has_admin": admin_user_exists()},
    )


@fastapi_app.get("/admin/register")
async def admin_register_get(request: Request):
    """Admin registration — served via React SPA."""
    return _react_index()


@fastapi_app.post("/admin/register")
async def admin_register_post(request: Request):
    """Admin registration (POST with form data) — legacy endpoint kept for compatibility."""
    if request.session.get("admin_logged_in"):
        return RedirectResponse(url="/const", status_code=302)
    form = await request.form()
    username = (form.get("username") or "").strip()
    password = form.get("password") or ""
    confirm  = form.get("confirm_password") or ""
    error = None
    success = None
    if not username or not password:
        error = "Username and password are required."
    elif password != confirm:
        error = "Passwords do not match."
    else:
        ok, msg = register_admin_user(username, password)
        if ok:
            success = "Admin account created! You can now log in."
        else:
            error = msg
    return templates.TemplateResponse(
        "admin_login.html",
        {"request": request, "error": error, "success": success,
         "register_mode": True, "has_admin": admin_user_exists()},
    )


@fastapi_app.post("/admin/logout")
async def admin_logout(request: Request):
    """Log out of the admin panel."""
    request.session.pop("admin_logged_in", None)
    request.session.pop("admin_username", None)
    return RedirectResponse(url="/admin/login", status_code=302)


# ─────────────────────────────────────────────────────────────────────────────
# JSON AUTH ENDPOINTS — consumed by the React frontend
# ─────────────────────────────────────────────────────────────────────────────

@fastapi_app.get("/admin/auth_status")
async def admin_auth_status(request: Request):
    """Return current admin session status as JSON (used by React SPA)."""
    logged_in = bool(request.session.get("admin_logged_in"))
    return JSONResponse({
        "logged_in": logged_in,
        "username": request.session.get("admin_username") if logged_in else None,
    })


@fastapi_app.get("/admin/has_admin")
async def admin_has_admin():
    """Return whether an admin account has been registered (used by React SPA)."""
    return JSONResponse({"has_admin": admin_user_exists()})


@fastapi_app.post("/admin/api/login")
async def admin_api_login(request: Request):
    """JSON admin login endpoint (used by React SPA)."""
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "Invalid JSON"}, status_code=400)

    username = (body.get("username") or "").strip()
    password = body.get("password") or ""

    if not admin_user_exists():
        if secrets.compare_digest(password, Config.ADMIN_PASSWORD):
            request.session["admin_logged_in"] = True
            request.session["admin_username"] = username or "admin"
            return JSONResponse({"success": True, "username": username or "admin"})
        return JSONResponse({"error": "Incorrect password."}, status_code=401)

    if verify_admin_user(username, password):
        request.session["admin_logged_in"] = True
        request.session["admin_username"] = username
        return JSONResponse({"success": True, "username": username})
    return JSONResponse({"error": "Incorrect username or password."}, status_code=401)


@fastapi_app.post("/admin/api/logout")
async def admin_api_logout(request: Request):
    """JSON admin logout endpoint (used by React SPA)."""
    request.session.pop("admin_logged_in", None)
    request.session.pop("admin_username", None)
    return JSONResponse({"success": True})


@fastapi_app.post("/admin/api/register")
async def admin_api_register(request: Request):
    """JSON admin registration endpoint (used by React SPA)."""
    if request.session.get("admin_logged_in"):
        return JSONResponse({"error": "Already logged in."}, status_code=400)
    try:
        body = await request.json()
    except Exception:
        return JSONResponse({"error": "Invalid JSON"}, status_code=400)

    username        = (body.get("username") or "").strip()
    password        = body.get("password") or ""
    confirm_password = body.get("confirm_password") or ""

    if not username or not password:
        return JSONResponse({"error": "Username and password are required."}, status_code=400)
    if password != confirm_password:
        return JSONResponse({"error": "Passwords do not match."}, status_code=400)

    ok, msg = register_admin_user(username, password)
    if ok:
        return JSONResponse({"success": True, "message": "Admin account created. You can now log in."})
    return JSONResponse({"error": msg}, status_code=400)

@fastapi_app.middleware("http")
async def track_admin_visitor(request: Request, call_next):
    """Record a visit to tracked pages (main site + admin page) for analytics."""
    response = await call_next(request)

    if request.url.path not in TRACKED_VISITOR_PATHS:
        return response

    ip = _get_real_ip(request)
    ua = request.headers.get("user-agent", "")

    # Try to get country from CDN/proxy headers first
    hdr_country, hdr_code = _get_country_from_headers(request)
    if hdr_country and ip not in ip_country_cache:
        ip_country_cache[ip] = {"country": hdr_country, "code": hdr_code, "city": "", "region": ""}
    elif ip not in ip_country_cache and _is_private_ip(ip):
        ip_country_cache[ip] = {"country": "Local", "code": "", "city": "", "region": ""}

    cached = ip_country_cache.get(ip, {})
    visitor_entry = {
        "ip": ip,
        "timestamp": time.time(),
        "user_agent": ua,
        "browser": _parse_browser(ua),
        "device": _parse_device(ua),
        "country": cached.get("country", ""),
        "country_code": cached.get("code", ""),
        "city": cached.get("city", ""),
        "region": cached.get("region", ""),
        "page": request.url.path,
    }
    with visitors_lock:
        visitors.append(visitor_entry)
    # Resolve country in the background if not cached yet
    if ip not in ip_country_cache:
        accept_lang = request.headers.get("accept-language", "")
        threading.Thread(
            target=_lookup_country_async, args=(ip, accept_lang), daemon=True
        ).start()
    # Debounced save to avoid a thread-per-visit under high traffic
    _schedule_visitor_save()

    return response


@fastapi_app.get("/admin/downloads")
@admin_required
async def admin_downloads_api(request: Request):
    """Return the complete download history for the admin page.

    Merges in-memory records (which have live progress) with persisted
    database records so that logs survive in-memory eviction and server
    restarts.  In-memory entries take precedence for active downloads.
    """
    _fields = (
        "id", "title", "url", "status", "percent", "filename",
        "file_size_hr", "created_at", "end_time", "error", "ip",
        "country", "country_code", "city", "region", "format", "type",
    )

    def _pick(d: dict) -> dict:
        return {k: d.get(k, "" if k in ("country", "country_code", "city", "region") else None) for k in _fields}

    # Start with all DB records (these persist after in-memory cleanup)
    merged: dict = {}
    try:
        with _db_lock:
            conn = _get_db()
            try:
                rows = _execute(conn, "SELECT id, data FROM downloads").fetchall()
            finally:
                conn.close()
        for row in rows:
            try:
                d = json.loads(row["data"])
                did = d.get("id") or row["id"]
                merged[did] = _pick(d)
            except Exception:
                pass
    except Exception as exc:
        logger.error(f"admin_downloads_api DB read error: {exc}")

    # Overlay in-memory records (they have live progress for active downloads)
    with downloads_lock:
        for did, d in downloads.items():
            merged[did] = _pick(d)

    history = sorted(merged.values(), key=lambda x: x.get("created_at") or 0, reverse=True)
    return JSONResponse(history)


@fastapi_app.get("/admin/visitors")
@admin_required
async def admin_visitors_api(request: Request):
    """Return visitor analytics for the admin page."""
    with visitors_lock:
        data = list(visitors)
    data.sort(key=lambda x: x.get("timestamp", 0), reverse=True)
    return JSONResponse(data)


def _get_persistent_stats() -> dict:
    """Query the database for all-time aggregate stats (accurate even after in-memory cleanup).

    Returns a dict with keys:
        total_downloads, daily_downloads, download_rate_per_day,
        total_site_visitors, daily_site_visitors
    """
    now = datetime.now()
    today_start = datetime.combine(now.date(), datetime.min.time()).timestamp()
    week_start = (now - timedelta(days=7)).timestamp()

    try:
        with _db_lock:
            conn = _get_db()
            try:
                # Downloads: parse created_at from the JSON blob
                dl_rows = _execute(conn, "SELECT data FROM downloads").fetchall()
                total_downloads = len(dl_rows)
                daily_downloads = 0
                week_downloads = 0
                for row in dl_rows:
                    try:
                        d = json.loads(row["data"])
                        created = d.get("created_at") or 0
                        if created >= today_start:
                            daily_downloads += 1
                        if created >= week_start:
                            week_downloads += 1
                    except Exception:
                        pass

                # Visitors: only count main-site visits (page == "/")
                # Legacy records without a "page" key were admin-page visits, so skip them.
                v_rows = _execute(conn, "SELECT data FROM visitors").fetchall()
                total_site_visitors = 0
                daily_site_visitors = 0
                for row in v_rows:
                    try:
                        v = json.loads(row["data"])
                        if v.get("page") == "/":
                            total_site_visitors += 1
                            if v.get("timestamp", 0) >= today_start:
                                daily_site_visitors += 1
                    except Exception:
                        pass
            finally:
                conn.close()
    except Exception as exc:
        logger.error(f"_get_persistent_stats error: {exc}")
        total_downloads = 0
        daily_downloads = 0
        week_downloads = 0
        total_site_visitors = 0
        daily_site_visitors = 0

    download_rate_per_day = round(week_downloads / 7.0, 2)

    return {
        "total_downloads": total_downloads,
        "daily_downloads": daily_downloads,
        "download_rate_per_day": download_rate_per_day,
        "total_site_visitors": total_site_visitors,
        "daily_site_visitors": daily_site_visitors,
    }


@fastapi_app.get("/admin/analytics")
@admin_required
async def admin_analytics_api(request: Request):
    """Return aggregated analytics including country totals, persistent stats,
    and deep-insight metrics (peak hours, format preferences, OS breakdown,
    success rate, repeat visitors, average file size).
    """
    # --- Collect full download list from DB + in-memory overlay ---
    db_downloads: dict = {}
    try:
        with _db_lock:
            conn = _get_db()
            try:
                rows = _execute(conn, "SELECT id, data FROM downloads").fetchall()
            finally:
                conn.close()
        for row in rows:
            try:
                d = json.loads(row["data"])
                db_downloads[d.get("id") or row["id"]] = d
            except Exception:
                pass
    except Exception:
        pass
    with downloads_lock:
        for did, d in downloads.items():
            db_downloads[did] = d
    dl_list = list(db_downloads.values())

    with visitors_lock:
        v_list = list(visitors)

    # Country totals for downloads
    dl_countries: dict = {}
    for d in dl_list:
        country = d.get("country") or "Unknown"
        code = d.get("country_code") or ""
        key = f"{country}||{code}"
        dl_countries[key] = dl_countries.get(key, 0) + 1

    # Country totals for visitors (main-site and admin page combined)
    v_countries: dict = {}
    for v in v_list:
        country = v.get("country") or "Unknown"
        code = v.get("country_code") or ""
        key = f"{country}||{code}"
        v_countries[key] = v_countries.get(key, 0) + 1

    # Unique countries represented across both downloads and visitors
    all_country_names = set()
    for k in list(dl_countries.keys()) + list(v_countries.keys()):
        name = k.split("||")[0]
        if name and name != "Unknown":
            all_country_names.add(name)

    # Persistent aggregate stats (queried from the database for accuracy)
    persistent = _get_persistent_stats()

    # --- Deep-insight analytics ---

    # 1. Peak download hours (0-23)
    peak_hours = [0] * 24
    for d in dl_list:
        ts = d.get("created_at")
        if ts:
            try:
                peak_hours[datetime.fromtimestamp(float(ts)).hour] += 1
            except Exception:
                pass

    # 2. Peak visitor hours (0-23)
    visitor_hours = [0] * 24
    for v in v_list:
        ts = v.get("timestamp")
        if ts:
            try:
                visitor_hours[datetime.fromtimestamp(float(ts)).hour] += 1
            except Exception:
                pass

    # 3. Format preferences (from download records)
    format_counts: dict = {}
    for d in dl_list:
        fmt = d.get("format") or "unknown"
        # Simplify long format strings to a readable label
        label = fmt.split("/")[0].split("+")[0].strip()[:30] if fmt else "unknown"
        format_counts[label] = format_counts.get(label, 0) + 1

    # 4. Success / failure rate
    total_dl = len(dl_list)
    completed_count = sum(1 for d in dl_list if d.get("status") == "completed")
    failed_count = sum(1 for d in dl_list if d.get("status") == "failed")
    cancelled_count = sum(1 for d in dl_list if d.get("status") == "cancelled")
    success_rate = round(100.0 * completed_count / total_dl, 1) if total_dl else 0

    # 5. Average file size (completed downloads only)
    sizes = []
    for d in dl_list:
        if d.get("status") == "completed":
            sz = d.get("file_size")
            if sz and isinstance(sz, (int, float)) and sz > 0:
                sizes.append(sz)
    avg_file_size = round(sum(sizes) / len(sizes)) if sizes else 0
    avg_file_size_hr = format_size(avg_file_size) if avg_file_size else "—"

    # 6. OS / device breakdown (from visitor user-agents)
    os_counts: dict = {}
    for v in v_list:
        ua = (v.get("user_agent") or "").lower()
        if "windows" in ua:
            os_name = "Windows"
        elif "macintosh" in ua or "mac os" in ua:
            os_name = "macOS"
        elif "android" in ua:
            os_name = "Android"
        elif "iphone" in ua or "ipad" in ua:
            os_name = "iOS"
        elif "linux" in ua:
            os_name = "Linux"
        elif "cros" in ua:
            os_name = "ChromeOS"
        else:
            os_name = "Other"
        os_counts[os_name] = os_counts.get(os_name, 0) + 1

    # 7. Device type breakdown (from visitor user-agents)
    device_counts: dict = {}
    for v in v_list:
        device = v.get("device") or _parse_device(v.get("user_agent") or "")
        device_counts[device] = device_counts.get(device, 0) + 1

    # 8. Repeat visitors (IPs seen more than once on the main site)
    ip_visit_counts: dict = {}
    for v in v_list:
        if v.get("page") == "/":
            ip_visit_counts[v.get("ip", "")] = ip_visit_counts.get(v.get("ip", ""), 0) + 1
    unique_ips = len(ip_visit_counts)
    repeat_ips = sum(1 for c in ip_visit_counts.values() if c > 1)
    repeat_rate = round(100.0 * repeat_ips / unique_ips, 1) if unique_ips else 0

    # 8. Downloads by day-of-week (Mon=0 … Sun=6)
    dow_downloads = [0] * 7
    for d in dl_list:
        ts = d.get("created_at")
        if ts:
            try:
                dow_downloads[datetime.fromtimestamp(float(ts)).weekday()] += 1
            except Exception:
                pass

    # 9. Review rating breakdown
    with reviews_lock:
        review_ratings = [0] * 5  # index 0 = 1-star, index 4 = 5-star
        for r in reviews:
            rt = r.get("rating")
            if isinstance(rt, int) and 1 <= rt <= 5:
                review_ratings[rt - 1] += 1
        total_reviews = len(reviews)

    return JSONResponse({
        "download_countries": [
            {"country": k.split("||")[0], "code": k.split("||")[1], "count": v}
            for k, v in sorted(dl_countries.items(), key=lambda x: x[1], reverse=True)
        ],
        "visitor_countries": [
            {"country": k.split("||")[0], "code": k.split("||")[1], "count": v}
            for k, v in sorted(v_countries.items(), key=lambda x: x[1], reverse=True)
        ],
        "unique_countries_total": len(all_country_names),
        "total_downloads": persistent["total_downloads"],
        "daily_downloads": persistent["daily_downloads"],
        "download_rate_per_day": persistent["download_rate_per_day"],
        "total_site_visitors": persistent["total_site_visitors"],
        "daily_site_visitors": persistent["daily_site_visitors"],
        # Deep insights
        "peak_hours": peak_hours,
        "visitor_hours": visitor_hours,
        "format_preferences": [
            {"format": k, "count": v}
            for k, v in sorted(format_counts.items(), key=lambda x: x[1], reverse=True)[:15]
        ],
        "success_rate": success_rate,
        "completed_count": completed_count,
        "failed_count": failed_count,
        "cancelled_count": cancelled_count,
        "avg_file_size": avg_file_size,
        "avg_file_size_hr": avg_file_size_hr,
        "os_breakdown": [
            {"os": k, "count": v}
            for k, v in sorted(os_counts.items(), key=lambda x: x[1], reverse=True)
        ],
        "device_breakdown": [
            {"device": k, "count": v}
            for k, v in sorted(device_counts.items(), key=lambda x: x[1], reverse=True)
        ],
        "unique_visitors": unique_ips,
        "repeat_visitors": repeat_ips,
        "repeat_rate": repeat_rate,
        "dow_downloads": dow_downloads,
        "review_ratings": review_ratings,
        "total_reviews": total_reviews,
    })


@fastapi_app.post("/admin/cancel_download/{download_id}")
@admin_required
async def admin_cancel_download(request: Request, download_id: str):
    """Cancel an active download (admin only)."""
    with downloads_lock:
        if download_id in downloads:
            if downloads[download_id]["status"] in ("starting", "fetching_info", "queued", "downloading"):
                downloads[download_id]["status"] = "cancelled"
                downloads[download_id]["end_time"] = time.time()
                emit_from_thread("cancelled", {"id": download_id}, room=download_id)
                logger.info(f"Admin cancelled download: {download_id}")
                threading.Thread(target=save_downloads_to_disk, daemon=True).start()
                return JSONResponse({"success": True})
            return JSONResponse({"error": "Download is not active"}, status_code=409)
    return JSONResponse({"error": "Download not found"}, status_code=404)


@fastapi_app.delete("/admin/delete_record/{download_id}")
@admin_required
async def admin_delete_record(request: Request, download_id: str):
    """Remove a download record from the history (admin only).
    Active downloads are automatically cancelled before deletion."""
    with downloads_lock:
        if download_id in downloads:
            status = downloads[download_id].get("status")
            if status in ("starting", "fetching_info", "queued", "downloading"):
                downloads[download_id]["status"] = "cancelled"
                downloads[download_id]["end_time"] = time.time()
                emit_from_thread("cancelled", {"id": download_id}, room=download_id)
                logger.info(f"Admin force-cancelled download for deletion: {download_id}")
            del downloads[download_id]

    # Delete from database as well so the record is permanently removed
    def _db_delete():
        try:
            with _db_lock:
                conn = _get_db()
                try:
                    _execute(conn, "DELETE FROM downloads WHERE id = ?", (download_id,))
                    conn.commit()
                finally:
                    conn.close()
        except Exception as exc:
            logger.error(f"DB delete error for {download_id}: {exc}")
    threading.Thread(target=_db_delete, daemon=True).start()
    logger.info(f"Admin deleted download record: {download_id}")
    return JSONResponse({"success": True})


@fastapi_app.delete("/admin/clear_all_downloads")
@admin_required
async def admin_clear_all_downloads(request: Request):
    """Clear ALL download records from both in-memory store and database (admin only).

    Active and queued downloads are cancelled before the records are removed so
    that in-flight workers receive the cancellation signal.
    """
    cancelled_ids: list[str] = []
    with downloads_lock:
        for did, d in list(downloads.items()):
            status = d.get("status")
            if status in ("starting", "fetching_info", "queued", "downloading"):
                cancelled_ids.append(did)
        downloads.clear()

    # Emit cancellation signals outside the lock to avoid deadlocks
    for did in cancelled_ids:
        emit_from_thread("cancelled", {"id": did}, room=did)

    def _db_clear():
        try:
            with _db_lock:
                conn = _get_db()
                try:
                    _execute(conn, "DELETE FROM downloads")
                    conn.commit()
                finally:
                    conn.close()
        except Exception as exc:
            logger.error(f"admin_clear_all_downloads DB error: {exc}")

    threading.Thread(target=_db_clear, daemon=True).start()
    logger.info(f"Admin cleared ALL download records ({len(cancelled_ids)} active/queued cancelled)")
    return JSONResponse({"success": True})


@fastapi_app.delete("/admin/clear_visitors")
@admin_required
async def admin_clear_visitors(request: Request):
    """Clear all visitor records (admin only)."""
    with visitors_lock:
        visitors.clear()
    threading.Thread(target=save_visitors_to_disk, daemon=True).start()
    logger.info("Admin cleared all visitor records")
    return JSONResponse({"success": True})


@fastapi_app.get("/admin/db/download")
@admin_required
async def admin_db_download(request: Request):
    """Download a database backup (admin only).

    When using PostgreSQL the backup is a JSON file; otherwise the raw SQLite
    file is returned.
    """
    # Flush in-memory state to disk before serving
    save_downloads_to_disk()
    save_visitors_to_disk()
    logger.info("Admin downloaded database backup")

    if USE_POSTGRES:
        # Export PostgreSQL data as JSON
        with _db_lock:
            conn = _get_db()
            try:
                dl_rows = _execute(conn, "SELECT id, data FROM downloads").fetchall()
                v_rows = _execute(conn, "SELECT data FROM visitors ORDER BY id").fetchall()
            finally:
                conn.close()
        backup = {
            "downloads": {r["id"]: json.loads(r["data"]) for r in dl_rows},
            "visitors": [json.loads(r["data"]) for r in v_rows],
        }
        return Response(
            content=json.dumps(backup, indent=2, default=str),
            media_type="application/json",
            headers={"Content-Disposition": "attachment; filename=admin_backup.json"},
        )
    else:
        if not os.path.exists(DB_PATH):
            return JSONResponse({"error": "Database file not found"}, status_code=404)
        return FileResponse(
            DB_PATH,
            filename="admin.db",
            media_type="application/x-sqlite3",
        )


@fastapi_app.post("/admin/db/upload")
@admin_required
async def admin_db_upload(request: Request):
    """Merge an uploaded backup into the live database (admin only).

    Accepts either a SQLite ``.db`` file (always supported for migration) or a
    JSON backup produced by the PostgreSQL download endpoint.  Downloads are
    merged using an *ignore-on-conflict* strategy (live records take
    precedence).  Visitors from both sources are combined, deduplicated by JSON
    content, and re-sorted chronologically.  After a successful merge the
    in-memory state is reloaded.
    """
    form = await request.form()
    db_file = form.get("db_file")
    if db_file is None:
        return JSONResponse({"error": "No file uploaded"}, status_code=400)
    if not hasattr(db_file, "read"):
        return JSONResponse({"error": "No file selected"}, status_code=400)

    content = await db_file.read()

    # ── Detect format ────────────────────────────────────────
    is_sqlite = len(content) >= 16 and content[:16] == b"SQLite format 3\x00"
    is_json = False
    backup_dl: list[tuple[str, str]] = []
    backup_visitors: list[str] = []

    if not is_sqlite:
        # Try to parse as JSON backup
        _bad_backup_msg = "Uploaded file is not a valid backup (expected SQLite or JSON)"
        try:
            parsed = json.loads(content)
            if isinstance(parsed, dict) and ("downloads" in parsed or "visitors" in parsed):
                is_json = True
                dl_data = parsed.get("downloads", {})
                for did, d in dl_data.items():
                    backup_dl.append((did, json.dumps(d, default=str)))
                v_data = parsed.get("visitors", [])
                for v in v_data:
                    backup_visitors.append(json.dumps(v, default=str))
            else:
                return JSONResponse({"error": _bad_backup_msg}, status_code=400)
        except (json.JSONDecodeError, ValueError):
            return JSONResponse({"error": _bad_backup_msg}, status_code=400)

    if is_sqlite:
        # Parse the SQLite backup file
        tmp_path = os.path.join(DATA_DIR, "upload_tmp.db")
        try:
            with open(tmp_path, "wb") as fh:
                fh.write(content)

            check_conn = sqlite3.connect(tmp_path)
            check_conn.row_factory = sqlite3.Row
            try:
                result = check_conn.execute("PRAGMA integrity_check").fetchone()
                if result[0] != "ok":
                    return JSONResponse({"error": "Uploaded database failed integrity check"}, status_code=400)
                try:
                    backup_dl_rows = check_conn.execute("SELECT id, data FROM downloads").fetchall()
                    backup_dl = [(r["id"], r["data"]) for r in backup_dl_rows]
                except sqlite3.OperationalError:
                    backup_dl = []
                try:
                    backup_v_rows = check_conn.execute("SELECT data FROM visitors ORDER BY rowid").fetchall()
                    backup_visitors = [r["data"] for r in backup_v_rows]
                except sqlite3.OperationalError:
                    backup_visitors = []
            finally:
                check_conn.close()
        finally:
            if os.path.exists(tmp_path):
                try:
                    os.remove(tmp_path)
                except OSError:
                    pass

    # ── Merge into live database ─────────────────────────────
    try:
        save_downloads_to_disk()
        save_visitors_to_disk()

        with _db_lock:
            conn = _get_db()
            try:
                for (did, data_json) in backup_dl:
                    if USE_POSTGRES:
                        cur = conn.cursor()
                        cur.execute(
                            "INSERT INTO downloads (id, data) VALUES (%s, %s) ON CONFLICT (id) DO NOTHING",
                            (did, data_json),
                        )
                    else:
                        conn.execute(
                            "INSERT OR IGNORE INTO downloads (id, data) VALUES (?, ?)",
                            (did, data_json),
                        )

                live_v_rows = _execute(conn, "SELECT data FROM visitors").fetchall()
                live_v_jsons = [r["data"] for r in live_v_rows]

                seen_visitor_jsons: set[str] = set()
                all_visitors = []
                for v_json in live_v_jsons + backup_visitors:
                    if v_json not in seen_visitor_jsons:
                        seen_visitor_jsons.add(v_json)
                        try:
                            all_visitors.append(json.loads(v_json))
                        except (json.JSONDecodeError, ValueError):
                            pass

                all_visitors.sort(key=lambda v: v.get("timestamp", 0))

                _execute(conn, "DELETE FROM visitors")
                for v in all_visitors:
                    _execute(
                        conn,
                        "INSERT INTO visitors (data) VALUES (?)",
                        (json.dumps(v, default=str),),
                    )

                conn.commit()
            finally:
                conn.close()

        global downloads, visitors
        with downloads_lock:
            downloads.clear()
        with visitors_lock:
            visitors.clear()
        load_persistence()

        added_dl = len(backup_dl)
        added_v = len(backup_visitors)
        logger.info(
            f"Admin merged database backup: {added_dl} download records and "
            f"{added_v} visitor records from backup processed"
        )
        return JSONResponse({
            "success": True,
            "message": (
                f"Database merged successfully. "
                f"Processed {added_dl} download record(s) and "
                f"{added_v} visitor record(s) from backup."
            ),
        })
    except Exception as exc:
        logger.error(f"DB upload error: {exc}")
        return JSONResponse({"error": f"Upload failed: {exc}"}, status_code=500)


# ── Cookie management (admin) ─────────────────────────────
@fastapi_app.get("/admin/cookies/status")
@admin_required
async def admin_cookies_status(request: Request):
    """Return whether a cookies file is currently configured."""
    exists = os.path.isfile(COOKIES_FILE)
    info: dict = {"has_cookies": exists}
    if exists:
        info["size"] = os.path.getsize(COOKIES_FILE)
        info["modified"] = os.path.getmtime(COOKIES_FILE)
    return JSONResponse(info)


@fastapi_app.post("/admin/cookies/upload")
@admin_required
async def admin_cookies_upload(request: Request):
    """Upload a Netscape-format cookies.txt file for yt-dlp."""
    form = await request.form()
    cookie_file = form.get("cookie_file")
    if cookie_file is None or not hasattr(cookie_file, "read"):
        return JSONResponse({"error": "No file uploaded"}, status_code=400)

    content = await cookie_file.read()
    if not content:
        return JSONResponse({"error": "Uploaded file is empty"}, status_code=400)

    # Basic validation: Netscape cookies files start with a comment or domain
    text = content.decode("utf-8", errors="replace")
    lines = [s for ln in text.splitlines() if (s := ln.strip()) and not s.startswith("#")]
    if not lines:
        return JSONResponse({"error": "Cookies file appears to be empty (no cookie entries)"}, status_code=400)

    os.makedirs(os.path.dirname(COOKIES_FILE), exist_ok=True)
    with open(COOKIES_FILE, "wb") as fh:
        fh.write(content)
    logger.info("Admin uploaded new cookies file (%d bytes)", len(content))
    return JSONResponse({"success": True, "message": "Cookies file uploaded successfully."})


@fastapi_app.delete("/admin/cookies")
@admin_required
async def admin_cookies_delete(request: Request):
    """Delete the current cookies file."""
    if os.path.isfile(COOKIES_FILE):
        os.remove(COOKIES_FILE)
        logger.info("Admin deleted cookies file")
        return JSONResponse({"success": True, "message": "Cookies file deleted."})
    return JSONResponse({"error": "No cookies file found"}, status_code=404)


@fastapi_app.post("/download_zip")
async def download_zip(
    request: Request,
    filenames: str = Form("[]"),
):
    """Package a selection of downloaded files into a ZIP and serve it."""
    filenames_json = filenames

    try:
        parsed_filenames = json.loads(filenames_json)
    except (json.JSONDecodeError, ValueError):
        return JSONResponse({"error": "Invalid filenames JSON"}, status_code=400)

    if not isinstance(parsed_filenames, list) or len(parsed_filenames) == 0:
        return JSONResponse({"error": "No files selected"}, status_code=400)

    filepairs = []
    for fn in parsed_filenames:
        if not fn:
            continue
        fp = os.path.join(DOWNLOAD_FOLDER, fn)
        if not os.path.abspath(fp).startswith(os.path.abspath(DOWNLOAD_FOLDER)):
            continue
        if os.path.isfile(fp):
            filepairs.append((fn, fp))

    if not filepairs:
        return JSONResponse({"error": "No valid files found"}, status_code=404)

    zip_filename = f"downloads_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip"
    zip_path     = os.path.join(DOWNLOAD_FOLDER, zip_filename)
    try:
        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            for fn, fp in filepairs:
                zf.write(fp, fn)
    except Exception as exc:
        logger.error("ZIP creation error: %s", exc)
        return JSONResponse({"error": f"Failed to create ZIP: {exc}"}, status_code=500)

    return FileResponse(
        zip_path,
        filename=zip_filename,
        media_type="application/zip",
    )


# =========================================================
# CV GENERATION MODULE
# =========================================================

_MAX_CV_LOGO_BYTES = 5 * 1024 * 1024  # 5 MB max for CV logo uploads
_TEMP_DIR_CLEANUP_DELAY_SECS = 60     # seconds before temp conversion dirs are removed

# ---------------------------------------------------------------------------
# Pure-Python PDF builder (fpdf2).  No LaTeX / LibreOffice required.
# ---------------------------------------------------------------------------
def _build_cv_pdf(
    output_path: str,
    *,
    name: str,
    email: str,
    phone: str = "",
    location: str = "",
    link: str = "",
    summary: str = "",
    experience: str = "",
    education: str = "",
    skills: str = "",
    projects: str = "",
    publications: str = "",
    logo_path: str = "",
    theme: str = "classic",
    layout: str = "chronological",
) -> None:
    """Build a professional single-file PDF CV using fpdf2.

    Supported themes: 'classic' (blue), 'modern' (dark header), 'minimal' (B&W),
    'executive' (navy/gold).

    ``layout`` controls section order:
    - ``chronological`` – experience first, then education, then skills
    - ``functional``    – skills first (prominently), then experience, then education
    """
    from fpdf import FPDF, XPos, YPos

    _NL = dict(new_x=XPos.LMARGIN, new_y=YPos.NEXT)  # replaces deprecated ln=True

    # ---- Colour palettes ----
    _THEMES = {
        "classic": {
            "dark":   (30,  30,  30),
            "accent": (37,  99, 235),   # blue-600
            "light":  (107, 114, 128),  # gray-500
            "header_bg": None,          # no filled header band
            "header_fg": (30, 30, 30),
        },
        "modern": {
            "dark":   (20,  20,  20),
            "accent": (15,  23,  42),   # slate-900 heading bar
            "light":  (100, 116, 139),  # slate-500
            "header_bg": (15, 23, 42),  # filled dark band
            "header_fg": (255, 255, 255),
        },
        "minimal": {
            "dark":   (0,   0,   0),
            "accent": (0,   0,   0),    # pure black
            "light":  (120, 120, 120),
            "header_bg": None,
            "header_fg": (0, 0, 0),
        },
        "executive": {
            "dark":   (22,  36,  71),   # deep navy
            "accent": (22,  36,  71),
            "light":  (100, 100, 130),
            "header_bg": (22, 36, 71),  # navy band
            "header_fg": (212, 175, 55),  # gold text
        },
        "creative": {
            "dark":   (30,  10,  60),   # deep purple
            "accent": (124, 58, 237),   # violet-600
            "light":  (139, 92, 246),   # violet-400
            "header_bg": (124, 58, 237),  # violet band
            "header_fg": (255, 255, 255),
        },
        "tech": {
            "dark":   (15,  23,  42),   # slate-900
            "accent": (16, 185, 129),   # emerald-500
            "light":  (100, 116, 139),  # slate-400
            "header_bg": (15, 23, 42),  # dark band
            "header_fg": (16, 185, 129),  # emerald text
        },
        "elegant": {
            "dark":   (30,  10,  20),   # very dark burgundy
            "accent": (157, 23,  77),   # rose-800 / burgundy
            "light":  (156, 100, 120),
            "header_bg": None,
            "header_fg": (157, 23, 77),
        },
        "vibrant": {
            "dark":   (20,  20,  20),
            "accent": (234, 88,  12),   # orange-600
            "light":  (107, 114, 128),
            "header_bg": (234, 88, 12),  # orange band
            "header_fg": (255, 255, 255),
        },
    }
    t = _THEMES.get(theme, _THEMES["classic"])
    DARK      = t["dark"]
    ACCENT    = t["accent"]
    LIGHT     = t["light"]
    HEADER_BG = t["header_bg"]
    HEADER_FG = t["header_fg"]

    # ---- Unicode font setup ----
    # Prefer DejaVu Sans (TTF) for full Unicode support so accented letters,
    # CJK characters, special symbols, etc. render correctly instead of showing
    # "?" (which happens with the Latin-1-only Helvetica core font).
    # Fall back to Helvetica + a best-effort transliteration map when the font
    # files are not present on the server.
    #
    # Check several common font locations across Linux distributions.
    _DEJAVU_CANDIDATES = [
        "/usr/share/fonts/truetype/dejavu",   # Debian/Ubuntu
        "/usr/share/fonts/dejavu",            # Fedora/RHEL
        "/usr/share/fonts/TTF",               # Arch Linux
        "/usr/share/fonts/dejavu-sans-fonts", # some Alpine variants
    ]
    _DEJAVU_DIR = next(
        (d for d in _DEJAVU_CANDIDATES if os.path.isdir(d)),
        "/usr/share/fonts/truetype/dejavu",   # default (may not exist)
    )
    _DEJAVU_REG   = os.path.join(_DEJAVU_DIR, "DejaVuSans.ttf")
    _DEJAVU_BOLD  = os.path.join(_DEJAVU_DIR, "DejaVuSans-Bold.ttf")
    _DEJAVU_OBLI  = os.path.join(_DEJAVU_DIR, "DejaVuSans-Oblique.ttf")
    _use_unicode  = all(os.path.isfile(p) for p in (_DEJAVU_REG, _DEJAVU_BOLD, _DEJAVU_OBLI))

    # Latin-1 fallback: map common Unicode punctuation/symbols to ASCII equivalents.
    _UNICODE_MAP = str.maketrans({
        # Dashes and hyphens
        "\u2013": " - ",   # en dash
        "\u2014": " - ",   # em dash
        "\u2015": " - ",   # horizontal bar
        "\u2010": "-",     # hyphen
        "\u2011": "-",     # non-breaking hyphen
        "\u2012": "-",     # figure dash
        "\u2212": "-",     # minus sign
        # Bullets and dots
        "\u2022": "*",     # bullet •
        "\u2023": "*",     # triangular bullet
        "\u25E6": "*",     # white bullet
        "\u2043": "-",     # hyphen bullet
        "\u00B7": "*",     # middle dot
        "\u2027": ".",     # hyphenation point
        # Quotation marks
        "\u2018": "'",     # left single quote
        "\u2019": "'",     # right single quote
        "\u201A": ",",     # single low-9 quotation
        "\u201B": "'",     # single high-reversed-9 quotation
        "\u201C": '"',     # left double quote
        "\u201D": '"',     # right double quote
        "\u201E": '"',     # double low-9 quotation
        "\u201F": '"',     # double high-reversed-9 quotation
        "\u2039": "'",     # single left angle quotation
        "\u203A": "'",     # single right angle quotation
        "\u00AB": '"',     # left guillemet
        "\u00BB": '"',     # right guillemet
        "\u2032": "'",     # prime
        "\u2033": '"',     # double prime
        "\u2034": "'''",   # triple prime
        # Spaces and whitespace
        "\u00A0": " ",     # non-breaking space
        "\u2002": " ",     # en space
        "\u2003": " ",     # em space
        "\u2004": " ",     # three-per-em space
        "\u2005": " ",     # four-per-em space
        "\u2009": " ",     # thin space
        "\u200A": " ",     # hair space
        "\u202F": " ",     # narrow no-break space
        "\u205F": " ",     # medium mathematical space
        "\u3000": " ",     # ideographic space
        "\u200B": "",      # zero-width space
        "\u200C": "",      # zero-width non-joiner
        "\u200D": "",      # zero-width joiner
        "\uFEFF": "",      # BOM / zero-width no-break space
        # Ellipsis and similar
        "\u2026": "...",   # ellipsis
        "\u22EF": "...",   # midline horizontal ellipsis
        # Miscellaneous symbols
        "\u2122": "(TM)",  # trademark sign
        "\u00AE": "(R)",   # registered sign
        "\u00A9": "(C)",   # copyright sign
        "\u2117": "(P)",   # sound recording copyright
        "\u2020": "+",     # dagger
        "\u2021": "++",    # double dagger
        "\u00B0": " deg",  # degree sign
        "\u00B1": "+/-",   # plus-minus sign
        "\u00D7": "x",     # multiplication sign
        "\u00F7": "/",     # division sign
        "\u2248": "~",     # almost equal
        "\u2260": "!=",    # not equal to
        "\u2264": "<=",    # less-than or equal
        "\u2265": ">=",    # greater-than or equal
        "\u221E": "inf",   # infinity
        "\u25B6": ">",     # play button / right-pointing triangle
        "\u2713": "v",     # check mark
        "\u2714": "v",     # heavy check mark
        "\u2717": "x",     # ballot x
        "\u2718": "x",     # heavy ballot x
    })

    import unicodedata as _udata

    def _safe(text: str) -> str:
        """Transliterate *text* to Latin-1, replacing unmapped chars with their
        closest ASCII equivalent using Unicode decomposition.

        Used only when DejaVu TTF fonts are unavailable and fpdf2 falls back to
        the Helvetica core font (Latin-1 encoding).  Accented letters like
        é, ü, ñ are handled directly by Latin-1; characters above U+00FF are
        first decomposed (NFKD) so their base letter is preserved ("Ż" → "Z"),
        avoiding spurious "?" placeholders.
        """
        # Step 1: apply explicit symbol-to-ASCII substitutions.
        text = text.translate(_UNICODE_MAP)
        # Step 2: NFKD-decompose and drop combining (non-spacing) marks so that
        # accented letters outside Latin-1 (e.g. Polish ż→z, ł→l, Czech ě→e)
        # are represented by their base character rather than "?".
        result = []
        for ch in text:
            if ord(ch) > 0xFF:
                # Try decomposition first
                normalized = _udata.normalize("NFKD", ch)
                for c in normalized:
                    if _udata.category(c) == "Mn":
                        continue  # skip combining marks
                    if ord(c) <= 0xFF:
                        result.append(c)
                    else:
                        result.append("?")
            else:
                result.append(ch)
        return "".join(result)

    def _t(text: str) -> str:
        """Pass *text* through unchanged (Unicode fonts handle it natively),
        or apply the Latin-1 safe conversion when falling back to Helvetica."""
        return text if _use_unicode else _safe(text)

    class CV(FPDF):
        def header(self):
            pass  # custom header drawn in body
        def footer(self):
            self.set_y(-12)
            self.set_font(_FONT, "I", 8)
            self.set_text_color(*LIGHT)
            self.cell(0, 8, f"Page {self.page_no()}", align="C")

    pdf = CV(orientation="P", unit="mm", format="A4")

    # Register the TTF font family so fpdf2 can embed it in the PDF output.
    if _use_unicode:
        pdf.add_font("DejaVu", "",   _DEJAVU_REG)
        pdf.add_font("DejaVu", "B",  _DEJAVU_BOLD)
        pdf.add_font("DejaVu", "I",  _DEJAVU_OBLI)
        pdf.add_font("DejaVu", "BI", _DEJAVU_OBLI)
        _FONT = "DejaVu"
    else:
        _FONT = "Helvetica"

    pdf.set_auto_page_break(auto=True, margin=18)
    pdf.add_page()
    pdf.set_margins(18, 18, 18)

    page_w = pdf.w - 36  # usable width (A4 210 mm − 18 mm × 2)

    # ---- Helper: section divider ----
    def section_heading(title: str):
        pdf.ln(4)
        pdf.set_font(_FONT, "B", 10)
        pdf.set_text_color(*ACCENT)
        pdf.cell(0, 7, _t(title.upper()), **_NL)
        pdf.set_draw_color(*ACCENT)
        pdf.set_line_width(0.4)
        pdf.line(18, pdf.get_y(), 18 + page_w, pdf.get_y())
        pdf.ln(2)
        pdf.set_text_color(*DARK)

    # ---- Helper: wrap long text ----
    def multi(txt: str, font_size: int = 9, style: str = "", indent: float = 0):
        pdf.set_font(_FONT, style, font_size)
        pdf.set_text_color(*DARK)
        if indent:
            pdf.set_x(18 + indent)
        pdf.multi_cell(page_w - indent, 5, _t(txt), **_NL)

    # ================================================================
    # HEADER BLOCK
    # ================================================================
    # Themed header: themes with HEADER_BG get a filled colour band.
    header_top_y = 14
    header_h = 28 + (5 if link else 0) + (5 if [p for p in (email, phone, location) if p] else 0)

    if HEADER_BG:
        pdf.set_fill_color(*HEADER_BG)
        pdf.rect(0, header_top_y - 2, pdf.w, header_h + 4, style="F")

    logo_w = 0.0
    if logo_path and os.path.isfile(logo_path):
        try:
            logo_w = 22.0
            pdf.image(logo_path, x=18, y=header_top_y, w=logo_w)
        except Exception:
            logo_w = 0.0

    x_text = 18 + (logo_w + 4 if logo_w else 0)
    w_text = page_w - (logo_w + 4 if logo_w else 0)

    name_fg = HEADER_FG if HEADER_BG else DARK
    pdf.set_xy(x_text, header_top_y)
    pdf.set_font(_FONT, "B", 20)
    pdf.set_text_color(*name_fg)
    pdf.cell(w_text, 10, _t(name), **_NL)

    # Contact line
    contact_parts = [p for p in (email, phone, location) if p]
    if contact_parts:
        pdf.set_x(x_text)
        pdf.set_font(_FONT, "", 9)
        contact_fg = HEADER_FG if HEADER_BG else LIGHT
        pdf.set_text_color(*contact_fg)
        pdf.cell(w_text, 5, _t("  |  ".join(contact_parts)), **_NL)

    if link:
        pdf.set_x(x_text)
        pdf.set_font(_FONT, "", 9)
        link_fg = HEADER_FG if HEADER_BG else ACCENT
        pdf.set_text_color(*link_fg)
        pdf.cell(w_text, 5, _t(link), **_NL)

    # Move below the header band
    pdf.ln(3)
    if HEADER_BG:
        # Ensure we're past the band before drawing the separator
        band_bottom = header_top_y + header_h + 4
        if pdf.get_y() < band_bottom:
            pdf.set_y(band_bottom)
    pdf.set_draw_color(*ACCENT)
    pdf.set_line_width(0.8)
    pdf.line(18, pdf.get_y(), 18 + page_w, pdf.get_y())
    pdf.ln(4)
    pdf.set_text_color(*DARK)

    # ================================================================
    # SUMMARY
    # ================================================================
    if summary.strip():
        section_heading("Professional Summary")
        multi(summary.strip())

    # ---- helpers for optional sections ----
    def _block_lines(block: str) -> list[str]:
        """Return non-empty stripped lines from an experience/education block."""
        return [l.strip() for l in block.split('\n') if l.strip()]

    def _render_skills():
        skill_list = [s.strip() for s in skills.split(",") if s.strip()]
        if not skill_list:
            return
        section_heading("Skills")
        pdf.set_font(_FONT, "", 9)
        pdf.set_text_color(*DARK)
        row_size = 5
        for i in range(0, len(skill_list), row_size):
            row = skill_list[i:i + row_size]
            pdf.set_x(18)
            pdf.cell(0, 5, _t("  *  ".join(row)), **_NL)
        pdf.ln(1)

    def _render_experience():
        if not experience.strip():
            return
        section_heading("Work Experience")
        for block in re.split(r'\n\s*\n', experience.strip()):
            block_lines = _block_lines(block)
            if not block_lines:
                continue
            pdf.set_font(_FONT, "B", 9)
            pdf.set_text_color(*DARK)
            pdf.multi_cell(page_w, 5, _t(block_lines[0]), **_NL)
            for bullet in block_lines[1:]:
                bullet_text = bullet.lstrip("*-\u2022\u2013\u2014 ").strip()
                if bullet_text:
                    pdf.set_font(_FONT, "", 9)
                    pdf.set_text_color(*DARK)
                    pdf.set_x(22)
                    pdf.multi_cell(page_w - 4, 5, _t(f"* {bullet_text}"), **_NL)
            pdf.ln(1)

    def _render_education():
        if not education.strip():
            return
        section_heading("Education")
        for block in re.split(r'\n\s*\n', education.strip()):
            block_lines = _block_lines(block)
            if not block_lines:
                continue
            pdf.set_font(_FONT, "B", 9)
            pdf.set_text_color(*DARK)
            pdf.multi_cell(page_w, 5, _t(block_lines[0]), **_NL)
            for extra in block_lines[1:]:
                if extra.strip():
                    pdf.set_font(_FONT, "", 9)
                    pdf.set_text_color(*LIGHT)
                    pdf.set_x(22)
                    pdf.multi_cell(page_w - 4, 5, _t(extra.strip()), **_NL)
            pdf.ln(1)

    def _render_projects():
        if not projects.strip():
            return
        section_heading("Projects")
        for block in re.split(r'\n\s*\n', projects.strip()):
            block_lines = _block_lines(block)
            if not block_lines:
                continue
            pdf.set_font(_FONT, "B", 9)
            pdf.set_text_color(*DARK)
            pdf.multi_cell(page_w, 5, _t(block_lines[0]), **_NL)
            for extra in block_lines[1:]:
                if extra.strip():
                    pdf.set_font(_FONT, "", 9)
                    pdf.set_text_color(*DARK)
                    pdf.set_x(22)
                    pdf.multi_cell(page_w - 4, 5, _t(extra.strip()), **_NL)
            pdf.ln(1)

    def _render_publications():
        if not publications.strip():
            return
        section_heading("Publications")
        for block in re.split(r'\n\s*\n', publications.strip()):
            block_lines = _block_lines(block)
            if not block_lines:
                continue
            pdf.set_font(_FONT, "I", 9)
            pdf.set_text_color(*DARK)
            pdf.multi_cell(page_w, 5, _t(block_lines[0]), **_NL)
            for extra in block_lines[1:]:
                if extra.strip():
                    pdf.set_font(_FONT, "", 9)
                    pdf.set_text_color(*LIGHT)
                    pdf.set_x(22)
                    pdf.multi_cell(page_w - 4, 5, _t(extra.strip()), **_NL)
            pdf.ln(1)

    # Section order depends on layout
    if layout == "functional":
        # Functional: skills prominently before experience
        _render_skills()
        _render_experience()
        _render_education()
    else:
        # Chronological (default): experience → education → skills
        _render_experience()
        _render_education()
        _render_skills()

    _render_projects()
    _render_publications()

    pdf.output(output_path)


@fastapi_app.post("/api/cv/generate")
async def api_cv_generate(
    request: Request,
    name: str = Form(""),
    email: str = Form(""),
    phone: str = Form(""),
    location: str = Form(""),
    link: str = Form(""),
    summary: str = Form(""),
    experience: str = Form(""),
    education: str = Form(""),
    skills: str = Form(""),
    projects: str = Form(""),
    publications: str = Form(""),
    logo: UploadFile = File(None),
    theme: str = Form("classic"),
    layout: str = Form("chronological"),
):
    """Generate a professional PDF CV using fpdf2 (pure Python, no LaTeX required)."""
    import tempfile

    try:
        from fpdf import FPDF  # noqa: F401
    except ImportError:
        return JSONResponse(
            {"error": "CV generation is not available: fpdf2 is not installed on this server."},
            status_code=503,
        )

    name = name.strip()
    email = email.strip()
    theme = (theme or "").strip().lower() or "classic"
    layout = (layout or "chronological").strip().lower()
    if layout not in ("chronological", "functional"):
        layout = "chronological"
    _VALID_THEMES = {
        "classic", "modern", "minimal", "executive",
        "creative", "tech", "elegant", "vibrant",
    }
    if theme not in _VALID_THEMES:
        theme = "classic"
    if not name or not email:
        return JSONResponse({"error": "Name and email are required."}, status_code=400)

    ip = _get_real_ip(request)

    tmpdir = tempfile.mkdtemp(prefix="cv_", dir=DOWNLOAD_FOLDER)
    try:
        # Save logo if provided
        logo_path = ""
        if logo and logo.filename:
            ext = os.path.splitext(logo.filename)[1].lower()
            if ext not in (".png", ".jpg", ".jpeg"):
                return JSONResponse({"error": "Logo must be PNG or JPG."}, status_code=400)
            logo_path = os.path.join(tmpdir, f"logo{ext}")
            content = await logo.read()
            if len(content) > _MAX_CV_LOGO_BYTES:
                return JSONResponse({"error": f"Logo file is too large (max {_MAX_CV_LOGO_BYTES // (1024*1024)} MB)."}, status_code=400)
            with open(logo_path, "wb") as f:
                f.write(content)

        output_pdf = os.path.join(tmpdir, "cv.pdf")

        # Build PDF directly using fpdf2 (no external tools required)
        _build_cv_pdf(
            output_pdf,
            name=name,
            email=email,
            phone=phone.strip(),
            location=location.strip(),
            link=link.strip(),
            summary=summary.strip(),
            experience=experience,
            education=education,
            skills=skills,
            projects=projects,
            publications=publications,
            logo_path=logo_path,
            theme=theme,
            layout=layout,
        )

        if not os.path.isfile(output_pdf):
            return JSONResponse({"error": "CV generation produced no output."}, status_code=500)

        # Track the CV generation as a download record
        record_id = str(uuid.uuid4())
        cached_geo = ip_country_cache.get(ip, {})
        with downloads_lock:
            downloads[record_id] = {
                "id": record_id,
                "url": "cv_generation",
                "title": f"CV: {name}",
                "safe_title": safe_filename(f"cv_{name}"),
                "status": "complete",
                "type": "cv_generation",
                "percent": 100,
                "created_at": time.time(),
                "end_time": time.time(),
                "filename": "cv.pdf",
                "ip": ip,
                "country": cached_geo.get("country", ""),
                "country_code": cached_geo.get("code", ""),
                "city": cached_geo.get("city", ""),
                "region": cached_geo.get("region", ""),
            }
        threading.Thread(target=save_downloads_to_disk, daemon=True).start()

        return FileResponse(
            output_pdf,
            filename="cv.pdf",
            media_type="application/pdf",
            background=None,
        )
    except Exception as exc:
        logger.error("CV generation error: %s", exc, exc_info=True)
        return JSONResponse({"error": f"CV generation failed: {exc}"}, status_code=500)
    finally:
        # Clean up temp directory after a short delay to allow FileResponse to stream
        def _delayed_rm(path, delay=30):
            time.sleep(delay)
            shutil.rmtree(path, ignore_errors=True)
        threading.Thread(target=_delayed_rm, args=(tmpdir,), daemon=True).start()


# =========================================================
# CV TXT EXPORT
# =========================================================

def _build_cv_txt(
    *,
    name: str,
    email: str,
    phone: str = "",
    location: str = "",
    link: str = "",
    summary: str = "",
    experience: str = "",
    education: str = "",
    skills: str = "",
    projects: str = "",
    publications: str = "",
    layout: str = "chronological",
) -> str:
    """Build a plain-text representation of the CV.

    ``layout`` controls section order:
    - ``chronological`` – standard order: summary, experience, education, skills, ...
    - ``functional``    – skills-first: summary, skills, experience, education, ...
    """
    lines: list[str] = []
    sep = "=" * 60

    # ── Header ──
    lines.append(name.upper())
    contact_parts = [p for p in (email, phone, location, link) if p]
    if contact_parts:
        lines.append("  |  ".join(contact_parts))
    lines.append(sep)
    lines.append("")

    def _section(title: str, body: str) -> None:
        if not body.strip():
            return
        lines.append(title.upper())
        lines.append("-" * len(title))
        lines.append(body.strip())
        lines.append("")

    if layout == "functional":
        _section("Professional Summary", summary)
        _section("Skills", skills)
        _section("Work Experience", experience)
        _section("Education", education)
        _section("Projects", projects)
        _section("Publications", publications)
    else:  # chronological (default)
        _section("Professional Summary", summary)
        _section("Work Experience", experience)
        _section("Education", education)
        _section("Skills", skills)
        _section("Projects", projects)
        _section("Publications", publications)

    return "\n".join(lines)


@fastapi_app.post("/api/cv/generate_txt")
async def api_cv_generate_txt(
    request: Request,
    name: str = Form(""),
    email: str = Form(""),
    phone: str = Form(""),
    location: str = Form(""),
    link: str = Form(""),
    summary: str = Form(""),
    experience: str = Form(""),
    education: str = Form(""),
    skills: str = Form(""),
    projects: str = Form(""),
    publications: str = Form(""),
    layout: str = Form("chronological"),
):
    """Generate a plain-text CV and return it as a downloadable .txt file."""
    name = name.strip()
    email = email.strip()
    if not name or not email:
        return JSONResponse({"error": "Name and email are required."}, status_code=400)

    layout = (layout or "chronological").strip().lower()
    if layout not in ("chronological", "functional"):
        layout = "chronological"

    cv_text = _build_cv_txt(
        name=name,
        email=email,
        phone=phone.strip(),
        location=location.strip(),
        link=link.strip(),
        summary=summary.strip(),
        experience=experience,
        education=education,
        skills=skills,
        projects=projects,
        publications=publications,
        layout=layout,
    )

    safe_name = "".join(c if c.isalnum() or c in " _-" else "_" for c in name).strip() or "cv"
    filename = f"cv_{safe_name}.txt"
    return Response(
        content=cv_text.encode("utf-8"),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


# =========================================================
# CV EXTRACTION MODULE
# =========================================================

_MAX_CV_UPLOAD_BYTES = 10 * 1024 * 1024  # 10 MB max for CV uploads

def _extract_text_from_pdf(path: str) -> str:
    """Extract plain text from a PDF using pdftotext (poppler-utils)."""
    import subprocess
    try:
        result = subprocess.run(
            ["pdftotext", "-layout", path, "-"],
            capture_output=True, text=True, timeout=30,
        )
        return result.stdout
    except Exception:
        return ""

def _extract_text_from_docx(path: str) -> str:
    """Extract plain text from a DOCX file using python-docx."""
    try:
        from docx import Document
        doc = Document(path)
        parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append("  |  ".join(cells))
        return "\n".join(parts)
    except Exception:
        return ""

def _parse_cv_text(text: str) -> dict:
    """Heuristically parse plain-text CV content into structured fields."""
    import re

    lines = [l.rstrip() for l in text.splitlines()]
    non_empty = [l for l in lines if l.strip()]

    def _find_email(t: str) -> str:
        m = re.search(r"[\w.+-]+@[\w.-]+\.[a-zA-Z]{2,}", t)
        return m.group(0) if m else ""

    def _find_phone(t: str) -> str:
        m = re.search(r"(?:\+?\d[\d\s\-().]{7,}\d)", t)
        return m.group(0).strip() if m else ""

    def _find_link(t: str) -> str:
        m = re.search(r"https?://[^\s]+|linkedin\.com/[^\s]+|github\.com/[^\s]+", t, re.I)
        return m.group(0).strip() if m else ""

    # Name: usually the first non-empty line that doesn't look like a heading keyword
    _HEADING_KEYWORDS = re.compile(
        r"^(curriculum vitae|cv|resume|profile|summary|objective|experience|education|skills|contact|references)\b",
        re.I,
    )
    name = ""
    for l in non_empty[:5]:
        if not _HEADING_KEYWORDS.match(l.strip()) and len(l.strip().split()) <= 6 and not _find_email(l):
            name = l.strip()
            break

    email    = _find_email(text)
    phone    = _find_phone(text)
    link     = _find_link(text)

    # Location: look for City/Country pattern near name/email lines
    location = ""
    loc_pat  = re.compile(r"[A-Z][a-zA-Z\s]+,\s*[A-Z][a-zA-Z\s]+")
    for l in non_empty[:15]:
        m = loc_pat.search(l)
        if m and not _find_email(l):
            location = m.group(0).strip()
            break

    # Section extraction helper
    _SEC_PATTERN = re.compile(
        r"^(summary|professional summary|objective|profile"
        r"|experience|work experience|employment|career"
        r"|education|academic|qualifications"
        r"|skills|technical skills|core competencies"
        r"|projects|personal projects"
        r"|publications|research|certificates?|certifications?)"
        r"[\s:]*$",
        re.I,
    )

    sections: dict[str, list[str]] = {}
    current_sec = None
    for line in lines:
        stripped = line.strip()
        m = _SEC_PATTERN.match(stripped)
        if m:
            key = m.group(1).lower()
            # Normalise key
            if "summary" in key or "objective" in key or "profile" in key:
                key = "summary"
            elif "experience" in key or "employment" in key or "career" in key:
                key = "experience"
            elif "education" in key or "academic" in key or "qualification" in key:
                key = "education"
            elif "skill" in key or "competenc" in key:
                key = "skills"
            elif "project" in key:
                key = "projects"
            elif "publication" in key or "research" in key or "certif" in key:
                key = "publications"
            current_sec = key
            sections.setdefault(key, [])
        elif current_sec is not None:
            sections[current_sec].append(line)

    def _sec(key: str) -> str:
        return "\n".join(sections.get(key, [])).strip()

    summary     = _sec("summary")
    experience  = _sec("experience")
    education   = _sec("education")
    publications = _sec("publications")
    projects    = _sec("projects")

    # Skills: try skills section first, otherwise look for comma-separated lines
    skills = _sec("skills")
    if not skills:
        for l in non_empty:
            if len(l.split(",")) >= 4:
                skills = l.strip()
                break

    return {
        "name": name,
        "email": email,
        "phone": phone,
        "location": location,
        "link": link,
        "summary": summary,
        "experience": experience,
        "education": education,
        "skills": skills,
        "projects": projects,
        "publications": publications,
    }


@fastapi_app.post("/api/cv/extract")
async def api_cv_extract(
    request: Request,
    file: UploadFile = File(...),
):
    """Extract CV field data from an uploaded PDF or DOCX file."""
    import tempfile

    filename = (file.filename or "").lower()
    if not (filename.endswith(".pdf") or filename.endswith(".docx") or filename.endswith(".doc")):
        return JSONResponse(
            {"error": "Only PDF and DOCX files are supported for CV extraction."},
            status_code=400,
        )

    content = await file.read()
    if len(content) > _MAX_CV_UPLOAD_BYTES:
        return JSONResponse(
            {"error": f"File is too large (max {_MAX_CV_UPLOAD_BYTES // (1024 * 1024)} MB)."},
            status_code=400,
        )

    tmpdir = tempfile.mkdtemp(prefix="cvext_")
    try:
        ext = ".pdf" if filename.endswith(".pdf") else ".docx"
        tmp_path = os.path.join(tmpdir, "upload" + ext)
        with open(tmp_path, "wb") as f:
            f.write(content)

        if ext == ".pdf":
            text = _extract_text_from_pdf(tmp_path)
        else:
            text = _extract_text_from_docx(tmp_path)

        if not text.strip():
            return JSONResponse(
                {"error": "Could not extract text from the file. The file may be image-based or encrypted."},
                status_code=422,
            )

        fields = _parse_cv_text(text)
        return JSONResponse({"fields": fields})

    except Exception as exc:
        logger.error("CV extraction error: %s", exc, exc_info=True)
        return JSONResponse({"error": f"CV extraction failed: {exc}"}, status_code=500)
    finally:
        shutil.rmtree(tmpdir, ignore_errors=True)


# =========================================================
# AI ASSISTANT MODULE
# =========================================================
# Free AI models are used to provide suggestions and enhancements
# for CV fields, text polish, and download URL analysis.
#
# Priority order:
#   1. Groq API (free tier, fast inference) — used when GROQ_API_KEY is set.
#   2. Hugging Face Inference API (free public models) — used when HF_TOKEN is
#      set or the model is public and the free anonymous tier is sufficient.
#   3. Rule-based offline suggestions — always available as a fallback; no
#      internet access or API key required.
# =========================================================

_GROQ_API_KEY  = os.environ.get("GROQ_API_KEY", "")
_HF_TOKEN      = os.environ.get("HF_TOKEN", "")
_AI_MODEL_GROQ = os.environ.get("AI_MODEL_GROQ", "llama3-8b-8192")
_AI_MODEL_HF   = os.environ.get("AI_MODEL_HF", "mistralai/Mistral-7B-Instruct-v0.3")

# ---------------------------------------------------------------------------
# Offline / rule-based CV suggestion engine (no API key required)
# ---------------------------------------------------------------------------

_CV_VERBS = [
    "Led", "Built", "Designed", "Developed", "Implemented", "Improved",
    "Managed", "Delivered", "Achieved", "Reduced", "Increased", "Launched",
    "Collaborated", "Mentored", "Optimised", "Automated", "Streamlined",
    "Architected", "Deployed", "Migrated", "Integrated", "Analysed",
    "Coordinated", "Established", "Transformed",
]

_WEAK_PHRASES = {
    r"\bresponsible for\b":       "focus on measurable achievements instead (e.g. 'Led …')",
    r"\bhelped\b":                "use a stronger action verb (e.g. 'Contributed to …', 'Collaborated on …')",
    r"\bworked on\b":             "be specific (e.g. 'Developed …', 'Implemented …')",
    r"\bwas involved in\b":       "describe your direct contribution",
    r"\bduties included?\b":      "list accomplishments, not duties",
    r"\bresponsibilities\b":      "focus on impact and results",
    r"\b(i am a|my)\b":           "avoid first-person pronouns in CVs",
    r"\b(very|really|quite)\b":   "remove filler adverbs for conciseness",
    r"\b(etc\.?|and so on)\b":    "be explicit — list every item",
}

_SUMMARY_TIPS = [
    "Start with a strong professional title and years of experience.",
    "Mention 2–3 core technical or domain competencies.",
    "Include a brief statement about the value you bring to employers.",
    "Keep the summary to 3–5 sentences (50–80 words).",
]

_EXPERIENCE_TIPS = [
    "Use the format: Company — Title — Start–End dates.",
    "Begin each bullet point with a strong past-tense action verb.",
    "Quantify results wherever possible (e.g. 'Reduced load time by 40%').",
    "Focus on impact: what changed because of your work?",
    "Separate each role with a blank line.",
]

_SKILLS_TIPS = [
    "Group skills by category (e.g. Languages, Frameworks, Tools, Cloud).",
    "List the most relevant skills first.",
    "Avoid rating skills (e.g. 'Python ★★★★☆') — just list them.",
    "Include both hard skills and relevant soft skills.",
]


def _rule_based_cv_suggestions(field: str, text: str) -> dict:
    """Return offline, rule-based suggestions for a given CV *field* and *text*.

    Returns a dict with keys:
      ``suggestions``  – list of plain-English improvement hints
      ``sample_verbs`` – list of action verbs to consider (for experience field)
    """
    import re as _re

    suggestions = []

    if not text.strip():
        tip_map = {
            "summary":     _SUMMARY_TIPS[:2],
            "experience":  _EXPERIENCE_TIPS[:2],
            "skills":      _SKILLS_TIPS[:2],
        }
        suggestions = tip_map.get(field, ["Add content to this section to make your CV stand out."])
        return {"suggestions": suggestions, "sample_verbs": []}

    text_lower = text.lower()

    # Field-specific structural checks
    if field == "summary":
        word_count = len(text.split())
        if word_count < 20:
            suggestions.append("Your summary is quite short. Aim for 50–80 words.")
        elif word_count > 120:
            suggestions.append("Your summary is too long. Keep it under 100 words.")
        for tip in _SUMMARY_TIPS:
            suggestions.append(tip)

    elif field == "experience":
        if not _re.search(r"\d{4}", text):
            suggestions.append("Include dates for each role (e.g. 2020–2024).")
        bullet_count = len(_re.findall(r"^[\*\-\u2022]", text, _re.MULTILINE))
        if bullet_count == 0:
            suggestions.append(
                "Add bullet points under each role to describe your achievements."
            )
        for tip in _EXPERIENCE_TIPS:
            suggestions.append(tip)

    elif field == "skills":
        skill_list = [s.strip() for s in text.split(",") if s.strip()]
        if len(skill_list) < 4:
            suggestions.append("Consider listing at least 6–10 skills.")
        if len(skill_list) > 30:
            suggestions.append("You have many skills listed — consider grouping them.")
        for tip in _SKILLS_TIPS:
            suggestions.append(tip)

    elif field == "education":
        if not _re.search(r"\d{4}", text):
            suggestions.append("Include graduation years for each qualification.")

    # Weak phrase detection (applies to all fields)
    for pattern, advice in _WEAK_PHRASES.items():
        if _re.search(pattern, text_lower):
            suggestions.append(f"Tip: {advice}.")

    # Deduplicate while preserving order
    seen: set = set()
    unique: list = []
    for s in suggestions:
        if s not in seen:
            seen.add(s)
            unique.append(s)

    return {
        "suggestions": unique[:8],
        "sample_verbs": _CV_VERBS[:10] if field == "experience" else [],
    }


# ---------------------------------------------------------------------------
# Groq API helper
# ---------------------------------------------------------------------------

def _call_groq(prompt: str, max_tokens: int = 400) -> str:
    """Call the Groq chat-completion API and return the assistant reply text.

    Returns an empty string on any error so callers can fall back gracefully.
    """
    import urllib.request
    import urllib.error
    import json as _json

    if not _GROQ_API_KEY:
        return ""
    payload = _json.dumps({
        "model": _AI_MODEL_GROQ,
        "messages": [{"role": "user", "content": prompt}],
        "max_tokens": max_tokens,
        "temperature": 0.5,
    }).encode()
    req = urllib.request.Request(
        "https://api.groq.com/openai/v1/chat/completions",
        data=payload,
        headers={
            "Content-Type": "application/json",
            "Authorization": f"Bearer {_GROQ_API_KEY}",
        },
        method="POST",
    )
    try:
        with urllib.request.urlopen(req, timeout=15) as resp:
            data = _json.loads(resp.read())
            return data["choices"][0]["message"]["content"].strip()
    except (urllib.error.URLError, _json.JSONDecodeError, KeyError, ValueError, OSError) as exc:
        logger.warning("Groq API call failed: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# Hugging Face Inference API helper
# ---------------------------------------------------------------------------

def _call_hf_inference(prompt: str, max_new_tokens: int = 300) -> str:
    """Call the Hugging Face Inference API (free, rate-limited) and return the
    generated text.  Returns an empty string on any error.
    """
    import urllib.request
    import urllib.error
    import json as _json

    url = f"https://api-inference.huggingface.co/models/{_AI_MODEL_HF}"
    payload = _json.dumps({
        "inputs": prompt,
        "parameters": {
            "max_new_tokens": max_new_tokens,
            "temperature": 0.5,
            "return_full_text": False,
        },
    }).encode()
    headers: dict = {"Content-Type": "application/json"}
    if _HF_TOKEN:
        headers["Authorization"] = f"Bearer {_HF_TOKEN}"
    req = urllib.request.Request(url, data=payload, headers=headers, method="POST")
    try:
        with urllib.request.urlopen(req, timeout=20) as resp:
            data = _json.loads(resp.read())
            if isinstance(data, list) and data:
                return (data[0].get("generated_text") or "").strip()
            return ""
    except (urllib.error.URLError, _json.JSONDecodeError, KeyError, ValueError, OSError) as exc:
        logger.warning("HF Inference API call failed: %s", exc)
        return ""


# ---------------------------------------------------------------------------
# /api/ai/cv_suggest  — main AI suggestion endpoint
# ---------------------------------------------------------------------------

class _AiCvRequest(BaseModel):
    field: str
    text: str = ""
    name: str = ""
    job_title: str = ""


@fastapi_app.post("/api/ai/cv_suggest")
async def api_ai_cv_suggest(body: _AiCvRequest):
    """Return AI-powered improvement suggestions for a CV field.

    Supported fields: ``summary``, ``experience``, ``education``, ``skills``,
    ``projects``, ``publications``.

    The endpoint tries (in order):
    1. Groq API  (if ``GROQ_API_KEY`` is set in the environment)
    2. Hugging Face Inference API  (if ``HF_TOKEN`` is set, or anonymously)
    3. Rule-based offline suggestions  (always available)

    Response JSON:
    ```json
    {
      "suggestions": ["hint 1", "hint 2", ...],
      "sample_verbs": ["Led", "Built", ...],
      "enhanced_text": "Optional AI-rewritten version of the text",
      "source": "groq" | "huggingface" | "offline"
    }
    ```
    """
    _VALID_FIELDS = {"summary", "experience", "education", "skills", "projects", "publications"}
    field = (body.field or "").strip().lower()
    if field not in _VALID_FIELDS:
        return JSONResponse({"error": f"Invalid field '{field}'."}, status_code=400)

    text = (body.text or "").strip()
    name = (body.name or "").strip()
    job_title = (body.job_title or "").strip()

    # Try Groq first (fastest, highest quality)
    if _GROQ_API_KEY:
        context = f"Name: {name}\nJob title: {job_title}\n" if (name or job_title) else ""
        prompt = (
            f"You are a professional CV/résumé writing coach. "
            f"Review the following '{field}' section of a CV and provide:\n"
            f"1. Up to 5 concise, actionable improvement suggestions (one per line, "
            f"prefixed with '- ').\n"
            f"2. A polished, rewritten version of the text (labelled 'REWRITE:').\n\n"
            f"{context}"
            f"TEXT:\n{text or '(empty)'}\n\n"
            f"Keep suggestions brief and specific."
        )
        ai_text = _call_groq(prompt, max_tokens=400)
        if ai_text:
            lines = ai_text.splitlines()
            suggestions = [
                l.lstrip("-– ").strip()
                for l in lines
                if l.strip().startswith(("-", "–", "*", "•")) and "REWRITE:" not in l
            ]
            rewrite = ""
            in_rewrite = False
            for l in lines:
                if l.strip().upper().startswith("REWRITE:"):
                    in_rewrite = True
                    rewrite = l[l.upper().find("REWRITE:") + 8:].strip()
                elif in_rewrite:
                    rewrite += "\n" + l
            rule = _rule_based_cv_suggestions(field, text)
            return JSONResponse({
                "suggestions": suggestions or rule["suggestions"],
                "sample_verbs": rule["sample_verbs"],
                "enhanced_text": rewrite.strip(),
                "source": "groq",
            })

    # Try Hugging Face Inference API
    if text:
        hf_prompt = (
            f"Improve the following '{field}' section of a professional CV. "
            f"Rewrite it to be more impactful and results-oriented:\n\n{text}"
        )
        hf_text = _call_hf_inference(hf_prompt, max_new_tokens=250)
        if hf_text:
            rule = _rule_based_cv_suggestions(field, text)
            return JSONResponse({
                "suggestions": rule["suggestions"],
                "sample_verbs": rule["sample_verbs"],
                "enhanced_text": hf_text,
                "source": "huggingface",
            })

    # Fallback: offline rule-based suggestions
    result = _rule_based_cv_suggestions(field, text)
    return JSONResponse({
        "suggestions": result["suggestions"],
        "sample_verbs": result["sample_verbs"],
        "enhanced_text": "",
        "source": "offline",
    })


# ---------------------------------------------------------------------------
# /api/ai/enhance_text  — general text polishing endpoint
# ---------------------------------------------------------------------------

class _AiEnhanceRequest(BaseModel):
    text: str
    context: str = "professional CV"


@fastapi_app.post("/api/ai/enhance_text")
async def api_ai_enhance_text(body: _AiEnhanceRequest):
    """Polish a block of text for clarity and professionalism.

    Uses Groq → HF Inference → offline echo (no-op) as the fallback chain.

    Response JSON:
    ```json
    {
      "original": "...",
      "enhanced": "...",
      "source": "groq" | "huggingface" | "offline"
    }
    ```
    """
    text = (body.text or "").strip()
    context = (body.context or "professional CV").strip()
    if not text:
        return JSONResponse({"error": "text is required."}, status_code=400)
    if len(text) > 5000:
        return JSONResponse({"error": "text is too long (max 5000 characters)."}, status_code=400)

    if _GROQ_API_KEY:
        prompt = (
            f"Rewrite the following text for use in a {context}. "
            f"Make it clearer, more concise, and more impactful. "
            f"Return only the rewritten text, no preamble:\n\n{text}"
        )
        enhanced = _call_groq(prompt, max_tokens=500)
        if enhanced:
            return JSONResponse({"original": text, "enhanced": enhanced, "source": "groq"})

    hf_prompt = (
        f"Rewrite the following for a {context}. Return only the improved text:\n\n{text}"
    )
    enhanced = _call_hf_inference(hf_prompt, max_new_tokens=300)
    if enhanced:
        return JSONResponse({"original": text, "enhanced": enhanced, "source": "huggingface"})

    # Offline: return original text unchanged
    return JSONResponse({"original": text, "enhanced": text, "source": "offline"})


# ---------------------------------------------------------------------------
# /api/ai/status  — reports which AI back-end is active
# ---------------------------------------------------------------------------

@fastapi_app.get("/api/ai/status")
async def api_ai_status():
    """Return information about the active AI backend.

    Response JSON:
    ```json
    {
      "groq_available": true,
      "hf_available": true,
      "offline_available": true,
      "active_backend": "groq",
      "model": "llama3-8b-8192"
    }
    ```
    """
    groq_ok = bool(_GROQ_API_KEY)
    hf_ok   = True  # always available (anonymous free tier)
    if groq_ok:
        active = "groq"
        model  = _AI_MODEL_GROQ
    else:
        active = "huggingface"
        model  = _AI_MODEL_HF
    return JSONResponse({
        "groq_available":    groq_ok,
        "hf_available":      hf_ok,
        "offline_available": True,
        "active_backend":    active,
        "model":             model,
    })


# =========================================================
# DOCUMENT CONVERSION MODULE
# =========================================================

_MAX_DOC_UPLOAD_BYTES = 50 * 1024 * 1024  # 50 MB max for doc conversion

_DOC_CONVERSIONS = {
    # source_ext → {target_format → output_ext}
    "pdf":  {"word": "docx", "excel": "xlsx", "jpeg": "jpg", "png": "png", "text": "txt"},
    "docx": {"pdf": "pdf", "text": "txt"},
    "doc":  {"pdf": "pdf", "text": "txt"},
    "xlsx": {"pdf": "pdf"},
    "xls":  {"pdf": "pdf"},
    "jpg":  {"pdf": "pdf"},
    "jpeg": {"pdf": "pdf"},
    "png":  {"pdf": "pdf"},
    "rtf":  {"pdf": "pdf", "text": "txt"},
    "txt":  {"pdf": "pdf"},
}

_LIBREOFFICE_FORMATS = {"docx", "doc", "xlsx", "xls", "pptx", "ppt", "odt", "ods", "rtf", "txt"}


def _convert_pdf_to_word(src: str, dst: str) -> None:
    """Convert PDF to DOCX using pdf2docx."""
    from pdf2docx import Converter
    cv = Converter(src)
    cv.convert(dst, start=0, end=None)
    cv.close()


def _convert_pdf_to_excel(src: str, dst: str) -> None:
    """Extract tables from PDF pages and write to an Excel workbook."""
    import tabula
    import openpyxl
    tables = tabula.read_pdf(src, pages="all", multiple_tables=True, silent=True)
    wb = openpyxl.Workbook()
    wb.remove(wb.active)
    if not tables:
        ws = wb.create_sheet("Sheet1")
        ws.append(["No tables found in the PDF."])
    else:
        for i, df in enumerate(tables):
            ws = wb.create_sheet(f"Table{i + 1}")
            # Header row
            ws.append(list(df.columns))
            for _, row in df.iterrows():
                ws.append([str(v) if v is not None else "" for v in row])
    wb.save(dst)


def _convert_pdf_to_image(src: str, dst_dir: str, fmt: str) -> list[str]:
    """Convert each PDF page to an image using pdftoppm (poppler)."""
    import subprocess
    fmt_flag = "jpeg" if fmt in ("jpg", "jpeg") else "png"
    prefix = os.path.join(dst_dir, "page")
    subprocess.run(
        ["pdftoppm", f"-{fmt_flag}", "-r", "150", src, prefix],
        check=True, capture_output=True, timeout=120,
    )
    # pdftoppm outputs page-1.jpg / page-01.jpg etc. depending on version
    out_files = sorted(
        f for f in os.listdir(dst_dir)
        if f.startswith("page") and (f.endswith(".jpg") or f.endswith(".jpeg") or f.endswith(".png"))
    )
    return [os.path.join(dst_dir, f) for f in out_files]


def _convert_to_pdf_libreoffice(src: str, dst_dir: str) -> str:
    """Convert Office documents to PDF using LibreOffice headless."""
    import subprocess
    subprocess.run(
        ["libreoffice", "--headless", "--convert-to", "pdf", "--outdir", dst_dir, src],
        check=True, capture_output=True, timeout=120,
    )
    base = os.path.splitext(os.path.basename(src))[0]
    return os.path.join(dst_dir, base + ".pdf")


def _convert_image_to_pdf(src: str, dst: str) -> None:
    """Convert a JPEG or PNG image to PDF using img2pdf."""
    import img2pdf
    with open(src, "rb") as img_f, open(dst, "wb") as pdf_f:
        pdf_f.write(img2pdf.convert(img_f))


@fastapi_app.post("/api/doc/convert")
async def api_doc_convert(
    request: Request,
    file: UploadFile = File(...),
    target: str = Form(...),
):
    """Convert a document/image between supported formats and return the result file."""
    import tempfile, zipfile

    filename = (file.filename or "upload").strip()
    src_ext  = os.path.splitext(filename)[1].lower().lstrip(".")
    target   = target.lower().strip()

    allowed_src = set(_DOC_CONVERSIONS.keys())
    if src_ext not in allowed_src:
        return JSONResponse(
            {"error": f"Unsupported source format '.{src_ext}'. Supported: {', '.join(sorted(allowed_src))}."},
            status_code=400,
        )
    allowed_targets = _DOC_CONVERSIONS.get(src_ext, {})
    if target not in allowed_targets:
        return JSONResponse(
            {"error": f"Cannot convert .{src_ext} to '{target}'. Supported targets: {', '.join(sorted(allowed_targets))}."},
            status_code=400,
        )
    out_ext = allowed_targets[target]

    content = await file.read()
    if len(content) > _MAX_DOC_UPLOAD_BYTES:
        return JSONResponse(
            {"error": f"File is too large (max {_MAX_DOC_UPLOAD_BYTES // (1024 * 1024)} MB)."},
            status_code=400,
        )

    tmpdir = tempfile.mkdtemp(prefix="docconv_")
    try:
        src_path = os.path.join(tmpdir, f"input.{src_ext}")
        with open(src_path, "wb") as f:
            f.write(content)

        base_name = os.path.splitext(filename)[0] or "converted"

        # ── PDF → Word ──────────────────────────────────────────────────────────
        if src_ext == "pdf" and target == "word":
            out_path = os.path.join(tmpdir, f"{base_name}.docx")
            await asyncio.get_event_loop().run_in_executor(
                None, _convert_pdf_to_word, src_path, out_path,
            )
            return FileResponse(out_path, filename=f"{base_name}.docx",
                                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        # ── PDF → Excel ─────────────────────────────────────────────────────────
        elif src_ext == "pdf" and target == "excel":
            out_path = os.path.join(tmpdir, f"{base_name}.xlsx")
            await asyncio.get_event_loop().run_in_executor(
                None, _convert_pdf_to_excel, src_path, out_path,
            )
            return FileResponse(out_path, filename=f"{base_name}.xlsx",
                                media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")

        # ── PDF → JPEG / PNG ────────────────────────────────────────────────────
        elif src_ext == "pdf" and target in ("jpeg", "png"):
            imgs_dir = os.path.join(tmpdir, "imgs")
            os.makedirs(imgs_dir, exist_ok=True)
            img_files = await asyncio.get_event_loop().run_in_executor(
                None, _convert_pdf_to_image, src_path, imgs_dir, target,
            )
            if not img_files:
                return JSONResponse({"error": "No pages found in PDF."}, status_code=422)
            if len(img_files) == 1:
                # Single page — return the image directly
                ext_out = "jpg" if target == "jpeg" else "png"
                mime    = "image/jpeg" if target == "jpeg" else "image/png"
                return FileResponse(img_files[0], filename=f"{base_name}.{ext_out}", media_type=mime)
            else:
                # Multiple pages — zip them up
                zip_path = os.path.join(tmpdir, f"{base_name}_pages.zip")
                ext_out  = "jpg" if target == "jpeg" else "png"
                with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
                    for i, p in enumerate(img_files, 1):
                        zf.write(p, f"page_{i:03d}.{ext_out}")
                return FileResponse(zip_path, filename=f"{base_name}_pages.zip",
                                    media_type="application/zip")

        # ── Office → PDF ────────────────────────────────────────────────────────
        elif target == "pdf" and src_ext in _LIBREOFFICE_FORMATS:
            out_path = await asyncio.get_event_loop().run_in_executor(
                None, _convert_to_pdf_libreoffice, src_path, tmpdir,
            )
            return FileResponse(out_path, filename=f"{base_name}.pdf",
                                media_type="application/pdf")

        # ── Image → PDF ─────────────────────────────────────────────────────────
        elif target == "pdf" and src_ext in ("jpg", "jpeg", "png"):
            out_path = os.path.join(tmpdir, f"{base_name}.pdf")
            await asyncio.get_event_loop().run_in_executor(
                None, _convert_image_to_pdf, src_path, out_path,
            )
            return FileResponse(out_path, filename=f"{base_name}.pdf",
                                media_type="application/pdf")

        # ── Any supported source → Plain Text ───────────────────────────────────
        elif target == "text":
            if src_ext == "pdf":
                raw = await asyncio.get_event_loop().run_in_executor(
                    None, _extract_text_from_pdf, src_path,
                )
            elif src_ext in ("docx", "doc", "odt"):
                raw = await asyncio.get_event_loop().run_in_executor(
                    None, _extract_text_from_docx_rich, src_path,
                )
            elif src_ext == "rtf":
                raw = await asyncio.get_event_loop().run_in_executor(
                    None, _extract_text_from_rtf, src_path,
                )
            else:  # txt
                raw = _extract_text_from_txt(src_path)
            text_out = _normalize_text_bullets(raw)
            out_path = os.path.join(tmpdir, f"{base_name}.txt")
            with open(out_path, "w", encoding="utf-8") as fh:
                fh.write(text_out)
            return FileResponse(out_path, filename=f"{base_name}.txt",
                                media_type="text/plain; charset=utf-8")

        else:
            return JSONResponse({"error": "Conversion not implemented."}, status_code=501)

    except Exception as exc:
        logger.error("Doc conversion error: %s", exc, exc_info=True)
        return JSONResponse({"error": f"Conversion failed: {exc}"}, status_code=500)
    finally:
        def _rm(p):
            import time as _time
            _time.sleep(_TEMP_DIR_CLEANUP_DELAY_SECS)
            shutil.rmtree(p, ignore_errors=True)
        threading.Thread(target=_rm, args=(tmpdir,), daemon=True).start()


# =========================================================
# DOCUMENT TEXT EXTRACTION MODULE
# =========================================================

_TEXT_EXTRACT_ACCEPT = {"pdf", "docx", "doc", "txt", "rtf", "odt"}
_TEXT_EXTRACT_MAX_BYTES = 20 * 1024 * 1024  # 20 MB

# Regex that matches a leading bullet character (possibly preceded by spaces/tabs)
_BULLET_CHARS_RE = re.compile(
    r"^([ \t]*)([•◦▸▹►▻✓✗✦✧❖○●■□▪▫\u2013\u2014]|-{1,2}|\*{1,2})\s+",
    re.MULTILINE,
)


def _normalize_text_bullets(text: str) -> str:
    """Normalize various bullet/list characters to '•', preserving indentation."""
    def _repl(m: re.Match) -> str:
        return m.group(1) + "• "
    return _BULLET_CHARS_RE.sub(_repl, text)


def _extract_text_from_docx_rich(path: str) -> str:
    """Extract text from DOCX preserving line breaks and normalising list bullets."""
    try:
        from docx import Document as _DocxDocument
        doc = _DocxDocument(path)
        lines: list[str] = []
        for para in doc.paragraphs:
            raw = para.text
            # Detect numbered/bulleted list paragraphs via paragraph style name
            style_name = (para.style.name or "").lower()
            is_list = "list" in style_name or (
                para._p.pPr is not None
                and para._p.pPr.find(
                    "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}numPr"
                ) is not None
            )
            if is_list and raw.strip():
                lines.append("• " + raw)
            else:
                lines.append(raw)
        # Include table content
        for tbl in doc.tables:
            for row in tbl.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    lines.append("  |  ".join(cells))
        return "\n".join(lines)
    except Exception:
        return ""


def _extract_text_from_txt(path: str) -> str:
    """Read a plain-text file, trying UTF-8 then latin-1 fallback."""
    for enc in ("utf-8", "latin-1"):
        try:
            with open(path, "r", encoding=enc) as fh:
                return fh.read()
        except UnicodeDecodeError:
            continue
    return ""


def _extract_text_from_rtf(path: str) -> str:
    """Extract plain text from an RTF file.

    Tries pypandoc first (most accurate); falls back to a simple regex-based
    RTF-tag stripper when pypandoc / pandoc is not available.
    """
    # Try pypandoc (requires pandoc binary)
    try:
        import pypandoc
        return pypandoc.convert_file(path, "plain", format="rtf")
    except Exception:
        pass

    # Simple regex fallback: strip RTF control words / groups
    try:
        import re as _re
        with open(path, "rb") as fh:
            raw = fh.read()
        # Decode best-effort
        text = raw.decode("utf-8", errors="replace")
        # Remove RTF header and embedded binary blobs
        text = _re.sub(r"\\bin\d+\s?[^\\{]*", "", text)
        # Strip control words (e.g. \rtf1, \b, \par, \pard …)
        text = _re.sub(r"\\[a-zA-Z]+[-]?\d*\s?", "", text)
        # Strip braces
        text = text.replace("{", "").replace("}", "")
        # Collapse whitespace
        text = _re.sub(r"\r\n|\r", "\n", text)
        text = _re.sub(r"\n{3,}", "\n\n", text).strip()
        return text
    except Exception:
        return ""


@fastapi_app.post("/api/doc/to_text")
async def api_doc_to_text(
    request: Request,
    file: UploadFile = File(...),
):
    """Extract clean plain text from a PDF, DOCX, DOC, ODT, or TXT file.

    The returned JSON contains:
      - ``text``: the extracted text (Unicode-safe, emojis preserved, bullets normalised)
      - ``filename``: original filename
      - ``truncated``: true if content was trimmed to 200 000 characters
    """
    import tempfile

    session = request.session
    if not session.get("user_id"):
        return JSONResponse({"error": "Not authenticated."}, status_code=401)

    filename = (file.filename or "upload").strip()
    src_ext = os.path.splitext(filename)[1].lower().lstrip(".")

    if src_ext not in _TEXT_EXTRACT_ACCEPT:
        return JSONResponse(
            {
                "error": (
                    f"Unsupported file type '.{src_ext}'. "
                    f"Supported: {', '.join(sorted(_TEXT_EXTRACT_ACCEPT))}."
                )
            },
            status_code=400,
        )

    content = await file.read()
    if len(content) > _TEXT_EXTRACT_MAX_BYTES:
        return JSONResponse(
            {"error": f"File too large (max {_TEXT_EXTRACT_MAX_BYTES // (1024 * 1024)} MB)."},
            status_code=400,
        )

    tmpdir = tempfile.mkdtemp(prefix="totext_")
    try:
        src_path = os.path.join(tmpdir, f"input.{src_ext}")
        with open(src_path, "wb") as fh:
            fh.write(content)

        if src_ext == "pdf":
            raw = _extract_text_from_pdf(src_path)
        elif src_ext in ("docx", "doc", "odt"):
            raw = _extract_text_from_docx_rich(src_path)
        elif src_ext == "rtf":
            raw = _extract_text_from_rtf(src_path)
        else:  # txt
            raw = _extract_text_from_txt(src_path)

        # Normalise bullets; emojis and line breaks are kept as-is
        text = _normalize_text_bullets(raw)

        _MAX_CHARS = 200_000
        truncated = len(text) > _MAX_CHARS
        if truncated:
            text = text[:_MAX_CHARS]

        return JSONResponse({"text": text, "filename": filename, "truncated": truncated})

    except Exception as exc:
        logger.error("Text extraction error: %s", exc, exc_info=True)
        return JSONResponse({"error": f"Extraction failed: {exc}"}, status_code=500)
    finally:
        def _rm(p):
            import time as _t
            _t.sleep(_TEMP_DIR_CLEANUP_DELAY_SECS)
            shutil.rmtree(p, ignore_errors=True)
        threading.Thread(target=_rm, args=(tmpdir,), daemon=True).start()


# =========================================================
# ATS CV SCANNING MODULE
# =========================================================

# Common stop-words to skip when extracting keywords
_ATS_STOP_WORDS = frozenset({
    "a", "an", "the", "and", "or", "but", "in", "on", "at", "to", "for",
    "of", "with", "is", "are", "was", "were", "be", "been", "being",
    "have", "has", "had", "do", "does", "did", "will", "would", "shall",
    "should", "may", "might", "can", "could", "must", "that", "this",
    "these", "those", "it", "its", "we", "you", "your", "they", "their",
    "our", "us", "he", "she", "him", "her", "from", "by", "as", "into",
    "about", "above", "after", "before", "between", "through", "during",
    "while", "if", "then", "than", "so", "not", "no", "also", "more",
    "other", "such", "each", "any", "all", "both", "few", "most", "some",
    "experience", "work", "job", "role", "position", "candidate",
    "required", "preferred", "minimum", "years", "year", "including",
    "ability", "strong", "knowledge", "good", "excellent", "team", "own",
})

# Common tech / domain skill tokens that should be captured as multi-word phrases
_ATS_TECH_PHRASES = [
    "machine learning", "deep learning", "natural language processing",
    "computer vision", "data science", "data analysis", "big data",
    "cloud computing", "software development", "software engineering",
    "project management", "product management", "agile methodology",
    "continuous integration", "continuous deployment", "ci/cd",
    "object oriented", "test driven development", "rest api", "restful api",
    "version control", "source control", "microsoft office",
]


def _extract_ats_keywords(text: str) -> list[str]:
    """Extract meaningful keywords/phrases from a job description."""
    text_lower = text.lower()

    # First capture multi-word tech phrases
    found: set[str] = set()
    for phrase in _ATS_TECH_PHRASES:
        if phrase in text_lower:
            found.add(phrase)

    # Then extract individual significant tokens
    tokens = re.findall(r"[a-zA-Z][a-zA-Z0-9+#./\-]{2,}", text)
    for tok in tokens:
        tok_lower = tok.lower().rstrip(".")
        if tok_lower not in _ATS_STOP_WORDS and len(tok_lower) >= 3:
            found.add(tok_lower)

    return sorted(found)


def _score_ats(cv_text: str, keywords: list[str]) -> dict:
    """Score CV text against a list of keywords. Returns matched/missing/score."""
    cv_lower = cv_text.lower()
    matched  = [kw for kw in keywords if kw in cv_lower]
    missing  = [kw for kw in keywords if kw not in cv_lower]
    total    = len(keywords) if keywords else 1
    score    = round(len(matched) / total * 100)

    # Build improvement tips based on the score
    tips: list[str] = []
    if score < 50:
        tips.append("Your CV matches fewer than half the job keywords — consider adding a tailored skills section.")
    if score < 75:
        tips.append("Add the missing keywords naturally into your experience or skills sections.")
    if missing:
        tips.append(f"Top missing keywords: {', '.join(missing[:10])}.")
    if score >= 80:
        tips.append("Strong keyword alignment — ensure each keyword is backed by concrete examples.")

    return {
        "score":   score,
        "matched": matched,
        "missing": missing,
        "tips":    tips,
    }


class _AtsRequest(BaseModel):
    cv_text:         str = ""
    job_description: str = ""


@fastapi_app.post("/api/cv/ats_scan")
async def api_cv_ats_scan(
    request: Request,
    file: UploadFile = File(None),
    job_description: str = Form(""),
    cv_text: str = Form(""),
):
    """Score a CV against a job description for ATS compatibility.

    Accepts either:
    - A multipart form with an optional ``file`` (PDF/DOCX) + ``job_description`` text, or
    - A JSON body: ``{ cv_text, job_description }``

    Returns ``{score, matched, missing, tips, keywords_total}``.
    """
    # Support JSON body too
    ct = request.headers.get("content-type", "")
    if "application/json" in ct:
        try:
            body = await request.json()
            cv_text         = body.get("cv_text", "")
            job_description = body.get("job_description", "")
        except Exception:
            return JSONResponse({"error": "Invalid JSON body."}, status_code=400)

    # If a file was uploaded, extract text from it
    if file is not None:
        import tempfile
        filename = (file.filename or "").lower()
        if not (filename.endswith(".pdf") or filename.endswith(".docx") or filename.endswith(".doc")):
            return JSONResponse(
                {"error": "Only PDF and DOCX files are supported for ATS scanning."},
                status_code=400,
            )
        content = await file.read()
        if len(content) > _MAX_CV_UPLOAD_BYTES:
            return JSONResponse({"error": "File too large (max 10 MB)."}, status_code=400)
        tmpdir = tempfile.mkdtemp(prefix="ats_")
        try:
            ext = ".pdf" if filename.endswith(".pdf") else ".docx"
            tmp_path = os.path.join(tmpdir, "upload" + ext)
            with open(tmp_path, "wb") as fh:
                fh.write(content)
            cv_text = _extract_text_from_pdf(tmp_path) if ext == ".pdf" else _extract_text_from_docx(tmp_path)
        finally:
            shutil.rmtree(tmpdir, ignore_errors=True)

    if not cv_text.strip():
        return JSONResponse({"error": "CV text is required."}, status_code=400)
    if not job_description.strip():
        return JSONResponse({"error": "Job description is required."}, status_code=400)

    keywords = _extract_ats_keywords(job_description)
    result   = _score_ats(cv_text, keywords)
    result["keywords_total"] = len(keywords)
    return JSONResponse(result)


# =========================================================
# USER AUTHENTICATION MODULE
# =========================================================

_APP_USER_PROXIMITY_KM = float(os.environ.get("APP_USER_PROXIMITY_KM", "50"))
# Base fare per km for airport pickup rides (env-configurable, $1/km as per platform standard)
_FARE_PER_KM = float(os.environ.get("FARE_PER_KM", "1.0"))

# In-memory driver location store: user_id → {lat, lng, ts, name, user_id, empty}
_driver_locations: dict[str, dict] = {}
_driver_loc_lock = Lock()
_DRIVER_LOC_TTL_SECS = 300  # 5 minutes

# In-memory map of socket sid → user_id for authenticated users
_sid_to_user: dict[str, str] = {}
_user_to_sid: dict[str, str] = {}
_socket_user_lock = Lock()


def _haversine_km(lat1: float, lon1: float, lat2: float, lon2: float) -> float:
    """Return the great-circle distance in kilometres between two lat/lng points."""
    R = 6371.0
    phi1, phi2 = math.radians(lat1), math.radians(lat2)
    dphi  = math.radians(lat2 - lat1)
    dlam  = math.radians(lon2 - lon1)
    a = math.sin(dphi / 2) ** 2 + math.cos(phi1) * math.cos(phi2) * math.sin(dlam / 2) ** 2
    return R * 2 * math.atan2(math.sqrt(a), math.sqrt(1 - a))


def _get_app_user(user_id: str) -> dict | None:
    """Fetch a user row from the database by user_id. Returns None if not found."""
    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("SELECT user_id,name,email,role,location_lat,location_lng,location_name,avatar_url,bio,created_at,public_key,can_post_properties FROM app_users WHERE user_id=%s", (user_id,))
                row = cur.fetchone()
            else:
                cur = conn.execute("SELECT user_id,name,email,role,location_lat,location_lng,location_name,avatar_url,bio,created_at,public_key,can_post_properties FROM app_users WHERE user_id=?", (user_id,))
                row = cur.fetchone()
            if row is None:
                return None
            keys = ["user_id", "name", "email", "role", "location_lat", "location_lng", "location_name", "avatar_url", "bio", "created_at", "public_key", "can_post_properties"]
            return dict(zip(keys, row))
        finally:
            conn.close()


def _get_app_user_by_email(email: str) -> dict | None:
    """Fetch a user row including password_hash by email."""
    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("SELECT user_id,name,email,password_hash,role,location_lat,location_lng,location_name,created_at FROM app_users WHERE email=%s", (email.lower(),))
                row = cur.fetchone()
            else:
                cur = conn.execute("SELECT user_id,name,email,password_hash,role,location_lat,location_lng,location_name,created_at FROM app_users WHERE email=?", (email.lower(),))
                row = cur.fetchone()
            if row is None:
                return None
            keys = ["user_id", "name", "email", "password_hash", "role", "location_lat", "location_lng", "location_name", "created_at"]
            return dict(zip(keys, row))
        finally:
            conn.close()


class _UserRegisterRequest(BaseModel):
    name:     str
    email:    str
    password: str
    role:     str = "passenger"  # "passenger" | "driver"


class _UserLoginRequest(BaseModel):
    email:       str
    password:    str
    remember_me: bool = False


class _MagicLinkRequest(BaseModel):
    email: str


class _DriverApplyRequest(BaseModel):
    vehicle_make:      str
    vehicle_model:     str
    vehicle_year:      int
    vehicle_color:     str
    license_plate:     str
    subscription_type: str = "monthly"  # "monthly" | "yearly"


class _DriverApproveRequest(BaseModel):
    approved: bool


class _AgentApplyRequest(BaseModel):
    full_name:      str
    email:         str
    phone:         str = ""
    agency_name:   str = ""
    license_number: str


class _AgentApproveRequest(BaseModel):
    approved: bool


class _StorePublicKeyRequest(BaseModel):
    public_key: str


class _UserLocationUpdate(BaseModel):
    lat:           float
    lng:           float
    location_name: str = ""


class _UserProfileDetailsUpdate(BaseModel):
    name: str = ""
    bio:  str = ""


@fastapi_app.post("/api/auth/register")
async def api_user_register(body: _UserRegisterRequest):
    """Register a new platform user (passenger or driver)."""
    name     = body.name.strip()
    email    = body.email.strip().lower()
    password = body.password
    role     = body.role if body.role in ("passenger", "driver") else "passenger"

    if not name or not email or not password:
        return JSONResponse({"error": "Name, email and password are required."}, status_code=400)
    if len(password) < 6:
        return JSONResponse({"error": "Password must be at least 6 characters."}, status_code=400)
    if not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
        return JSONResponse({"error": "Invalid email address."}, status_code=400)

    user_id      = str(uuid.uuid4())
    pw_hash      = generate_password_hash(password)
    created_at   = datetime.now(timezone.utc).isoformat()

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                try:
                    cur.execute(
                        "INSERT INTO app_users (user_id,name,email,password_hash,role,created_at) VALUES (%s,%s,%s,%s,%s,%s)",
                        (user_id, name, email, pw_hash, role, created_at),
                    )
                except Exception:
                    conn.rollback()
                    return JSONResponse({"error": "Email already registered."}, status_code=409)
            else:
                try:
                    conn.execute(
                        "INSERT INTO app_users (user_id,name,email,password_hash,role,created_at) VALUES (?,?,?,?,?,?)",
                        (user_id, name, email, pw_hash, role, created_at),
                    )
                except Exception:
                    return JSONResponse({"error": "Email already registered."}, status_code=409)
            conn.commit()
        finally:
            conn.close()

    return JSONResponse({
        "ok":        True,
        "user_id":   user_id,
        "name":      name,
        "email":     email,
        "role":      role,
        "created_at": created_at,
    }, status_code=201)


@fastapi_app.post("/api/auth/login")
async def api_user_login(request: Request, body: _UserLoginRequest):
    """Login as a platform user. Sets a session cookie."""
    email    = body.email.strip().lower()
    password = body.password

    user = _get_app_user_by_email(email)
    if user is None or not check_password_hash(user["password_hash"], password):
        return JSONResponse({"error": "Invalid email or password."}, status_code=401)

    request.session["app_user_id"] = user["user_id"]
    if body.remember_me:
        # Extend session lifetime to 30 days for "Remember Me"
        request.session["remember_me"] = True
    return JSONResponse({
        "ok":      True,
        "user_id": user["user_id"],
        "name":    user["name"],
        "email":   user["email"],
        "role":    user["role"],
        "created_at": user.get("created_at", ""),
    })


@fastapi_app.post("/api/auth/logout")
async def api_user_logout(request: Request):
    """Logout the current platform user."""
    request.session.pop("app_user_id", None)
    return JSONResponse({"ok": True})


@fastapi_app.get("/api/auth/me")
async def api_user_me(request: Request):
    """Return the currently logged-in user's profile."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Not logged in."}, status_code=401)
    user = _get_app_user(user_id)
    if user is None:
        request.session.pop("app_user_id", None)
        return JSONResponse({"error": "User not found."}, status_code=404)
    return JSONResponse(user)


@fastapi_app.get("/api/user/dashboard")
async def api_user_dashboard(request: Request):
    """Return aggregated dashboard data for the logged-in user.

    Combines profile details, ride statistics, and recent ride history so
    the frontend UserDashboard page can render a personalised summary in a
    single request.
    """
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    user = _get_app_user(user_id)
    if user is None:
        request.session.pop("app_user_id", None)
        return JSONResponse({"error": "User not found."}, status_code=404)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["ride_id", "origin", "destination", "departure",
                    "seats", "status", "created_at"]
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "SELECT ride_id,origin,destination,departure,seats,status,created_at "
                    "FROM rides WHERE user_id=%s ORDER BY created_at DESC LIMIT 5",
                    (user_id,),
                )
                rows = cur.fetchall()
                cur.execute(
                    "SELECT COUNT(*) FROM rides WHERE user_id=%s", (user_id,)
                )
                total_rides = (cur.fetchone() or [0])[0]
                cur.execute(
                    "SELECT COUNT(*) FROM rides WHERE user_id=%s AND status='open'",
                    (user_id,),
                )
                open_rides = (cur.fetchone() or [0])[0]
            else:
                cur = conn.execute(
                    "SELECT ride_id,origin,destination,departure,seats,status,created_at "
                    "FROM rides WHERE user_id=? ORDER BY created_at DESC LIMIT 5",
                    (user_id,),
                )
                rows = cur.fetchall()
                cur = conn.execute(
                    "SELECT COUNT(*) FROM rides WHERE user_id=?", (user_id,)
                )
                total_rides = (cur.fetchone() or [0])[0]
                cur = conn.execute(
                    "SELECT COUNT(*) FROM rides WHERE user_id=? AND status='open'",
                    (user_id,),
                )
                open_rides = (cur.fetchone() or [0])[0]
        finally:
            conn.close()

    recent_rides = [dict(zip(cols, r)) for r in rows]
    return JSONResponse({
        "user": user,
        "stats": {
            "total_rides": total_rides,
            "open_rides": open_rides,
        },
        "recent_rides": recent_rides,
    })


@fastapi_app.put("/api/auth/profile")
async def api_user_update_profile(request: Request, body: _UserLocationUpdate):
    """Update the logged-in user's location."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Not logged in."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "UPDATE app_users SET location_lat=%s, location_lng=%s, location_name=%s WHERE user_id=%s",
                    (body.lat, body.lng, body.location_name.strip(), user_id),
                )
            else:
                conn.execute(
                    "UPDATE app_users SET location_lat=?, location_lng=?, location_name=? WHERE user_id=?",
                    (body.lat, body.lng, body.location_name.strip(), user_id),
                )
            conn.commit()
        finally:
            conn.close()

    return JSONResponse({"ok": True})


@fastapi_app.put("/api/auth/profile/details")
async def api_user_update_profile_details(request: Request, body: _UserProfileDetailsUpdate):
    """Update the logged-in user's name and bio."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Not logged in."}, status_code=401)

    name = body.name.strip()
    bio  = body.bio.strip()[:500]  # cap bio at 500 chars

    updates = {}
    if name:
        updates["name"] = name
    if bio is not None:
        updates["bio"] = bio

    if not updates:
        return JSONResponse({"ok": True})

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                for field, value in updates.items():
                    cur.execute(f"UPDATE app_users SET {field}=%s WHERE user_id=%s", (value, user_id))
            else:
                for field, value in updates.items():
                    conn.execute(f"UPDATE app_users SET {field}=? WHERE user_id=?", (value, user_id))
            conn.commit()
        finally:
            conn.close()

    user = _get_app_user(user_id)
    return JSONResponse({"ok": True, "user": user})


@fastapi_app.post("/api/auth/profile/avatar")
async def api_user_upload_avatar(request: Request, file: UploadFile = File(...)):
    """Upload a profile avatar image for the logged-in user."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Not logged in."}, status_code=401)

    # Validate content type
    allowed_types = {"image/jpeg", "image/png", "image/gif", "image/webp"}
    ct = (file.content_type or "").lower()
    if ct not in allowed_types:
        return JSONResponse({"error": "Only JPEG, PNG, GIF and WebP images are allowed."}, status_code=400)

    ext_map = {"image/jpeg": "jpg", "image/png": "png", "image/gif": "gif", "image/webp": "webp"}
    ext = ext_map.get(ct, "jpg")

    data = await file.read()
    if len(data) > 5 * 1024 * 1024:  # 5 MB limit
        return JSONResponse({"error": "Avatar image must be under 5 MB."}, status_code=400)

    filename = f"{user_id}.{ext}"
    avatar_path = os.path.join(AVATARS_DIR, filename)
    with open(avatar_path, "wb") as fh:
        fh.write(data)

    # When S3 is configured, upload the avatar to the media bucket and serve
    # via the /api/avatars/ endpoint (which redirects to a presigned URL).
    if _S3_ENABLED:
        s3_key = f"avatars/{filename}"
        _s3_upload_bytes(data, s3_key, ct)
        avatar_url = f"/api/avatars/{filename}"
    else:
        avatar_url = f"/static/avatars/{filename}"

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("UPDATE app_users SET avatar_url=%s WHERE user_id=%s", (avatar_url, user_id))
            else:
                conn.execute("UPDATE app_users SET avatar_url=? WHERE user_id=?", (avatar_url, user_id))
            conn.commit()
        finally:
            conn.close()

    return JSONResponse({"ok": True, "avatar_url": avatar_url})


@fastapi_app.get("/api/avatars/{filename}")
async def api_serve_avatar(filename: str):
    """Serve a user avatar from the media bucket (S3) or the local filesystem.

    When S3 is configured the client is redirected to a presigned URL so the
    image is streamed directly from the bucket.  Falls back to the locally
    cached copy when S3 is unavailable or not configured.
    """
    # Sanitise filename to prevent path-traversal
    safe_name = os.path.basename(filename)
    if not safe_name or safe_name != filename:
        raise HTTPException(status_code=400, detail="Invalid filename.")

    if _S3_ENABLED:
        presigned = _s3_presigned_url(f"avatars/{safe_name}")
        if presigned:
            return RedirectResponse(presigned)

    local_path = os.path.join(AVATARS_DIR, safe_name)
    if os.path.exists(local_path):
        mime = mimetypes.guess_type(safe_name)[0] or "image/jpeg"
        return FileResponse(local_path, media_type=mime)

    raise HTTPException(status_code=404, detail="Avatar not found.")


# ── Notifications ──────────────────────────────────────────────────────────────

def _create_notification(user_id: str, notif_type: str, title: str, body: str) -> str:
    """Insert a notification row for a user and return the notif_id."""
    notif_id  = str(uuid.uuid4())
    created   = datetime.now(timezone.utc).isoformat()
    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "INSERT INTO notifications (notif_id,user_id,type,title,body,created_at) VALUES (%s,%s,%s,%s,%s,%s)",
                    (notif_id, user_id, notif_type, title, body, created),
                )
            else:
                conn.execute(
                    "INSERT INTO notifications (notif_id,user_id,type,title,body,created_at) VALUES (?,?,?,?,?,?)",
                    (notif_id, user_id, notif_type, title, body, created),
                )
            conn.commit()
        finally:
            conn.close()
    # Persist notification record to bucket
    _bucket_write_json("notifications", "notification", notif_id, {
        "notif_id": notif_id,
        "user_id": user_id,
        "type": notif_type,
        "title": title,
        "body": body,
        "read_status": False,
        "created_at": created,
    })
    return notif_id


@fastapi_app.get("/api/notifications")
async def api_get_notifications(request: Request):
    """Return notifications for the logged-in user (most recent first, max 50)."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    cols = ["notif_id", "type", "title", "body", "read", "created_at"]
    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "SELECT notif_id,type,title,body,read,created_at FROM notifications WHERE user_id=%s ORDER BY created_at DESC LIMIT 50",
                    (user_id,),
                )
                rows = cur.fetchall()
            else:
                cur = conn.execute(
                    "SELECT notif_id,type,title,body,read,created_at FROM notifications WHERE user_id=? ORDER BY created_at DESC LIMIT 50",
                    (user_id,),
                )
                rows = cur.fetchall()
        finally:
            conn.close()

    notifs = [dict(zip(cols, r)) for r in rows]
    unread = sum(1 for n in notifs if not n["read"])
    return JSONResponse({"notifications": notifs, "unread": unread})


@fastapi_app.post("/api/notifications/{notif_id}/read")
async def api_mark_notification_read(request: Request, notif_id: str):
    """Mark a single notification as read."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "UPDATE notifications SET read=1 WHERE notif_id=%s AND user_id=%s",
                    (notif_id, user_id),
                )
            else:
                conn.execute(
                    "UPDATE notifications SET read=1 WHERE notif_id=? AND user_id=?",
                    (notif_id, user_id),
                )
            conn.commit()
        finally:
            conn.close()

    return JSONResponse({"ok": True})


@fastapi_app.post("/api/notifications/read_all")
async def api_mark_all_notifications_read(request: Request):
    """Mark all notifications for the logged-in user as read."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("UPDATE notifications SET read=1 WHERE user_id=%s", (user_id,))
            else:
                conn.execute("UPDATE notifications SET read=1 WHERE user_id=?", (user_id,))
            conn.commit()
        finally:
            conn.close()

    return JSONResponse({"ok": True})


# ── Platform stats ─────────────────────────────────────────────────────────────

@fastapi_app.get("/api/platform_stats")
async def api_platform_stats():
    """Return aggregated platform-wide statistics.

    Counts are computed live from the database and the result is also
    persisted to the bucket under /stats/ for historical reference.
    """
    with _db_lock:
        conn = _get_db()
        try:
            ph = "%s" if USE_POSTGRES else "?"
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("SELECT COUNT(*) FROM rides")
                total_rides = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM rides WHERE status='open'")
                open_rides = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM properties")
                total_properties = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM properties WHERE status='active'")
                active_properties = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM driver_applications WHERE status='approved'")
                registered_drivers = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM driver_applications WHERE status='pending'")
                pending_driver_apps = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM app_users")
                total_users = cur.fetchone()[0]
                cur.execute("SELECT COUNT(*) FROM notifications")
                total_notifications = cur.fetchone()[0]
            else:
                total_rides = conn.execute("SELECT COUNT(*) FROM rides").fetchone()[0]
                open_rides = conn.execute("SELECT COUNT(*) FROM rides WHERE status='open'").fetchone()[0]
                total_properties = conn.execute("SELECT COUNT(*) FROM properties").fetchone()[0]
                active_properties = conn.execute("SELECT COUNT(*) FROM properties WHERE status='active'").fetchone()[0]
                registered_drivers = conn.execute("SELECT COUNT(*) FROM driver_applications WHERE status='approved'").fetchone()[0]
                pending_driver_apps = conn.execute("SELECT COUNT(*) FROM driver_applications WHERE status='pending'").fetchone()[0]
                total_users = conn.execute("SELECT COUNT(*) FROM app_users").fetchone()[0]
                total_notifications = conn.execute("SELECT COUNT(*) FROM notifications").fetchone()[0]
        finally:
            conn.close()

    stats = {
        "total_rides": total_rides,
        "open_rides": open_rides,
        "total_properties": total_properties,
        "active_properties": active_properties,
        "registered_drivers": registered_drivers,
        "pending_driver_applications": pending_driver_apps,
        "total_users": total_users,
        "total_notifications": total_notifications,
        "generated_at": datetime.now(timezone.utc).isoformat(),
    }

    # Persist daily stats snapshot to bucket under /stats/
    today = datetime.now(timezone.utc).strftime("%Y%m%d")
    _bucket_write_json("stats", "stats", today, stats)

    return JSONResponse(stats)


# ── Ride chat messages (persistent) ────────────────────────────────────────────

_CHAT_MEDIA_PREFIX = "chat_media/"


def _resolve_chat_media(messages: list[dict]) -> list[dict]:
    """Replace S3 object keys in *messages* ``media_data`` fields with presigned URLs.

    When S3 is not configured or the key cannot be resolved the field is left
    unchanged so callers always receive a usable list.
    """
    if not _S3_ENABLED:
        return messages
    for msg in messages:
        md = msg.get("media_data")
        if md and isinstance(md, str) and md.startswith(_CHAT_MEDIA_PREFIX):
            url = _s3_presigned_url(md)
            if url:
                msg["media_data"] = url
    return messages


@fastapi_app.get("/api/rides/{ride_id}/chat")
async def api_get_ride_chat(request: Request, ride_id: str):
    """Return persisted chat messages for a ride (most recent last, max 200)."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["msg_id", "ride_id", "sender_name", "sender_role", "text",
                    "media_type", "media_data", "lat", "lng", "ts", "created_at"]
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "SELECT msg_id,ride_id,sender_name,sender_role,text,media_type,media_data,lat,lng,ts,created_at"
                    " FROM ride_chat_messages WHERE ride_id=%s ORDER BY ts ASC LIMIT 200",
                    (ride_id,),
                )
                rows = cur.fetchall()
            else:
                cur = conn.execute(
                    "SELECT msg_id,ride_id,sender_name,sender_role,text,media_type,media_data,lat,lng,ts,created_at"
                    " FROM ride_chat_messages WHERE ride_id=? ORDER BY ts ASC LIMIT 200",
                    (ride_id,),
                )
                rows = cur.fetchall()
        finally:
            conn.close()

    messages = [dict(zip(cols, row)) for row in rows]
    _resolve_chat_media(messages)
    return JSONResponse({"messages": messages})


@fastapi_app.get("/api/rides/chat/inbox")
async def api_ride_chat_inbox(request: Request):
    """Return latest chat message per ride that involves the current user's rides.

    Returns conversations grouped by ride_id, ordered by most-recent message.
    """
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    user = _get_app_user(user_id)
    if not user:
        return JSONResponse({"error": "User not found."}, status_code=404)

    sender_name = user["name"]

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                # Rides posted by the user
                cur.execute(
                    "SELECT ride_id, origin, destination FROM rides WHERE user_id=%s",
                    (user_id,),
                )
                poster_rides = {r[0]: {"origin": r[1], "destination": r[2]} for r in cur.fetchall()}
                # Latest message per ride that the user is involved in
                cur.execute(
                    """
                    SELECT DISTINCT ON (ride_id) msg_id, ride_id, sender_name, sender_role,
                           text, media_type, ts
                    FROM ride_chat_messages
                    WHERE ride_id = ANY(
                        SELECT ride_id FROM rides WHERE user_id = %s
                    ) OR sender_name = %s
                    ORDER BY ride_id, ts DESC
                    """,
                    (user_id, sender_name),
                )
                rows = cur.fetchall()
            else:
                # Rides posted by the user
                cur = conn.execute(
                    "SELECT ride_id, origin, destination FROM rides WHERE user_id=?",
                    (user_id,),
                )
                poster_rides = {r[0]: {"origin": r[1], "destination": r[2]} for r in cur.fetchall()}
                cur = conn.execute(
                    """
                    SELECT msg_id, ride_id, sender_name, sender_role, text, media_type, ts
                    FROM ride_chat_messages
                    WHERE ride_id IN (SELECT ride_id FROM rides WHERE user_id=?)
                       OR sender_name=?
                    GROUP BY ride_id
                    HAVING ts = MAX(ts)
                    ORDER BY ts DESC
                    LIMIT 50
                    """,
                    (user_id, sender_name),
                )
                rows = cur.fetchall()
        finally:
            conn.close()

    cols = ["msg_id", "ride_id", "sender_name", "sender_role", "text", "media_type", "ts"]
    conversations = []
    for row in rows:
        d = dict(zip(cols, row))
        d["ride_info"] = poster_rides.get(d["ride_id"], {})
        d["is_mine"] = (d["sender_name"] == sender_name)
        conversations.append(d)

    return JSONResponse({"conversations": conversations})


# ── Magic link (passwordless) ──────────────────────────────────────────────────

# In-memory store: token → {email, expires_at}
_magic_link_tokens: dict = {}
_magic_link_lock = threading.Lock()
_MAGIC_LINK_TTL_SECONDS = 900  # 15 minutes


@fastapi_app.post("/api/auth/magic_link")
async def api_magic_link_request(body: _MagicLinkRequest):
    """Generate a one-time magic-link token for passwordless login.

    In a production deployment this token would be emailed to the user.
    The endpoint returns the token in the response for demo / testing purposes.
    """
    email = body.email.strip().lower()
    if not email or not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
        return JSONResponse({"error": "Invalid email address."}, status_code=400)

    user = _get_app_user_by_email(email)
    if user is None:
        # Don't reveal whether the address is registered
        return JSONResponse({"ok": True, "message": "If that address is registered, a login link has been sent."})

    token = secrets.token_urlsafe(32)
    expires_at = time.time() + _MAGIC_LINK_TTL_SECONDS
    with _magic_link_lock:
        _magic_link_tokens[token] = {"email": email, "expires_at": expires_at}

    # In a real app: send email here.  For now, return token so the UI can demonstrate the flow.
    return JSONResponse({
        "ok":     True,
        "token":  token,
        "message": "Magic link generated. Check your email (demo: token returned in response).",
    })


@fastapi_app.post("/api/auth/magic_link/verify")
async def api_magic_link_verify(request: Request):
    """Verify a magic-link token and log the user in."""
    data = await request.json()
    token = (data.get("token") or "").strip()
    if not token:
        return JSONResponse({"error": "Token required."}, status_code=400)

    with _magic_link_lock:
        entry = _magic_link_tokens.get(token)
        if entry is None or time.time() > entry["expires_at"]:
            _magic_link_tokens.pop(token, None)
            return JSONResponse({"error": "Invalid or expired token."}, status_code=401)
        del _magic_link_tokens[token]  # single-use

    user = _get_app_user_by_email(entry["email"])
    if user is None:
        return JSONResponse({"error": "User not found."}, status_code=404)

    request.session["app_user_id"] = user["user_id"]
    return JSONResponse({
        "ok":        True,
        "user_id":   user["user_id"],
        "name":      user["name"],
        "email":     user["email"],
        "role":      user["role"],
        "created_at": user.get("created_at", ""),
    })


# ── Driver registration & approval ────────────────────────────────────────────

@fastapi_app.post("/api/auth/driver_apply")
async def api_driver_apply(request: Request, body: _DriverApplyRequest):
    """Submit a driver-role application."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    user = _get_app_user(user_id)
    if user is None:
        return JSONResponse({"error": "User not found."}, status_code=404)

    # Basic validation
    if not body.vehicle_make.strip() or not body.vehicle_model.strip():
        return JSONResponse({"error": "Vehicle make and model are required."}, status_code=400)
    if body.vehicle_year < 1900 or body.vehicle_year > datetime.now().year + 1:
        return JSONResponse({"error": "Invalid vehicle year."}, status_code=400)
    if not body.license_plate.strip():
        return JSONResponse({"error": "License plate is required."}, status_code=400)
    subscription_type = body.subscription_type if body.subscription_type in ("monthly", "yearly") else "monthly"

    app_id   = str(uuid.uuid4())
    created  = datetime.now(timezone.utc).isoformat()

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    """INSERT INTO driver_applications
                       (app_id, user_id, vehicle_make, vehicle_model, vehicle_year, vehicle_color, license_plate, subscription_type, status, created_at)
                       VALUES (%s,%s,%s,%s,%s,%s,%s,%s,'pending',%s)
                       ON CONFLICT (user_id) DO UPDATE SET
                         vehicle_make=EXCLUDED.vehicle_make, vehicle_model=EXCLUDED.vehicle_model,
                         vehicle_year=EXCLUDED.vehicle_year, vehicle_color=EXCLUDED.vehicle_color,
                         license_plate=EXCLUDED.license_plate, subscription_type=EXCLUDED.subscription_type,
                         status='pending', created_at=EXCLUDED.created_at""",
                    (app_id, user_id, body.vehicle_make.strip(), body.vehicle_model.strip(),
                     body.vehicle_year, body.vehicle_color.strip(), body.license_plate.strip().upper(),
                     subscription_type, created),
                )
            else:
                conn.execute(
                    """INSERT OR REPLACE INTO driver_applications
                       (app_id, user_id, vehicle_make, vehicle_model, vehicle_year, vehicle_color, license_plate, subscription_type, status, created_at)
                       VALUES (?,?,?,?,?,?,?,?,'pending',?)""",
                    (app_id, user_id, body.vehicle_make.strip(), body.vehicle_model.strip(),
                     body.vehicle_year, body.vehicle_color.strip(), body.license_plate.strip().upper(),
                     subscription_type, created),
                )
            conn.commit()
        finally:
            conn.close()

    # Persist driver application to bucket under /driver_reg/pending/
    _bucket_write_json("driver_reg/pending", "driver_reg", app_id, {
        "app_id": app_id,
        "user_id": user_id,
        "vehicle_make": body.vehicle_make.strip(),
        "vehicle_model": body.vehicle_model.strip(),
        "vehicle_year": body.vehicle_year,
        "vehicle_color": body.vehicle_color.strip(),
        "license_plate": body.license_plate.strip().upper(),
        "subscription_type": subscription_type,
        "status": "pending",
        "created_at": created,
    })

    return JSONResponse({"ok": True, "app_id": app_id}, status_code=201)


@fastapi_app.get("/api/auth/driver_application")
async def api_driver_application_status(request: Request):
    """Return the current user's driver application status."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["app_id", "user_id", "vehicle_make", "vehicle_model", "vehicle_year",
                    "vehicle_color", "license_plate", "subscription_type", "status", "created_at"]
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "SELECT app_id,user_id,vehicle_make,vehicle_model,vehicle_year,vehicle_color,license_plate,subscription_type,status,created_at FROM driver_applications WHERE user_id=%s",
                    (user_id,),
                )
                row = cur.fetchone()
            else:
                cur = conn.execute(
                    "SELECT app_id,user_id,vehicle_make,vehicle_model,vehicle_year,vehicle_color,license_plate,subscription_type,status,created_at FROM driver_applications WHERE user_id=?",
                    (user_id,),
                )
                row = cur.fetchone()
        finally:
            conn.close()

    if row is None:
        return JSONResponse({"application": None})
    return JSONResponse({"application": dict(zip(cols, row))})


@fastapi_app.get("/api/admin/driver_applications")
async def api_admin_driver_applications(request: Request):
    """Return all pending driver applications (admin only)."""
    if not request.session.get("admin_user"):
        return JSONResponse({"error": "Admin login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["app_id", "user_id", "vehicle_make", "vehicle_model", "vehicle_year",
                    "vehicle_color", "license_plate", "status", "created_at", "name", "email"]
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    """SELECT da.app_id, da.user_id, da.vehicle_make, da.vehicle_model,
                              da.vehicle_year, da.vehicle_color, da.license_plate, da.status,
                              da.created_at, au.name, au.email
                       FROM driver_applications da
                       JOIN app_users au ON da.user_id = au.user_id
                       ORDER BY da.created_at DESC"""
                )
                rows = cur.fetchall()
            else:
                cur = conn.execute(
                    """SELECT da.app_id, da.user_id, da.vehicle_make, da.vehicle_model,
                              da.vehicle_year, da.vehicle_color, da.license_plate, da.status,
                              da.created_at, au.name, au.email
                       FROM driver_applications da
                       JOIN app_users au ON da.user_id = au.user_id
                       ORDER BY da.created_at DESC"""
                )
                rows = cur.fetchall()
        finally:
            conn.close()

    return JSONResponse({"applications": [dict(zip(cols, r)) for r in rows]})


@fastapi_app.post("/api/admin/driver_applications/{app_id}/approve")
async def api_admin_driver_approve(request: Request, app_id: str, body: _DriverApproveRequest):
    """Approve or reject a driver application (admin only)."""
    if not request.session.get("admin_user"):
        return JSONResponse({"error": "Admin login required."}, status_code=401)

    new_status = "approved" if body.approved else "rejected"

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("SELECT user_id FROM driver_applications WHERE app_id=%s", (app_id,))
                row = cur.fetchone()
            else:
                cur = conn.execute("SELECT user_id FROM driver_applications WHERE app_id=?", (app_id,))
                row = cur.fetchone()

            if row is None:
                return JSONResponse({"error": "Application not found."}, status_code=404)

            target_user_id = row[0]

            if USE_POSTGRES:
                cur.execute("UPDATE driver_applications SET status=%s WHERE app_id=%s", (new_status, app_id))
                if body.approved:
                    cur.execute("UPDATE app_users SET role='driver' WHERE user_id=%s", (target_user_id,))
            else:
                conn.execute("UPDATE driver_applications SET status=? WHERE app_id=?", (new_status, app_id))
                if body.approved:
                    conn.execute("UPDATE app_users SET role='driver' WHERE user_id=?", (target_user_id,))
            conn.commit()
        finally:
            conn.close()

    # Send in-app notification to the affected user
    reviewed_at = datetime.now(timezone.utc).isoformat()
    if body.approved:
        _create_notification(
            target_user_id,
            "driver_approved",
            "🎉 Driver Application Approved",
            "Congratulations! Your driver application has been approved. You can now post rides and use Driver Alerts.",
        )
        # Also emit real-time socket event if the user is connected
        with _socket_user_lock:
            sid = _user_to_sid.get(target_user_id)
        if sid:
            asyncio.ensure_future(sio.emit("notification", {
                "type":  "driver_approved",
                "title": "🎉 Driver Application Approved",
                "body":  "Your driver application has been approved!",
            }, room=sid))
        # Persist approved record to bucket under /driver_reg/verified/
        _bucket_write_json("driver_reg/verified", "driver_reg", app_id, {
            "app_id": app_id,
            "user_id": target_user_id,
            "status": "approved",
            "reviewed_at": reviewed_at,
        })
    else:
        _create_notification(
            target_user_id,
            "driver_rejected",
            "❌ Driver Application Rejected",
            "Unfortunately, your driver application was not approved this time. You may re-apply with updated details.",
        )
        # Persist rejected record to bucket under /driver_reg/pending/ (status update)
        _bucket_write_json("driver_reg/pending", "driver_reg", app_id, {
            "app_id": app_id,
            "user_id": target_user_id,
            "status": new_status,
            "reviewed_at": reviewed_at,
        })

    return JSONResponse({"ok": True, "status": new_status})


# ── Public key (E2E encryption) ────────────────────────────────────────────────

@fastapi_app.put("/api/auth/public_key")
async def api_store_public_key(request: Request, body: _StorePublicKeyRequest):
    """Store the authenticated user's public key for E2E encryption."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    pk = body.public_key.strip()
    if not pk:
        return JSONResponse({"error": "public_key is required."}, status_code=400)
    with _db_lock:
        conn = _get_db()
        try:
            _execute(
                conn,
                "UPDATE app_users SET public_key=? WHERE user_id=?"
                if not USE_POSTGRES else
                "UPDATE app_users SET public_key=%s WHERE user_id=%s",
                (pk, user_id),
            )
            conn.commit()
        finally:
            conn.close()
    return JSONResponse({"ok": True})


@fastapi_app.get("/api/users/{user_id}/public_key")
async def api_get_user_public_key(request: Request, user_id: str):
    """Return the public key for a given user (requires auth)."""
    caller = request.session.get("app_user_id")
    if not caller:
        return JSONResponse({"error": "Login required."}, status_code=401)
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT public_key FROM app_users WHERE user_id=?"
                if not USE_POSTGRES else
                "SELECT public_key FROM app_users WHERE user_id=%s",
                (user_id,),
            )
            row = cur.fetchone()
        finally:
            conn.close()
    if row is None:
        return JSONResponse({"error": "User not found."}, status_code=404)
    pk = row["public_key"] if USE_POSTGRES else row[0]
    return JSONResponse({"user_id": user_id, "public_key": pk})


# ── Agent registration ─────────────────────────────────────────────────────────

@fastapi_app.post("/api/agent_applications")
async def api_agent_apply(request: Request, body: _AgentApplyRequest):
    """Submit an agent registration application."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    full_name      = body.full_name.strip()
    email          = body.email.strip().lower()
    license_number = body.license_number.strip()

    if not full_name:
        return JSONResponse({"error": "Full name is required."}, status_code=400)
    if not email or not re.match(r"^[^@\s]+@[^@\s]+\.[^@\s]+$", email):
        return JSONResponse({"error": "A valid email address is required."}, status_code=400)
    if not license_number:
        return JSONResponse({"error": "License or identification number is required."}, status_code=400)

    app_id  = str(uuid.uuid4())
    created = datetime.now(timezone.utc).isoformat()

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    """INSERT INTO agent_applications
                       (app_id, user_id, full_name, email, phone, agency_name, license_number, status, created_at)
                       VALUES (%s,%s,%s,%s,%s,%s,%s,'pending',%s)
                       ON CONFLICT (user_id) DO UPDATE SET
                         full_name=EXCLUDED.full_name, email=EXCLUDED.email,
                         phone=EXCLUDED.phone, agency_name=EXCLUDED.agency_name,
                         license_number=EXCLUDED.license_number,
                         status='pending', created_at=EXCLUDED.created_at""",
                    (app_id, user_id, full_name, email,
                     body.phone.strip(), body.agency_name.strip(), license_number, created),
                )
            else:
                conn.execute(
                    """INSERT OR REPLACE INTO agent_applications
                       (app_id, user_id, full_name, email, phone, agency_name, license_number, status, created_at)
                       VALUES (?,?,?,?,?,?,?,'pending',?)""",
                    (app_id, user_id, full_name, email,
                     body.phone.strip(), body.agency_name.strip(), license_number, created),
                )
            conn.commit()
        finally:
            conn.close()

    _bucket_write_json("agent_reg/pending", "agent_reg", app_id, {
        "app_id":          app_id,
        "user_id":         user_id,
        "full_name":       full_name,
        "email":           email,
        "phone":           body.phone.strip(),
        "agency_name":     body.agency_name.strip(),
        "license_number":  license_number,
        "status":          "pending",
        "created_at":      created,
    })

    return JSONResponse({"ok": True, "app_id": app_id}, status_code=201)


@fastapi_app.get("/api/agent_applications/status")
async def api_agent_application_status(request: Request):
    """Return the current user's agent application status."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["app_id", "user_id", "full_name", "email", "phone",
                    "agency_name", "license_number", "status", "created_at"]
            cur = _execute(
                conn,
                "SELECT app_id,user_id,full_name,email,phone,agency_name,license_number,status,created_at FROM agent_applications WHERE user_id=?"
                if not USE_POSTGRES else
                "SELECT app_id,user_id,full_name,email,phone,agency_name,license_number,status,created_at FROM agent_applications WHERE user_id=%s",
                (user_id,),
            )
            row = cur.fetchone()
        finally:
            conn.close()

    if row is None:
        return JSONResponse({"application": None})
    return JSONResponse({"application": dict(zip(cols, row))})


@fastapi_app.get("/api/admin/agent_applications")
async def api_admin_agent_applications(request: Request):
    """Return all agent applications (admin only)."""
    if not request.session.get("admin_user"):
        return JSONResponse({"error": "Admin login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["app_id", "user_id", "full_name", "email", "phone",
                    "agency_name", "license_number", "status", "created_at"]
            cur = _execute(
                conn,
                "SELECT app_id,user_id,full_name,email,phone,agency_name,license_number,status,created_at FROM agent_applications ORDER BY created_at DESC",
            )
            rows = cur.fetchall()
        finally:
            conn.close()

    return JSONResponse({"applications": [dict(zip(cols, r)) for r in rows]})


@fastapi_app.post("/api/admin/agent_applications/{app_id}/approve")
async def api_admin_agent_approve(request: Request, app_id: str, body: _AgentApproveRequest):
    """Approve or reject an agent application (admin only)."""
    if not request.session.get("admin_user"):
        return JSONResponse({"error": "Admin login required."}, status_code=401)

    new_status  = "approved" if body.approved else "rejected"
    reviewed_at = datetime.now(timezone.utc).isoformat()

    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT user_id,full_name FROM agent_applications WHERE app_id=?"
                if not USE_POSTGRES else
                "SELECT user_id,full_name FROM agent_applications WHERE app_id=%s",
                (app_id,),
            )
            row = cur.fetchone()
            if row is None:
                return JSONResponse({"error": "Application not found."}, status_code=404)
            target_user_id = row["user_id"] if USE_POSTGRES else row[0]
            applicant_name = row["full_name"] if USE_POSTGRES else row[1]

            _execute(
                conn,
                "UPDATE agent_applications SET status=? WHERE app_id=?"
                if not USE_POSTGRES else
                "UPDATE agent_applications SET status=%s WHERE app_id=%s",
                (new_status, app_id),
            )

            if body.approved:
                _execute(
                    conn,
                    "UPDATE app_users SET can_post_properties=1 WHERE user_id=?"
                    if not USE_POSTGRES else
                    "UPDATE app_users SET can_post_properties=1 WHERE user_id=%s",
                    (target_user_id,),
                )
            else:
                _execute(
                    conn,
                    "UPDATE app_users SET can_post_properties=0 WHERE user_id=?"
                    if not USE_POSTGRES else
                    "UPDATE app_users SET can_post_properties=0 WHERE user_id=%s",
                    (target_user_id,),
                )
            conn.commit()
        finally:
            conn.close()

    # Send notification to applicant
    if body.approved:
        _create_notification(
            target_user_id,
            "agent_approved",
            "Agent Application Approved",
            "Congratulations! Your agent registration has been approved. You can now post properties.",
        )
        _bucket_write_json("agent_reg/verified", "agent_reg", app_id, {
            "app_id": app_id, "user_id": target_user_id,
            "applicant_name": applicant_name,
            "status": "approved", "reviewed_at": reviewed_at,
        })
    else:
        _create_notification(
            target_user_id,
            "agent_rejected",
            "Agent Application Rejected",
            "Your agent registration application was not approved. Please contact support for more information.",
        )
        _bucket_write_json("agent_reg/pending", "agent_reg", app_id, {
            "app_id": app_id, "user_id": target_user_id,
            "status": "rejected", "reviewed_at": reviewed_at,
        })

    return JSONResponse({"ok": True, "status": new_status})


# ── Ride history ───────────────────────────────────────────────────────────────

@fastapi_app.get("/api/rides/history")
async def api_ride_history(request: Request):
    """Return all rides associated with the logged-in user."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["ride_id", "user_id", "driver_name", "origin", "destination",
                    "departure", "seats", "notes", "status", "created_at"]
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "SELECT ride_id,user_id,driver_name,origin,destination,departure,seats,notes,status,created_at FROM rides WHERE user_id=%s ORDER BY created_at DESC",
                    (user_id,),
                )
                rows = cur.fetchall()
            else:
                cur = conn.execute(
                    "SELECT ride_id,user_id,driver_name,origin,destination,departure,seats,notes,status,created_at FROM rides WHERE user_id=? ORDER BY created_at DESC",
                    (user_id,),
                )
                rows = cur.fetchall()
        finally:
            conn.close()

    rides = [dict(zip(cols, r)) for r in rows]
    return JSONResponse({"rides": rides})

# =========================================================
# RIDE SHARING MODULE
# =========================================================

class _RidePostRequest(BaseModel):
    origin:      str
    destination: str
    departure:   str
    seats:       int = 1
    notes:       str = ""
    origin_lat:  float | None = None
    origin_lng:  float | None = None
    dest_lat:    float | None = None
    dest_lng:    float | None = None
    fare:        float | None = None
    ride_type:   str = "airport"


class _RideJoinRequest(BaseModel):
    ride_id: str


@fastapi_app.post("/api/rides/post")
async def api_ride_post(request: Request, body: _RidePostRequest):
    """Post a new airport pickup ride offer. Only verified drivers may post rides."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required to post a ride."}, status_code=401)

    user = _get_app_user(user_id)
    if user is None:
        return JSONResponse({"error": "User not found."}, status_code=404)

    # Only verified (approved) drivers may post airport pickup rides
    if user.get("role") != "driver":
        return JSONResponse(
            {"error": "Only verified drivers can post airport pickup rides. "
                      "Register as a driver from your profile to gain access."},
            status_code=403,
        )

    origin      = body.origin.strip()
    destination = body.destination.strip()
    departure   = body.departure.strip()

    if not origin or not destination or not departure:
        return JSONResponse({"error": "Origin, destination and departure are required."}, status_code=400)
    if body.seats < 1 or body.seats > 20:
        return JSONResponse({"error": "Seats must be between 1 and 20."}, status_code=400)

    # Auto-calculate fare if coordinates provided and fare not given
    fare = body.fare
    if fare is None and body.origin_lat is not None and body.dest_lat is not None:
        dist_km = _haversine_km(body.origin_lat, body.origin_lng, body.dest_lat, body.dest_lng)
        fare = round(dist_km * _FARE_PER_KM, 2)

    ride_type = body.ride_type if body.ride_type in ("airport", "standard") else "airport"

    ride_id    = str(uuid.uuid4())
    created_at = datetime.now(timezone.utc).isoformat()

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    """INSERT INTO rides (ride_id,user_id,driver_name,origin,destination,origin_lat,origin_lng,dest_lat,dest_lng,fare,departure,seats,notes,status,created_at,ride_type)
                       VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,'open',%s,%s)""",
                    (ride_id, user_id, user["name"], origin, destination,
                     body.origin_lat, body.origin_lng, body.dest_lat, body.dest_lng, fare,
                     departure, body.seats, body.notes.strip(), created_at, ride_type),
                )
            else:
                conn.execute(
                    """INSERT INTO rides (ride_id,user_id,driver_name,origin,destination,origin_lat,origin_lng,dest_lat,dest_lng,fare,departure,seats,notes,status,created_at,ride_type)
                       VALUES (?,?,?,?,?,?,?,?,?,?,?,?,?,'open',?,?)""",
                    (ride_id, user_id, user["name"], origin, destination,
                     body.origin_lat, body.origin_lng, body.dest_lat, body.dest_lng, fare,
                     departure, body.seats, body.notes.strip(), created_at, ride_type),
                )
            conn.commit()
        finally:
            conn.close()

    ride_data = {
        "ride_id":     ride_id,
        "driver_name": user["name"],
        "origin":      origin,
        "destination": destination,
        "fare":        fare,
        "departure":   departure,
        "seats":       body.seats,
        "notes":       body.notes.strip(),
        "created_at":  created_at,
        "ride_type":   ride_type,
    }

    # Notify all connected users via Socket.IO
    asyncio.ensure_future(sio.emit("new_ride", ride_data))

    # Persist ride record to bucket under /rides/
    _bucket_write_json("rides", "ride", ride_id, {
        **ride_data,
        "user_id": user_id,
        "status": "open",
        "origin_lat": body.origin_lat,
        "origin_lng": body.origin_lng,
        "dest_lat": body.dest_lat,
        "dest_lng": body.dest_lng,
    })

    return JSONResponse({"ok": True, "ride_id": ride_id}, status_code=201)


@fastapi_app.get("/api/rides/list")
async def api_rides_list(status: str | None = None):
    """Return rides ordered by departure time.

    Query param ``status`` can be 'open', 'taken', 'cancelled', or omitted to
    return open and taken rides (everything visible to passengers).
    """
    _VALID_RIDE_STATUSES = ("open", "taken", "cancelled")
    # Default: show open and taken rides so passengers can see status tags.
    if status and status in _VALID_RIDE_STATUSES:
        status_filter = [status]
    else:
        status_filter = ["open", "taken"]

    # Build parameterised placeholders from a fixed count — no user data
    # enters the SQL string itself, so this is safe from injection.
    n = len(status_filter)
    pg_placeholders  = ",".join(["%s"] * n)
    sql_placeholders = ",".join(["?"]  * n)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["ride_id", "user_id", "driver_name", "origin", "destination",
                    "origin_lat", "origin_lng", "dest_lat", "dest_lng", "fare",
                    "departure", "seats", "notes", "status", "created_at", "ride_type"]
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "SELECT ride_id,user_id,driver_name,origin,destination,origin_lat,origin_lng,dest_lat,dest_lng,fare,departure,seats,notes,status,created_at,COALESCE(ride_type,'airport')"
                    f" FROM rides WHERE status IN ({pg_placeholders}) ORDER BY departure ASC LIMIT 200",
                    status_filter,
                )
                rows = cur.fetchall()
            else:
                cur = conn.execute(
                    "SELECT ride_id,user_id,driver_name,origin,destination,origin_lat,origin_lng,dest_lat,dest_lng,fare,departure,seats,notes,status,created_at,COALESCE(ride_type,'airport')"
                    f" FROM rides WHERE status IN ({sql_placeholders}) ORDER BY departure ASC LIMIT 200",
                    status_filter,
                )
                rows = cur.fetchall()
        finally:
            conn.close()

    rides = [dict(zip(cols, row)) for row in rows]
    return JSONResponse({"rides": rides})


@fastapi_app.get("/api/rides/calculate_fare")
async def api_calculate_fare(
    origin_lat: float, origin_lng: float,
    dest_lat: float, dest_lng: float,
):
    """Calculate estimated fare for an airport pickup from origin to destination.

    Returns fare in local currency units at the configured rate per km.
    """
    dist_km = _haversine_km(origin_lat, origin_lng, dest_lat, dest_lng)
    fare    = round(dist_km * _FARE_PER_KM, 2)
    return JSONResponse({"dist_km": round(dist_km, 2), "fare": fare, "rate_per_km": _FARE_PER_KM})


@fastapi_app.get("/api/rides/shared_fare")
async def api_shared_fare(
    total_fare: float,
    total_seats: int,
    booked_seats: int,
):
    """Calculate the cost for a passenger booking *booked_seats* on a ride.

    Rules:
    - Full vehicle (booked_seats == total_seats): passenger pays total_fare.
    - Shared ride (booked_seats < total_seats): cost is proportional to
      booked_seats / total_seats, giving a discount for sharing.
    - Returns per-seat price and total amount owed.
    """
    if total_seats < 1:
        return JSONResponse({"error": "total_seats must be at least 1."}, status_code=400)
    if booked_seats < 1 or booked_seats > total_seats:
        return JSONResponse({"error": "booked_seats must be between 1 and total_seats."}, status_code=400)
    if total_fare < 0:
        return JSONResponse({"error": "total_fare must be non-negative."}, status_code=400)

    per_seat_cost   = round(total_fare / total_seats, 2)
    amount_owed     = round(per_seat_cost * booked_seats, 2)
    is_full_vehicle = booked_seats == total_seats

    return JSONResponse({
        "total_fare":      round(total_fare, 2),
        "total_seats":     total_seats,
        "booked_seats":    booked_seats,
        "per_seat_cost":   per_seat_cost,
        "amount_owed":     amount_owed,
        "is_full_vehicle": is_full_vehicle,
        "rate_per_km":     _FARE_PER_KM,
    })


@fastapi_app.delete("/api/rides/{ride_id}")
async def api_ride_cancel(request: Request, ride_id: str):
    """Cancel a ride (only the poster can cancel it)."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("SELECT user_id FROM rides WHERE ride_id=%s", (ride_id,))
                row = cur.fetchone()
            else:
                cur = conn.execute("SELECT user_id FROM rides WHERE ride_id=?", (ride_id,))
                row = cur.fetchone()

            if row is None:
                return JSONResponse({"error": "Ride not found."}, status_code=404)
            if row[0] != user_id:
                return JSONResponse({"error": "Not authorised."}, status_code=403)

            if USE_POSTGRES:
                cur.execute("UPDATE rides SET status='cancelled' WHERE ride_id=%s", (ride_id,))
            else:
                conn.execute("UPDATE rides SET status='cancelled' WHERE ride_id=?", (ride_id,))
            conn.commit()
        finally:
            conn.close()

    asyncio.ensure_future(sio.emit("ride_cancelled", {"ride_id": ride_id}))
    # Persist status update to bucket
    _bucket_write_json("rides", "ride", ride_id, {"ride_id": ride_id, "status": "cancelled"})
    return JSONResponse({"ok": True})


@fastapi_app.post("/api/rides/{ride_id}/take")
async def api_ride_take(request: Request, ride_id: str):
    """Mark a ride as taken (only the poster can confirm this)."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute("SELECT user_id, status FROM rides WHERE ride_id=%s", (ride_id,))
                row = cur.fetchone()
            else:
                cur = conn.execute("SELECT user_id, status FROM rides WHERE ride_id=?", (ride_id,))
                row = cur.fetchone()

            if row is None:
                return JSONResponse({"error": "Ride not found."}, status_code=404)
            if row[0] != user_id:
                return JSONResponse({"error": "Not authorised."}, status_code=403)
            if row[1] == "cancelled":
                return JSONResponse({"error": "Cannot mark a cancelled ride as taken."}, status_code=409)

            if USE_POSTGRES:
                cur.execute("UPDATE rides SET status='taken' WHERE ride_id=%s", (ride_id,))
            else:
                conn.execute("UPDATE rides SET status='taken' WHERE ride_id=?", (ride_id,))
            conn.commit()
        finally:
            conn.close()

    asyncio.ensure_future(sio.emit("ride_taken", {"ride_id": ride_id}))
    # Persist status update to bucket; a "taken" ride moves into history
    completed_at = datetime.now(timezone.utc).isoformat()
    _bucket_write_json("rides", "ride", ride_id, {"ride_id": ride_id, "status": "taken"})
    _bucket_write_json("history", "history", ride_id, {
        "ride_id": ride_id,
        "user_id": user_id,
        "status": "taken",
        "completed_at": completed_at,
    })
    # Notify the ride poster
    _create_notification(
        user_id,
        "ride_taken",
        "✅ Your Ride Has Been Taken",
        f"Good news! Your ride has been marked as taken.",
    )
    return JSONResponse({"ok": True})


@fastapi_app.get("/api/admin/rides")
async def api_admin_rides(request: Request):
    """Return ride-sharing statistics for the admin dashboard."""
    user_id = request.session.get("admin_user")
    if not user_id:
        return JSONResponse({"error": "Admin login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cols = ["ride_id", "user_id", "driver_name", "origin", "destination",
                    "departure", "seats", "notes", "status", "created_at"]
            if USE_POSTGRES:
                cur = conn.cursor()
                cur.execute(
                    "SELECT ride_id,user_id,driver_name,origin,destination,departure,seats,notes,status,created_at FROM rides ORDER BY created_at DESC LIMIT 500"
                )
                rows = cur.fetchall()
                cur.execute("SELECT status, COUNT(*) FROM rides GROUP BY status")
                counts_rows = cur.fetchall()
            else:
                cur = conn.execute(
                    "SELECT ride_id,user_id,driver_name,origin,destination,departure,seats,notes,status,created_at FROM rides ORDER BY created_at DESC LIMIT 500"
                )
                rows = cur.fetchall()
                counts_cur = conn.execute("SELECT status, COUNT(*) FROM rides GROUP BY status")
                counts_rows = counts_cur.fetchall()
        finally:
            conn.close()

    rides = [dict(zip(cols, row)) for row in rows]
    counts = {r[0]: r[1] for r in counts_rows}
    total = sum(counts.values())
    return JSONResponse({
        "rides": rides,
        "stats": {
            "total": total,
            "open": counts.get("open", 0),
            "taken": counts.get("taken", 0),
            "cancelled": counts.get("cancelled", 0),
        },
    })


# =========================================================
# DRIVER GEOLOCATION MODULE
# =========================================================

class _DriverLocationUpdate(BaseModel):
    lat:   float
    lng:   float
    empty: bool = True  # True = car is empty / available
    seats: int  = 0     # Number of empty seats available


@fastapi_app.post("/api/driver/location")
async def api_driver_location(request: Request, body: _DriverLocationUpdate):
    """Driver broadcasts their current location. Notifies nearby connected users.

    Requires an approved driver application (verified driver).
    """
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    user = _get_app_user(user_id)
    if user is None or user.get("role") != "driver":
        return JSONResponse({"error": "Only registered drivers can broadcast location."}, status_code=403)

    ts = time.time()
    with _driver_loc_lock:
        _driver_locations[user_id] = {
            "user_id":  user_id,
            "name":     user["name"],
            "lat":      body.lat,
            "lng":      body.lng,
            "empty":    body.empty,
            "seats":    body.seats,
            "verified": True,
            "ts":       ts,
        }

    # Notify nearby users who have shared their location.
    # We include the driver's coordinates in the event so the client-side can
    # do its own proximity check before surfacing the alert.
    seats_label = f"{body.seats} empty seat{'s' if body.seats != 1 else ''}" if body.empty and body.seats > 0 else ("has empty seats" if body.empty else "is occupied")
    notification = {
        "driver_id":   user_id,
        "driver_name": user["name"],
        "lat":         body.lat,
        "lng":         body.lng,
        "empty":       body.empty,
        "seats":       body.seats,
        "radius_km":   _APP_USER_PROXIMITY_KM,
        "message":     f"Driver {user['name']} is nearby and {seats_label}!",
    }
    # Emit to users who have previously identified themselves via socket.
    # For each identified socket, only notify if the user has a stored location
    # within APP_USER_PROXIMITY_KM; otherwise fall back to broadcasting.
    targets_sent = False
    with _socket_user_lock:
        identified_sids = list(_sid_to_user.items())

    async def _notify():
        nonlocal targets_sent
        for sid, uid in identified_sids:
            u = _get_app_user(uid)
            # Only notify passengers — skip all drivers (including the
            # broadcasting driver, who is also a driver role).
            if u is None or u.get("role") == "driver":
                continue
            if u.get("location_lat") is not None and u.get("location_lng") is not None:
                dist = _haversine_km(body.lat, body.lng, u["location_lat"], u["location_lng"])
                if dist <= _APP_USER_PROXIMITY_KM:
                    await sio.emit("driver_nearby", notification, room=sid)
                    targets_sent = True
        if not targets_sent:
            # No location-aware passengers connected — broadcast to all identified
            # passengers (those without a stored location), skipping all drivers.
            for sid, uid in identified_sids:
                u = _get_app_user(uid)
                if u is None or u.get("role") == "driver":
                    continue
                await sio.emit("driver_nearby", notification, room=sid)
                targets_sent = True
        if not targets_sent:
            # Absolutely no identified passengers — broadcast to all anonymous sockets
            await sio.emit("driver_nearby", notification)

    asyncio.ensure_future(_notify())

    return JSONResponse({"ok": True, "ts": ts})


@fastapi_app.get("/api/driver/nearby")
async def api_driver_nearby(request: Request, lat: float, lng: float, radius_km: float = 10.0):
    """Return drivers who are within *radius_km* of the given coordinates."""
    now = time.time()
    with _driver_loc_lock:
        # Evict stale entries
        stale = [uid for uid, d in _driver_locations.items() if now - d["ts"] > _DRIVER_LOC_TTL_SECS]
        for uid in stale:
            del _driver_locations[uid]

        nearby = []
        for d in _driver_locations.values():
            dist = _haversine_km(lat, lng, d["lat"], d["lng"])
            if dist <= radius_km:
                nearby.append({**d, "distance_km": round(dist, 2), "ts": None})

    return JSONResponse({"drivers": nearby})


@fastapi_app.get("/api/driver/locations")
async def api_driver_locations():
    """Return all active (non-stale) driver locations (public)."""
    now = time.time()
    with _driver_loc_lock:
        active = [
            {**d, "ts": None}
            for d in _driver_locations.values()
            if now - d["ts"] <= _DRIVER_LOC_TTL_SECS
        ]
    return JSONResponse({"drivers": active})


@fastapi_app.get("/api/unified_map/nearby")
async def api_unified_map_nearby(lat: float, lng: float, radius_km: float = 25.0, mode: str = "drivers"):
    """Return nearby drivers or properties sorted by distance, for the unified map."""
    if mode == "drivers":
        now = time.time()
        with _driver_loc_lock:
            active = [
                {**d, "ts": None}
                for d in _driver_locations.values()
                if now - d["ts"] <= _DRIVER_LOC_TTL_SECS
            ]
        items = []
        for d in active:
            dist = _haversine_km(lat, lng, d["lat"], d["lng"])
            if dist <= radius_km:
                items.append({
                    "id": d.get("user_id"),
                    "name": d.get("name", "Driver"),
                    "lat": d["lat"],
                    "lng": d["lng"],
                    "distance_km": round(dist, 2),
                    "empty": d.get("empty", True),
                    "seats": d.get("seats", 0),
                    "vehicle": d.get("vehicle", ""),
                    "avatar": d.get("avatar", "🚗"),
                    "rating": d.get("rating"),
                })
        items.sort(key=lambda x: x["distance_km"])
        return JSONResponse({"items": items, "mode": mode})
    else:
        # Properties mode
        _seed_properties_if_empty()
        with _db_lock:
            conn = _get_db()
            try:
                cols = ["property_id","title","description","price","address","lat","lng","images_json","status","owner_user_id","created_at"]
                cur = _execute(conn, f"SELECT {','.join(cols)} FROM properties ORDER BY created_at DESC")
                rows = cur.fetchall()
            finally:
                conn.close()
        items = []
        for row in rows:
            p = dict(zip(["property_id","title","description","price","address","lat","lng","images_json","status","owner_user_id","created_at"], row))
            if p["lat"] is None or p["lng"] is None:
                continue
            dist = _haversine_km(lat, lng, p["lat"], p["lng"])
            if dist <= radius_km:
                try:
                    p["images"] = json.loads(p["images_json"] or "[]")
                except Exception:
                    p["images"] = []
                p["cover_image"] = p["images"][0] if p["images"] else None
                p["distance_km"] = round(dist, 2)
                items.append(p)
        items.sort(key=lambda x: x["distance_km"])
        return JSONResponse({"items": items, "mode": mode})


# =========================================================
# REAL ESTATE AGENT ENDPOINTS
# =========================================================

# In-memory agent store seeded with demo data (supplements DB when empty).
_DEMO_AGENTS_SEED = [
    {"agent_id": "agent-1", "user_id": None, "name": "Alice Johnson", "avatar": "👩", "bio": "10 years experience in London residential.", "email": "alice@example.com", "phone": "+44 20 0001", "availability_status": "available", "lat": 51.515, "lng": -0.082},
    {"agent_id": "agent-2", "user_id": None, "name": "Bob Williams",  "avatar": "👨", "bio": "Commercial and residential specialist.", "email": "bob@example.com", "phone": "+44 20 0002", "availability_status": "busy",      "lat": 51.499, "lng": -0.121},
    {"agent_id": "agent-3", "user_id": None, "name": "Carol Davis",   "avatar": "👩", "bio": "Award-winning lettings agent.", "email": "carol@example.com", "phone": "+44 20 0003", "availability_status": "available", "lat": 51.527, "lng": -0.108},
    {"agent_id": "agent-4", "user_id": None, "name": "Dan Brown",     "avatar": "👨", "bio": "New-build and off-plan expert.", "email": "dan@example.com", "phone": "+44 20 0004", "availability_status": "offline",   "lat": 51.487, "lng": -0.059},
    {"agent_id": "agent-5", "user_id": None, "name": "Eva Martinez",  "avatar": "👩", "bio": "South London property specialist.", "email": "eva@example.com", "phone": "+44 20 0005", "availability_status": "available", "lat": 51.503, "lng": -0.095},
    {"agent_id": "agent-6", "user_id": None, "name": "Frank Lee",     "avatar": "👨", "bio": "Investment and buy-to-let advisor.", "email": "frank@example.com", "phone": "+44 20 0006", "availability_status": "busy",      "lat": 51.532, "lng": -0.072},
    {"agent_id": "agent-7", "user_id": None, "name": "Grace Kim",     "avatar": "👩", "bio": "First-time buyer support specialist.", "email": "grace@example.com", "phone": "+44 20 0007", "availability_status": "available", "lat": 51.497, "lng": -0.143},
    {"agent_id": "agent-8", "user_id": None, "name": "Henry Chen",    "avatar": "👨", "bio": "Luxury and high-end properties.", "email": "henry@example.com", "phone": "+44 20 0008", "availability_status": "available", "lat": 51.519, "lng": -0.128},
]
_VALID_AGENT_STATUSES = {"available", "busy", "offline"}

def _seed_agents_if_empty():
    """Seed the real_estate_agents table with demo data if it is empty."""
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(conn, "SELECT COUNT(*) FROM real_estate_agents")
            count = (cur.fetchone() or [0])[0]
            if count == 0:
                now = datetime.utcnow().isoformat()
                for a in _DEMO_AGENTS_SEED:
                    try:
                        if USE_POSTGRES:
                            conn.cursor().execute(
                                "INSERT INTO real_estate_agents (agent_id,user_id,name,avatar,bio,email,phone,availability_status,lat,lng,created_at) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) ON CONFLICT DO NOTHING",
                                (a["agent_id"], a["user_id"], a["name"], a["avatar"], a["bio"], a["email"], a["phone"], a["availability_status"], a["lat"], a["lng"], now),
                            )
                        else:
                            conn.execute(
                                "INSERT OR IGNORE INTO real_estate_agents (agent_id,user_id,name,avatar,bio,email,phone,availability_status,lat,lng,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                                (a["agent_id"], a["user_id"], a["name"], a["avatar"], a["bio"], a["email"], a["phone"], a["availability_status"], a["lat"], a["lng"], now),
                            )
                    except Exception:
                        pass
                conn.commit()
        finally:
            conn.close()


def _get_agent_row(agent_id: str) -> dict | None:
    """Fetch one agent by agent_id, including computed review stats."""
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT agent_id,user_id,name,avatar,bio,email,phone,availability_status,lat,lng,created_at FROM real_estate_agents WHERE agent_id=?"
                if not USE_POSTGRES else
                "SELECT agent_id,user_id,name,avatar,bio,email,phone,availability_status,lat,lng,created_at FROM real_estate_agents WHERE agent_id=%s",
                (agent_id,),
            )
            cols = ["agent_id","user_id","name","avatar","bio","email","phone","availability_status","lat","lng","created_at"]
            row = cur.fetchone()
            if not row:
                return None
            agent = dict(zip(cols, row))
            # Compute review stats
            cur2 = _execute(
                conn,
                "SELECT COUNT(*), COALESCE(AVG(rating),0) FROM agent_reviews WHERE agent_id=?"
                if not USE_POSTGRES else
                "SELECT COUNT(*), COALESCE(AVG(rating),0) FROM agent_reviews WHERE agent_id=%s",
                (agent_id,),
            )
            r = cur2.fetchone()
            agent["review_count"] = int(r[0]) if r else 0
            agent["avg_rating"]   = round(float(r[1]), 1) if r else 0.0
            # Like count
            cur3 = _execute(
                conn,
                "SELECT COUNT(*) FROM agent_likes WHERE agent_id=?"
                if not USE_POSTGRES else
                "SELECT COUNT(*) FROM agent_likes WHERE agent_id=%s",
                (agent_id,),
            )
            r3 = cur3.fetchone()
            agent["like_count"] = int(r3[0]) if r3 else 0
            return agent
        finally:
            conn.close()


@fastapi_app.get("/api/agents")
async def api_list_agents(status: str | None = None):
    """List real estate agents, optionally filtered by availability_status."""
    _seed_agents_if_empty()
    with _db_lock:
        conn = _get_db()
        try:
            if status and status in _VALID_AGENT_STATUSES:
                cur = _execute(
                    conn,
                    "SELECT agent_id,user_id,name,avatar,bio,email,phone,availability_status,lat,lng,created_at FROM real_estate_agents WHERE availability_status=? ORDER BY name"
                    if not USE_POSTGRES else
                    "SELECT agent_id,user_id,name,avatar,bio,email,phone,availability_status,lat,lng,created_at FROM real_estate_agents WHERE availability_status=%s ORDER BY name",
                    (status,),
                )
            else:
                cur = _execute(
                    conn,
                    "SELECT agent_id,user_id,name,avatar,bio,email,phone,availability_status,lat,lng,created_at FROM real_estate_agents ORDER BY CASE availability_status WHEN 'available' THEN 0 WHEN 'busy' THEN 1 ELSE 2 END, name",
                )
            cols = ["agent_id","user_id","name","avatar","bio","email","phone","availability_status","lat","lng","created_at"]
            agents = [dict(zip(cols, r)) for r in cur.fetchall()]
        finally:
            conn.close()
    # Attach review/like stats
    for a in agents:
        row = _get_agent_row(a["agent_id"])
        if row:
            a["review_count"] = row["review_count"]
            a["avg_rating"]   = row["avg_rating"]
            a["like_count"]   = row["like_count"]
    return JSONResponse({"agents": agents})


@fastapi_app.get("/api/agents/{agent_id}")
async def api_get_agent(agent_id: str):
    """Get agent profile detail including reviews."""
    _seed_agents_if_empty()
    agent = _get_agent_row(agent_id)
    if not agent:
        return JSONResponse({"error": "Agent not found."}, status_code=404)
    # Fetch reviews
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT review_id,reviewer_name,rating,text,created_at FROM agent_reviews WHERE agent_id=? ORDER BY created_at DESC LIMIT 50"
                if not USE_POSTGRES else
                "SELECT review_id,reviewer_name,rating,text,created_at FROM agent_reviews WHERE agent_id=%s ORDER BY created_at DESC LIMIT 50",
                (agent_id,),
            )
            rcols = ["review_id","reviewer_name","rating","text","created_at"]
            agent["reviews"] = [dict(zip(rcols, r)) for r in cur.fetchall()]
        finally:
            conn.close()
    return JSONResponse({"agent": agent})


class _AgentStatusUpdate(BaseModel):
    status: str

@fastapi_app.put("/api/agents/{agent_id}/status")
async def api_update_agent_status(request: Request, agent_id: str, body: _AgentStatusUpdate):
    """Update an agent's availability status. Requires login (agent updates own status, or admin)."""
    user_id = request.session.get("app_user_id")
    is_admin = request.session.get("admin_logged_in") or request.session.get("admin_user")
    if not user_id and not is_admin:
        return JSONResponse({"error": "Login required."}, status_code=401)
    status = body.status.strip().lower()
    if status not in _VALID_AGENT_STATUSES:
        return JSONResponse({"error": f"Invalid status. Must be one of: {', '.join(_VALID_AGENT_STATUSES)}"}, status_code=400)
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT agent_id, user_id FROM real_estate_agents WHERE agent_id=?"
                if not USE_POSTGRES else
                "SELECT agent_id, user_id FROM real_estate_agents WHERE agent_id=%s",
                (agent_id,),
            )
            row = cur.fetchone()
            if not row:
                conn.close()
                return JSONResponse({"error": "Agent not found."}, status_code=404)
            # Only allow if admin OR the agent's linked user_id matches
            if not is_admin and row[1] != user_id:
                conn.close()
                return JSONResponse({"error": "Forbidden."}, status_code=403)
            _execute(
                conn,
                "UPDATE real_estate_agents SET availability_status=? WHERE agent_id=?"
                if not USE_POSTGRES else
                "UPDATE real_estate_agents SET availability_status=%s WHERE agent_id=%s",
                (status, agent_id),
            )
            conn.commit()
        finally:
            conn.close()
    # Broadcast status change via socket
    asyncio.ensure_future(sio.emit("agent_status_changed", {"agent_id": agent_id, "status": status}))
    return JSONResponse({"ok": True, "status": status})


class _AgentReviewRequest(BaseModel):
    rating: int
    text: str = ""

@fastapi_app.post("/api/agents/{agent_id}/review")
async def api_submit_agent_review(request: Request, agent_id: str, body: _AgentReviewRequest):
    """Submit a review for an agent. Requires login."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    if not (1 <= body.rating <= 5):
        return JSONResponse({"error": "Rating must be between 1 and 5."}, status_code=400)
    user = _get_app_user(user_id)
    if not user:
        return JSONResponse({"error": "User not found."}, status_code=404)
    agent = _get_agent_row(agent_id)
    if not agent:
        return JSONResponse({"error": "Agent not found."}, status_code=404)
    review_id  = str(uuid.uuid4())
    now        = datetime.utcnow().isoformat()
    text       = str(body.text or "").strip()[:500]
    reviewer_name = user["name"]
    with _db_lock:
        conn = _get_db()
        try:
            if USE_POSTGRES:
                conn.cursor().execute(
                    "INSERT INTO agent_reviews (review_id,agent_id,reviewer_user_id,reviewer_name,rating,text,created_at) VALUES (%s,%s,%s,%s,%s,%s,%s)",
                    (review_id, agent_id, user_id, reviewer_name, body.rating, text, now),
                )
            else:
                conn.execute(
                    "INSERT INTO agent_reviews (review_id,agent_id,reviewer_user_id,reviewer_name,rating,text,created_at) VALUES (?,?,?,?,?,?,?)",
                    (review_id, agent_id, user_id, reviewer_name, body.rating, text, now),
                )
            conn.commit()
        finally:
            conn.close()
    # Create notification for agent's linked user if any
    if agent.get("user_id"):
        _create_notification(
            agent["user_id"],
            "agent_review",
            f"New review from {reviewer_name}",
            f"{reviewer_name} gave you {body.rating}★" + (f": {text[:80]}" if text else ""),
        )
    return JSONResponse({"ok": True, "review_id": review_id})


@fastapi_app.post("/api/agents/{agent_id}/like")
async def api_like_agent(request: Request, agent_id: str):
    """Toggle a like for an agent. Requires login."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    agent = _get_agent_row(agent_id)
    if not agent:
        return JSONResponse({"error": "Agent not found."}, status_code=404)
    now = datetime.utcnow().isoformat()
    liked = False
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT id FROM agent_likes WHERE agent_id=? AND user_id=?"
                if not USE_POSTGRES else
                "SELECT id FROM agent_likes WHERE agent_id=%s AND user_id=%s",
                (agent_id, user_id),
            )
            existing = cur.fetchone()
            if existing:
                _execute(
                    conn,
                    "DELETE FROM agent_likes WHERE agent_id=? AND user_id=?"
                    if not USE_POSTGRES else
                    "DELETE FROM agent_likes WHERE agent_id=%s AND user_id=%s",
                    (agent_id, user_id),
                )
                liked = False
            else:
                if USE_POSTGRES:
                    conn.cursor().execute(
                        "INSERT INTO agent_likes (agent_id,user_id,created_at) VALUES (%s,%s,%s) ON CONFLICT DO NOTHING",
                        (agent_id, user_id, now),
                    )
                else:
                    conn.execute(
                        "INSERT OR IGNORE INTO agent_likes (agent_id,user_id,created_at) VALUES (?,?,?)",
                        (agent_id, user_id, now),
                    )
                liked = True
            conn.commit()
        finally:
            conn.close()
    # Notify agent's linked user if liked
    user = _get_app_user(user_id)
    if liked and agent.get("user_id") and user:
        _create_notification(
            agent["user_id"],
            "agent_like",
            "Someone liked your profile",
            f"{user['name']} liked your agent profile.",
        )
    updated = _get_agent_row(agent_id)
    return JSONResponse({"ok": True, "liked": liked, "like_count": updated["like_count"] if updated else 0})


@fastapi_app.get("/api/agents/{agent_id}/chat")
async def api_get_agent_chat(request: Request, agent_id: str):
    """Get chat message history between the current user and an agent."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT msg_id,agent_id,user_id,sender_role,text,ts,created_at FROM agent_chat_messages WHERE agent_id=? AND user_id=? ORDER BY ts ASC LIMIT 200"
                if not USE_POSTGRES else
                "SELECT msg_id,agent_id,user_id,sender_role,text,ts,created_at FROM agent_chat_messages WHERE agent_id=%s AND user_id=%s ORDER BY ts ASC LIMIT 200",
                (agent_id, user_id),
            )
            cols = ["msg_id","agent_id","user_id","sender_role","text","ts","created_at"]
            messages = [dict(zip(cols, r)) for r in cur.fetchall()]
        finally:
            conn.close()
    return JSONResponse({"messages": messages})


# =========================================================
# PROPERTY ENDPOINTS
# =========================================================

_VALID_PROPERTY_STATUSES = {"active", "sold", "rented"}

# Demo property seed data (linked to the demo agents)
_DEMO_PROPERTIES_SEED = [
    {
        "property_id": "prop-1",
        "title": "Modern 2-Bed Flat – Shoreditch",
        "description": "A bright, modern two-bedroom flat in the heart of Shoreditch. Open-plan kitchen/living area, private balcony, and secure parking.",
        "price": 2200,
        "address": "12 Oak Street, London E1",
        "lat": 51.522, "lng": -0.074,
        "images_json": '["https://images.unsplash.com/photo-1558618666-fcd25c85cd64?w=800","https://images.unsplash.com/photo-1484154218962-a197022b5858?w=800"]',
        "status": "active",
        "agent_ids": ["agent-1", "agent-3"],
    },
    {
        "property_id": "prop-2",
        "title": "Spacious 3-Bed House – Islington",
        "description": "Beautiful Victorian terrace with three double bedrooms, two bathrooms, and a south-facing garden. Walking distance to Angel Tube.",
        "price": 3500,
        "address": "8 Maple Avenue, London N1",
        "lat": 51.533, "lng": -0.103,
        "images_json": '["https://images.unsplash.com/photo-1568605114967-8130f3a36994?w=800","https://images.unsplash.com/photo-1570129477492-45c003edd2be?w=800"]',
        "status": "active",
        "agent_ids": ["agent-2", "agent-5"],
    },
    {
        "property_id": "prop-3",
        "title": "Studio Apartment – Canary Wharf",
        "description": "Compact but stylish studio in Canary Wharf. Floor-to-ceiling windows with river views, gym access included, concierge service.",
        "price": 1600,
        "address": "34 Harbour Way, London E14",
        "lat": 51.503, "lng": -0.017,
        "images_json": '["https://images.unsplash.com/photo-1502672260266-1c1ef2d93688?w=800","https://images.unsplash.com/photo-1493809842364-78817add7ffb?w=800"]',
        "status": "active",
        "agent_ids": ["agent-7"],
    },
    {
        "property_id": "prop-4",
        "title": "Luxury 4-Bed Villa – Richmond",
        "description": "Stunning detached family home close to Richmond Park. Four bedrooms, three reception rooms, double garage, and landscaped garden.",
        "price": 6500,
        "address": "5 Elm Close, Richmond TW10",
        "lat": 51.461, "lng": -0.301,
        "images_json": '["https://images.unsplash.com/photo-1613977257363-707ba9348227?w=800","https://images.unsplash.com/photo-1512917774080-9991f1c4c750?w=800"]',
        "status": "active",
        "agent_ids": ["agent-8", "agent-6", "agent-1"],
    },
    {
        "property_id": "prop-5",
        "title": "1-Bed Flat – Clapham",
        "description": "Well-presented one-bedroom flat in popular Clapham. Recently refurbished, modern kitchen, private roof terrace, near tube and park.",
        "price": 1800,
        "address": "21 Birch Lane, London SW4",
        "lat": 51.461, "lng": -0.138,
        "images_json": '["https://images.unsplash.com/photo-1586023492125-27b2c045efd7?w=800","https://images.unsplash.com/photo-1505691938895-1758d7feb511?w=800"]',
        "status": "active",
        "agent_ids": ["agent-4", "agent-2"],
    },
    {
        "property_id": "prop-6",
        "title": "2-Bed New Build – Stratford",
        "description": "Brand-new two-bedroom apartment in Stratford development. High spec throughout, balcony overlooking park, excellent transport links.",
        "price": 2400,
        "address": "9 Cedar Court, London E20",
        "lat": 51.541, "lng": -0.002,
        "images_json": '["https://images.unsplash.com/photo-1545324418-cc1a3fa10c00?w=800","https://images.unsplash.com/photo-1554995207-c18c203602cb?w=800"]',
        "status": "sold",
        "agent_ids": ["agent-3", "agent-8"],
    },
]


def _seed_properties_if_empty():
    """Seed the properties table with demo data if it is empty."""
    _seed_agents_if_empty()
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(conn, "SELECT COUNT(*) FROM properties")
            count = (cur.fetchone() or [0])[0]
            if count == 0:
                now = datetime.utcnow().isoformat()
                for p in _DEMO_PROPERTIES_SEED:
                    try:
                        if USE_POSTGRES:
                            conn.cursor().execute(
                                "INSERT INTO properties (property_id,title,description,price,address,lat,lng,images_json,status,owner_user_id,created_at) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s) ON CONFLICT DO NOTHING",
                                (p["property_id"], p["title"], p["description"], p["price"], p["address"], p["lat"], p["lng"], p["images_json"], p["status"], None, now),
                            )
                        else:
                            conn.execute(
                                "INSERT OR IGNORE INTO properties (property_id,title,description,price,address,lat,lng,images_json,status,owner_user_id,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                                (p["property_id"], p["title"], p["description"], p["price"], p["address"], p["lat"], p["lng"], p["images_json"], p["status"], None, now),
                            )
                        for pos, agent_id in enumerate(p.get("agent_ids", [])):
                            try:
                                if USE_POSTGRES:
                                    conn.cursor().execute(
                                        "INSERT INTO property_agents (property_id,agent_id,position) VALUES (%s,%s,%s) ON CONFLICT DO NOTHING",
                                        (p["property_id"], agent_id, pos),
                                    )
                                else:
                                    conn.execute(
                                        "INSERT OR IGNORE INTO property_agents (property_id,agent_id,position) VALUES (?,?,?)",
                                        (p["property_id"], agent_id, pos),
                                    )
                            except Exception:
                                pass
                    except Exception:
                        pass
                conn.commit()
        finally:
            conn.close()


def _get_property_agents(property_id: str) -> list[dict]:
    """Return up to 4 agent objects linked to a property, with review stats."""
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT agent_id FROM property_agents WHERE property_id=? ORDER BY position ASC LIMIT 4"
                if not USE_POSTGRES else
                "SELECT agent_id FROM property_agents WHERE property_id=%s ORDER BY position ASC LIMIT 4",
                (property_id,),
            )
            agent_ids = [r[0] for r in cur.fetchall()]
        finally:
            conn.close()
    return [a for a in (_get_agent_row(aid) for aid in agent_ids) if a]


def _get_property_row(property_id: str, with_agents: bool = False) -> dict | None:
    """Fetch one property by property_id."""
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT property_id,title,description,price,address,lat,lng,images_json,status,owner_user_id,created_at FROM properties WHERE property_id=?"
                if not USE_POSTGRES else
                "SELECT property_id,title,description,price,address,lat,lng,images_json,status,owner_user_id,created_at FROM properties WHERE property_id=%s",
                (property_id,),
            )
            cols = ["property_id","title","description","price","address","lat","lng","images_json","status","owner_user_id","created_at"]
            row = cur.fetchone()
            if not row:
                return None
            prop = dict(zip(cols, row))
        finally:
            conn.close()
    try:
        prop["images"] = json.loads(prop["images_json"] or "[]")
    except Exception:
        prop["images"] = []
    if with_agents:
        prop["agents"] = _get_property_agents(property_id)
    return prop


class _PropertyCreateRequest(BaseModel):
    title: str
    description: str = ""
    price: float = 0.0
    address: str = ""
    lat: float | None = None
    lng: float | None = None
    images: list[str] = []
    agent_ids: list[str] = []
    status: str = "active"


class _PropertyUpdateRequest(BaseModel):
    title: str | None = None
    description: str | None = None
    price: float | None = None
    address: str | None = None
    lat: float | None = None
    lng: float | None = None
    images: list[str] | None = None
    agent_ids: list[str] | None = None
    status: str | None = None


@fastapi_app.get("/api/properties")
async def api_list_properties(
    status: str | None = None,
    min_lat: float | None = None,
    max_lat: float | None = None,
    min_lng: float | None = None,
    max_lng: float | None = None,
):
    """List properties, optionally filtered by status and/or map bounds."""
    _seed_properties_if_empty()
    with _db_lock:
        conn = _get_db()
        try:
            ph = "%s" if USE_POSTGRES else "?"
            parts = []
            params = []
            if status:
                parts.append(f"status={ph}")
                params.append(status)
            if min_lat is not None:
                parts.append(f"lat>={ph}")
                params.append(min_lat)
            if max_lat is not None:
                parts.append(f"lat<={ph}")
                params.append(max_lat)
            if min_lng is not None:
                parts.append(f"lng>={ph}")
                params.append(min_lng)
            if max_lng is not None:
                parts.append(f"lng<={ph}")
                params.append(max_lng)
            where = (" WHERE " + " AND ".join(parts)) if parts else ""
            cols = ["property_id","title","description","price","address","lat","lng","images_json","status","owner_user_id","created_at"]
            cur = _execute(conn, f"SELECT {','.join(cols)} FROM properties{where} ORDER BY created_at DESC", params)
            rows = cur.fetchall()
        finally:
            conn.close()

    properties = []
    for row in rows:
        p = dict(zip(["property_id","title","description","price","address","lat","lng","images_json","status","owner_user_id","created_at"], row))
        try:
            p["images"] = json.loads(p["images_json"] or "[]")
        except Exception:
            p["images"] = []
        # Return first image as cover
        p["cover_image"] = p["images"][0] if p["images"] else None
        properties.append(p)
    return JSONResponse({"properties": properties})


@fastapi_app.get("/api/properties/{property_id}/map_preview")
async def api_property_map_preview(property_id: str):
    """Return minimal property data sufficient to render a teaser map pin.

    This endpoint is intentionally unauthenticated so that external links and
    search-result cards can show a focused map preview (location pin + basic
    context) to non-logged-in visitors.  Only non-sensitive fields are
    returned – full description, agent profiles, and contact details require
    authentication via the full /api/properties/{property_id} endpoint.
    """
    _seed_properties_if_empty()
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT property_id,title,address,lat,lng,status FROM properties WHERE property_id=? LIMIT 1"
                if not USE_POSTGRES else
                "SELECT property_id,title,address,lat,lng,status FROM properties WHERE property_id=%s LIMIT 1",
                (property_id,),
            )
            row = cur.fetchone()
        finally:
            conn.close()
    if not row:
        return JSONResponse({"error": "Property not found."}, status_code=404)
    cols = ["property_id", "title", "address", "lat", "lng", "status"]
    return JSONResponse({"preview": dict(zip(cols, row))})


@fastapi_app.get("/api/properties/{property_id}")
async def api_get_property(property_id: str):
    """Get full details of a property including linked agents."""
    prop = _get_property_row(property_id, with_agents=True)
    if not prop:
        return JSONResponse({"error": "Property not found."}, status_code=404)
    return JSONResponse({"property": prop})


@fastapi_app.post("/api/properties")
async def api_create_property(request: Request, body: _PropertyCreateRequest):
    """Create a new property listing (approved agents only)."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    # Check that the user has agent posting permission
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT can_post_properties FROM app_users WHERE user_id=?"
                if not USE_POSTGRES else
                "SELECT can_post_properties FROM app_users WHERE user_id=%s",
                (user_id,),
            )
            row = cur.fetchone()
        finally:
            conn.close()
    can_post = (row["can_post_properties"] if USE_POSTGRES else (row[0] if row else 0)) or 0
    if not can_post:
        return JSONResponse({"error": "Only approved agents can post properties."}, status_code=403)
    title = body.title.strip()[:200]
    if not title:
        return JSONResponse({"error": "Title is required."}, status_code=400)
    if body.lat is None or body.lng is None:
        return JSONResponse({"error": "Property location (lat/lng) is required."}, status_code=400)
    if body.status not in _VALID_PROPERTY_STATUSES:
        return JSONResponse({"error": f"Invalid status. Must be one of: {', '.join(sorted(_VALID_PROPERTY_STATUSES))}"}, status_code=400)
    agent_ids = list(dict.fromkeys(body.agent_ids))[:4]  # deduplicate, max 4
    property_id = str(uuid.uuid4())
    now = datetime.now(timezone.utc).isoformat()
    images_json = json.dumps(body.images[:20])

    with _db_lock:
        conn = _get_db()
        try:
            _execute(
                conn,
                "INSERT INTO properties (property_id,title,description,price,address,lat,lng,images_json,status,owner_user_id,created_at) VALUES (?,?,?,?,?,?,?,?,?,?,?)"
                if not USE_POSTGRES else
                "INSERT INTO properties (property_id,title,description,price,address,lat,lng,images_json,status,owner_user_id,created_at) VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)",
                (property_id, title, body.description.strip(), body.price, body.address.strip(), body.lat, body.lng, images_json, body.status, user_id, now),
            )
            for pos, agent_id in enumerate(agent_ids):
                try:
                    _execute(
                        conn,
                        "INSERT OR IGNORE INTO property_agents (property_id,agent_id,position) VALUES (?,?,?)"
                        if not USE_POSTGRES else
                        "INSERT INTO property_agents (property_id,agent_id,position) VALUES (%s,%s,%s) ON CONFLICT DO NOTHING",
                        (property_id, agent_id, pos),
                    )
                except Exception:
                    pass
            conn.commit()
        finally:
            conn.close()

    prop = _get_property_row(property_id, with_agents=True)
    # Persist property record to bucket under /properties/
    if prop:
        _bucket_write_json("properties", "property", property_id, prop)
    return JSONResponse({"ok": True, "property": prop}, status_code=201)


@fastapi_app.put("/api/properties/{property_id}")
async def api_update_property(request: Request, property_id: str, body: _PropertyUpdateRequest):
    """Update a property (owner or admin only)."""
    user_id = request.session.get("app_user_id")
    admin   = request.session.get("admin_user")
    if not user_id and not admin:
        return JSONResponse({"error": "Login required."}, status_code=401)

    prop = _get_property_row(property_id)
    if not prop:
        return JSONResponse({"error": "Property not found."}, status_code=404)
    if not admin and prop["owner_user_id"] != user_id:
        return JSONResponse({"error": "Access denied."}, status_code=403)

    if body.status is not None and body.status not in _VALID_PROPERTY_STATUSES:
        return JSONResponse({"error": f"Invalid status."}, status_code=400)

    updates: dict = {}
    if body.title       is not None: updates["title"]       = body.title.strip()[:200]
    if body.description is not None: updates["description"] = body.description.strip()
    if body.price       is not None: updates["price"]       = body.price
    if body.address     is not None: updates["address"]     = body.address.strip()
    if body.lat         is not None: updates["lat"]         = body.lat
    if body.lng         is not None: updates["lng"]         = body.lng
    if body.status      is not None: updates["status"]      = body.status
    if body.images      is not None: updates["images_json"] = json.dumps(body.images[:20])

    if updates:
        ph = "%s" if USE_POSTGRES else "?"
        set_clause = ", ".join(f"{k}={ph}" for k in updates)
        vals = list(updates.values()) + [property_id]
        with _db_lock:
            conn = _get_db()
            try:
                _execute(
                    conn,
                    f"UPDATE properties SET {set_clause} WHERE property_id={'%s' if USE_POSTGRES else '?'}",
                    vals,
                )
                conn.commit()
            finally:
                conn.close()

    if body.agent_ids is not None:
        agent_ids = list(dict.fromkeys(body.agent_ids))[:4]
        with _db_lock:
            conn = _get_db()
            try:
                _execute(
                    conn,
                    "DELETE FROM property_agents WHERE property_id=?"
                    if not USE_POSTGRES else
                    "DELETE FROM property_agents WHERE property_id=%s",
                    (property_id,),
                )
                for pos, agent_id in enumerate(agent_ids):
                    try:
                        _execute(
                            conn,
                            "INSERT OR IGNORE INTO property_agents (property_id,agent_id,position) VALUES (?,?,?)"
                            if not USE_POSTGRES else
                            "INSERT INTO property_agents (property_id,agent_id,position) VALUES (%s,%s,%s) ON CONFLICT DO NOTHING",
                            (property_id, agent_id, pos),
                        )
                    except Exception:
                        pass
                conn.commit()
            finally:
                conn.close()

    prop = _get_property_row(property_id, with_agents=True)
    # Persist updated property record to bucket (overwrites previous snapshot)
    if prop:
        _bucket_write_json("properties", "property", property_id, prop)
    return JSONResponse({"ok": True, "property": prop})


@fastapi_app.delete("/api/properties/{property_id}")
async def api_delete_property(request: Request, property_id: str):
    """Delete a property listing (owner or admin only)."""
    user_id = request.session.get("app_user_id")
    admin   = request.session.get("admin_user")
    if not user_id and not admin:
        return JSONResponse({"error": "Login required."}, status_code=401)

    prop = _get_property_row(property_id)
    if not prop:
        return JSONResponse({"error": "Property not found."}, status_code=404)
    if not admin and prop["owner_user_id"] != user_id:
        return JSONResponse({"error": "Access denied."}, status_code=403)

    with _db_lock:
        conn = _get_db()
        try:
            _execute(
                conn,
                "DELETE FROM property_agents WHERE property_id=?"
                if not USE_POSTGRES else
                "DELETE FROM property_agents WHERE property_id=%s",
                (property_id,),
            )
            _execute(
                conn,
                "DELETE FROM properties WHERE property_id=?"
                if not USE_POSTGRES else
                "DELETE FROM properties WHERE property_id=%s",
                (property_id,),
            )
            conn.commit()
        finally:
            conn.close()
    return JSONResponse({"ok": True})


# ── Property Conversations (Inbox) ─────────────────────────────────────────────

class _PropConvStartRequest(BaseModel):
    property_id: str
    agent_id: str


class _PropMsgSendRequest(BaseModel):
    conv_id: str
    content: str


def _get_property_conversation(conv_id: str) -> dict | None:
    """Fetch a property conversation by conv_id."""
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT conv_id,property_id,user_id,agent_id,unread_user,unread_agent,created_at FROM property_conversations WHERE conv_id=?"
                if not USE_POSTGRES else
                "SELECT conv_id,property_id,user_id,agent_id,unread_user,unread_agent,created_at FROM property_conversations WHERE conv_id=%s",
                (conv_id,),
            )
            row = cur.fetchone()
            if not row:
                return None
            return dict(zip(["conv_id","property_id","user_id","agent_id","unread_user","unread_agent","created_at"], row))
        finally:
            conn.close()


def _find_or_create_property_conversation(property_id: str, user_id: str, agent_id: str) -> dict:
    """Find or create a property conversation for user↔agent on a property."""
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT conv_id,property_id,user_id,agent_id,unread_user,unread_agent,created_at FROM property_conversations WHERE property_id=? AND user_id=? AND agent_id=?"
                if not USE_POSTGRES else
                "SELECT conv_id,property_id,user_id,agent_id,unread_user,unread_agent,created_at FROM property_conversations WHERE property_id=%s AND user_id=%s AND agent_id=%s",
                (property_id, user_id, agent_id),
            )
            row = cur.fetchone()
            if row:
                return dict(zip(["conv_id","property_id","user_id","agent_id","unread_user","unread_agent","created_at"], row))
            conv_id = str(uuid.uuid4())
            now = datetime.now(timezone.utc).isoformat()
            _execute(
                conn,
                "INSERT INTO property_conversations (conv_id,property_id,user_id,agent_id,unread_user,unread_agent,created_at) VALUES (?,?,?,?,0,0,?)"
                if not USE_POSTGRES else
                "INSERT INTO property_conversations (conv_id,property_id,user_id,agent_id,unread_user,unread_agent,created_at) VALUES (%s,%s,%s,%s,0,0,%s)",
                (conv_id, property_id, user_id, agent_id, now),
            )
            conn.commit()
            return {"conv_id": conv_id, "property_id": property_id, "user_id": user_id, "agent_id": agent_id, "unread_user": 0, "unread_agent": 0, "created_at": now}
        finally:
            conn.close()


@fastapi_app.get("/api/property_conversations")
async def api_list_property_conversations(request: Request):
    """List all property inbox conversations for the current user (as buyer or agent)."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT conv_id,property_id,user_id,agent_id,unread_user,unread_agent,created_at FROM property_conversations WHERE user_id=? OR agent_id=?"
                if not USE_POSTGRES else
                "SELECT conv_id,property_id,user_id,agent_id,unread_user,unread_agent,created_at FROM property_conversations WHERE user_id=%s OR agent_id=%s",
                (user_id, user_id),
            )
            rows = cur.fetchall()
        finally:
            conn.close()

    conversations = []
    for row in rows:
        conv = dict(zip(["conv_id","property_id","user_id","agent_id","unread_user","unread_agent","created_at"], row))
        prop = _get_property_row(conv["property_id"])
        agent = _get_agent_row(conv["agent_id"])
        other_user = _get_app_user(conv["user_id"])
        unread = conv["unread_user"] if conv["user_id"] == user_id else conv["unread_agent"]

        # Last message
        with _db_lock:
            conn = _get_db()
            try:
                cur = _execute(
                    conn,
                    "SELECT msg_id,sender_id,sender_role,content,ts FROM property_messages WHERE conv_id=? ORDER BY ts DESC LIMIT 1"
                    if not USE_POSTGRES else
                    "SELECT msg_id,sender_id,sender_role,content,ts FROM property_messages WHERE conv_id=%s ORDER BY ts DESC LIMIT 1",
                    (conv["conv_id"],),
                )
                lm = cur.fetchone()
            finally:
                conn.close()

        last_msg = None
        if lm:
            last_msg = dict(zip(["msg_id","sender_id","sender_role","content","ts"], lm))

        conversations.append({
            "conv_id":      conv["conv_id"],
            "property":     {"property_id": conv["property_id"], "title": prop["title"] if prop else conv["property_id"], "cover_image": prop["images"][0] if prop and prop.get("images") else None},
            "agent":        {"agent_id": conv["agent_id"], "name": agent["name"] if agent else conv["agent_id"], "avatar": agent["avatar"] if agent else "👤"},
            "other_user":   {"user_id": conv["user_id"], "name": other_user["name"] if other_user else conv["user_id"]},
            "unread_count": unread,
            "last_message": last_msg,
            "created_at":   conv["created_at"],
            "role":         "user" if conv["user_id"] == user_id else "agent",
        })

    conversations.sort(key=lambda c: (c["last_message"]["ts"] if c["last_message"] else 0), reverse=True)
    return JSONResponse({"conversations": conversations})


@fastapi_app.post("/api/property_conversations")
async def api_start_property_conversation(request: Request, body: _PropConvStartRequest):
    """Find or create a property conversation (user contacting an agent about a property)."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    prop = _get_property_row(body.property_id)
    if not prop:
        return JSONResponse({"error": "Property not found."}, status_code=404)
    agent = _get_agent_row(body.agent_id)
    if not agent:
        return JSONResponse({"error": "Agent not found."}, status_code=404)
    conv = _find_or_create_property_conversation(body.property_id, user_id, body.agent_id)
    return JSONResponse({"conv": conv})


@fastapi_app.get("/api/property_conversations/{conv_id}/messages")
async def api_get_property_messages(request: Request, conv_id: str):
    """Get messages in a property conversation."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    conv = _get_property_conversation(conv_id)
    if not conv:
        return JSONResponse({"error": "Conversation not found."}, status_code=404)
    if user_id not in (conv["user_id"], conv["agent_id"]):
        return JSONResponse({"error": "Access denied."}, status_code=403)

    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT msg_id,conv_id,sender_id,sender_role,content,ts,created_at FROM property_messages WHERE conv_id=? ORDER BY ts ASC LIMIT 200"
                if not USE_POSTGRES else
                "SELECT msg_id,conv_id,sender_id,sender_role,content,ts,created_at FROM property_messages WHERE conv_id=%s ORDER BY ts ASC LIMIT 200",
                (conv_id,),
            )
            cols = ["msg_id","conv_id","sender_id","sender_role","content","ts","created_at"]
            messages = [dict(zip(cols, r)) for r in cur.fetchall()]
        finally:
            conn.close()

    return JSONResponse({"messages": messages, "conv": conv})


@fastapi_app.post("/api/property_messages")
async def api_send_property_message(request: Request, body: _PropMsgSendRequest):
    """Send a message in a property conversation."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    content = body.content.strip()[:2000]
    if not content:
        return JSONResponse({"error": "Message cannot be empty."}, status_code=400)
    conv = _get_property_conversation(body.conv_id)
    if not conv:
        return JSONResponse({"error": "Conversation not found."}, status_code=404)
    if user_id not in (conv["user_id"], conv["agent_id"]):
        return JSONResponse({"error": "Access denied."}, status_code=403)

    msg_id      = str(uuid.uuid4())
    ts          = time.time()
    now         = datetime.now(timezone.utc).isoformat()
    sender_role = "user" if user_id == conv["user_id"] else "agent"

    with _db_lock:
        conn = _get_db()
        try:
            _execute(
                conn,
                "INSERT OR IGNORE INTO property_messages (msg_id,conv_id,sender_id,sender_role,content,ts,created_at) VALUES (?,?,?,?,?,?,?)"
                if not USE_POSTGRES else
                "INSERT INTO property_messages (msg_id,conv_id,sender_id,sender_role,content,ts,created_at) VALUES (%s,%s,%s,%s,%s,%s,%s) ON CONFLICT (msg_id) DO NOTHING",
                (msg_id, body.conv_id, user_id, sender_role, content, ts, now),
            )
            # Increment unread for the other participant
            if sender_role == "user":
                _execute(
                    conn,
                    "UPDATE property_conversations SET unread_agent=unread_agent+1 WHERE conv_id=?"
                    if not USE_POSTGRES else
                    "UPDATE property_conversations SET unread_agent=unread_agent+1 WHERE conv_id=%s",
                    (body.conv_id,),
                )
            else:
                _execute(
                    conn,
                    "UPDATE property_conversations SET unread_user=unread_user+1 WHERE conv_id=?"
                    if not USE_POSTGRES else
                    "UPDATE property_conversations SET unread_user=unread_user+1 WHERE conv_id=%s",
                    (body.conv_id,),
                )
            conn.commit()
        finally:
            conn.close()

    msg = {
        "msg_id":      msg_id,
        "conv_id":     body.conv_id,
        "sender_id":   user_id,
        "sender_role": sender_role,
        "content":     content,
        "ts":          ts,
    }

    # Real-time delivery to conversation room
    room = f"prop_conv_{body.conv_id}"
    await sio.emit("property_message", msg, room=room)

    # Notify the other participant if online
    other_id = conv["agent_id"] if sender_role == "user" else conv["user_id"]
    with _socket_user_lock:
        other_sid = _user_to_sid.get(other_id)
    if other_sid:
        me = _get_app_user(user_id)
        sender_name = me["name"] if me else "Someone"
        prop = _get_property_row(conv["property_id"])
        prop_title = prop["title"] if prop else "a property"
        await sio.emit(
            "property_message_notification",
            {"conv_id": body.conv_id, "from": sender_name, "preview": content[:80], "property_title": prop_title},
            room=other_sid,
        )

    return JSONResponse({"ok": True, "message": msg})


@fastapi_app.post("/api/property_conversations/{conv_id}/read")
async def api_property_conversation_read(request: Request, conv_id: str):
    """Mark messages in a property conversation as read for the current user."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    conv = _get_property_conversation(conv_id)
    if not conv:
        return JSONResponse({"error": "Conversation not found."}, status_code=404)
    if user_id not in (conv["user_id"], conv["agent_id"]):
        return JSONResponse({"error": "Access denied."}, status_code=403)

    with _db_lock:
        conn = _get_db()
        try:
            if user_id == conv["user_id"]:
                _execute(
                    conn,
                    "UPDATE property_conversations SET unread_user=0 WHERE conv_id=?"
                    if not USE_POSTGRES else
                    "UPDATE property_conversations SET unread_user=0 WHERE conv_id=%s",
                    (conv_id,),
                )
            else:
                _execute(
                    conn,
                    "UPDATE property_conversations SET unread_agent=0 WHERE conv_id=?"
                    if not USE_POSTGRES else
                    "UPDATE property_conversations SET unread_agent=0 WHERE conv_id=%s",
                    (conv_id,),
                )
            conn.commit()
        finally:
            conn.close()

    room = f"prop_conv_{conv_id}"
    await sio.emit("property_conv_read", {"conv_id": conv_id, "reader_id": user_id}, room=room)
    return JSONResponse({"ok": True})


# =========================================================
# DIRECT MESSAGING (DM) ENDPOINTS
# =========================================================

class _DMSendRequest(BaseModel):
    conv_id: str
    content: str
    reply_to_id: str | None = None

class _DMStartRequest(BaseModel):
    other_user_id: str


def _get_dm_conversation(conv_id: str) -> dict | None:
    """Return the conversation row as a dict or None."""
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT conv_id,user1_id,user2_id,unread_u1,unread_u2,created_at FROM dm_conversations WHERE conv_id=?"
                if not USE_POSTGRES else
                "SELECT conv_id,user1_id,user2_id,unread_u1,unread_u2,created_at FROM dm_conversations WHERE conv_id=%s",
                (conv_id,),
            )
            row = cur.fetchone()
            if not row:
                return None
            return dict(zip(["conv_id","user1_id","user2_id","unread_u1","unread_u2","created_at"], row))
        finally:
            conn.close()


def _dm_increment_unread(conn, conv: dict, recipient_id: str) -> None:
    """Increment the unread counter for `recipient_id` in a DM conversation (within an open connection)."""
    if recipient_id == conv["user1_id"]:
        if USE_POSTGRES:
            conn.cursor().execute("UPDATE dm_conversations SET unread_u1=unread_u1+1 WHERE conv_id=%s", (conv["conv_id"],))
        else:
            conn.execute("UPDATE dm_conversations SET unread_u1=unread_u1+1 WHERE conv_id=?", (conv["conv_id"],))
    else:
        if USE_POSTGRES:
            conn.cursor().execute("UPDATE dm_conversations SET unread_u2=unread_u2+1 WHERE conv_id=%s", (conv["conv_id"],))
        else:
            conn.execute("UPDATE dm_conversations SET unread_u2=unread_u2+1 WHERE conv_id=?", (conv["conv_id"],))


def _dm_reset_unread(conn, conv: dict, reader_id: str) -> None:
    """Reset the unread counter for `reader_id` in a DM conversation (within an open connection)."""
    if reader_id == conv["user1_id"]:
        if USE_POSTGRES:
            conn.cursor().execute("UPDATE dm_conversations SET unread_u1=0 WHERE conv_id=%s", (conv["conv_id"],))
        else:
            conn.execute("UPDATE dm_conversations SET unread_u1=0 WHERE conv_id=?", (conv["conv_id"],))
    else:
        if USE_POSTGRES:
            conn.cursor().execute("UPDATE dm_conversations SET unread_u2=0 WHERE conv_id=%s", (conv["conv_id"],))
        else:
            conn.execute("UPDATE dm_conversations SET unread_u2=0 WHERE conv_id=?", (conv["conv_id"],))


def _find_or_create_conversation(user_a: str, user_b: str) -> dict:
    """Find an existing conversation between two users, or create one."""
    # Canonical order: smaller user_id is user1 so we get a unique pair
    u1, u2 = (user_a, user_b) if user_a < user_b else (user_b, user_a)
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT conv_id,user1_id,user2_id,unread_u1,unread_u2,created_at FROM dm_conversations WHERE user1_id=? AND user2_id=?"
                if not USE_POSTGRES else
                "SELECT conv_id,user1_id,user2_id,unread_u1,unread_u2,created_at FROM dm_conversations WHERE user1_id=%s AND user2_id=%s",
                (u1, u2),
            )
            row = cur.fetchone()
            if row:
                return dict(zip(["conv_id","user1_id","user2_id","unread_u1","unread_u2","created_at"], row))
            # Create new
            conv_id = str(uuid.uuid4())
            now = datetime.now(timezone.utc).isoformat()
            _execute(
                conn,
                "INSERT INTO dm_conversations (conv_id,user1_id,user2_id,unread_u1,unread_u2,created_at) VALUES (?,?,?,0,0,?)"
                if not USE_POSTGRES else
                "INSERT INTO dm_conversations (conv_id,user1_id,user2_id,unread_u1,unread_u2,created_at) VALUES (%s,%s,%s,0,0,%s)",
                (conv_id, u1, u2, now),
            )
            conn.commit()
            return {"conv_id": conv_id, "user1_id": u1, "user2_id": u2, "unread_u1": 0, "unread_u2": 0, "created_at": now}
        finally:
            conn.close()


@fastapi_app.get("/api/dm/conversations")
async def api_dm_list_conversations(request: Request):
    """List all DM conversations for the current user, with last-message preview."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)

    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT conv_id,user1_id,user2_id,unread_u1,unread_u2,created_at FROM dm_conversations WHERE user1_id=? OR user2_id=?"
                if not USE_POSTGRES else
                "SELECT conv_id,user1_id,user2_id,unread_u1,unread_u2,created_at FROM dm_conversations WHERE user1_id=%s OR user2_id=%s",
                (user_id, user_id),
            )
            rows = cur.fetchall()
        finally:
            conn.close()

    conversations = []
    for row in rows:
        conv = dict(zip(["conv_id","user1_id","user2_id","unread_u1","unread_u2","created_at"], row))
        other_id = conv["user2_id"] if conv["user1_id"] == user_id else conv["user1_id"]
        other = _get_app_user(other_id)
        unread = conv["unread_u1"] if conv["user1_id"] == user_id else conv["unread_u2"]

        # Get last message
        with _db_lock:
            conn = _get_db()
            try:
                cur = _execute(
                    conn,
                    "SELECT msg_id,sender_id,content,status,reply_to_id,ts FROM dm_messages WHERE conv_id=? ORDER BY ts DESC LIMIT 1"
                    if not USE_POSTGRES else
                    "SELECT msg_id,sender_id,content,status,reply_to_id,ts FROM dm_messages WHERE conv_id=%s ORDER BY ts DESC LIMIT 1",
                    (conv["conv_id"],),
                )
                lm = cur.fetchone()
            finally:
                conn.close()

        last_msg = None
        if lm:
            last_msg = dict(zip(["msg_id","sender_id","content","status","reply_to_id","ts"], lm))

        conversations.append({
            "conv_id":      conv["conv_id"],
            "other_user":   {"user_id": other_id, "name": other["name"] if other else other_id, "online_status": "offline"},
            "unread_count": unread,
            "last_message": last_msg,
            "created_at":   conv["created_at"],
        })

    # Sort by last message timestamp (most recent first)
    conversations.sort(key=lambda c: (c["last_message"]["ts"] if c["last_message"] else 0), reverse=True)
    return JSONResponse({"conversations": conversations})


@fastapi_app.post("/api/dm/conversations")
async def api_dm_start_conversation(request: Request, body: _DMStartRequest):
    """Find or create a DM conversation with another user."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    other_id = body.other_user_id.strip()
    if not other_id or other_id == user_id:
        return JSONResponse({"error": "Invalid user."}, status_code=400)
    other = _get_app_user(other_id)
    if not other:
        return JSONResponse({"error": "User not found."}, status_code=404)
    conv = _find_or_create_conversation(user_id, other_id)
    return JSONResponse({"conv": conv})


@fastapi_app.get("/api/dm/conversations/{conv_id}/messages")
async def api_dm_get_messages(request: Request, conv_id: str):
    """Get messages in a DM conversation (requires auth and participation)."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    conv = _get_dm_conversation(conv_id)
    if not conv:
        return JSONResponse({"error": "Conversation not found."}, status_code=404)
    if user_id not in (conv["user1_id"], conv["user2_id"]):
        return JSONResponse({"error": "Access denied."}, status_code=403)

    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT msg_id,conv_id,sender_id,content,status,reply_to_id,ts,created_at FROM dm_messages WHERE conv_id=? ORDER BY ts ASC LIMIT 200"
                if not USE_POSTGRES else
                "SELECT msg_id,conv_id,sender_id,content,status,reply_to_id,ts,created_at FROM dm_messages WHERE conv_id=%s ORDER BY ts ASC LIMIT 200",
                (conv_id,),
            )
            cols = ["msg_id","conv_id","sender_id","content","status","reply_to_id","ts","created_at"]
            messages = [dict(zip(cols, r)) for r in cur.fetchall()]
        finally:
            conn.close()

    return JSONResponse({"messages": messages, "conv": conv})


@fastapi_app.post("/api/dm/send")
async def api_dm_send(request: Request, body: _DMSendRequest):
    """Send a DM message in a conversation."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    content = body.content.strip()[:4000]  # 4000 chars to accommodate E2E encrypted payloads (base64)
    if not content:
        return JSONResponse({"error": "Message cannot be empty."}, status_code=400)
    conv = _get_dm_conversation(body.conv_id)
    if not conv:
        return JSONResponse({"error": "Conversation not found."}, status_code=404)
    if user_id not in (conv["user1_id"], conv["user2_id"]):
        return JSONResponse({"error": "Access denied."}, status_code=403)

    msg_id   = str(uuid.uuid4())
    ts       = time.time()
    now      = datetime.now(timezone.utc).isoformat()
    reply_to = body.reply_to_id or None
    other_id = conv["user2_id"] if conv["user1_id"] == user_id else conv["user1_id"]

    with _db_lock:
        conn = _get_db()
        try:
            _execute(
                conn,
                "INSERT OR IGNORE INTO dm_messages (msg_id,conv_id,sender_id,content,status,reply_to_id,ts,created_at) VALUES (?,?,?,?,'sent',?,?,?)"
                if not USE_POSTGRES else
                "INSERT INTO dm_messages (msg_id,conv_id,sender_id,content,status,reply_to_id,ts,created_at) VALUES (%s,%s,%s,%s,'sent',%s,%s,%s) ON CONFLICT (msg_id) DO NOTHING",
                (msg_id, body.conv_id, user_id, content, reply_to, ts, now),
            )
            # Increment unread for the other participant
            _dm_increment_unread(conn, conv, other_id)
            conn.commit()
        finally:
            conn.close()

    msg = {
        "msg_id":      msg_id,
        "conv_id":     body.conv_id,
        "sender_id":   user_id,
        "content":     content,
        "status":      "sent",
        "reply_to_id": reply_to,
        "ts":          ts,
    }

    # Push real-time delivery to the conversation room
    room = f"dm_{body.conv_id}"
    await sio.emit("dm_message", msg, room=room)

    # Notify the other user if they are online (not in the room)
    with _socket_user_lock:
        other_sid = _user_to_sid.get(other_id)
    if other_sid:
        me = _get_app_user(user_id)
        sender_name = me["name"] if me else "Someone"
        await sio.emit(
            "dm_notification",
            {"conv_id": body.conv_id, "from": sender_name, "preview": content[:80]},
            room=other_sid,
        )

    return JSONResponse({"ok": True, "message": msg})


@fastapi_app.post("/api/dm/read/{conv_id}")
async def api_dm_mark_read(request: Request, conv_id: str):
    """Mark all messages in a conversation as read for the current user."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    conv = _get_dm_conversation(conv_id)
    if not conv:
        return JSONResponse({"error": "Conversation not found."}, status_code=404)
    if user_id not in (conv["user1_id"], conv["user2_id"]):
        return JSONResponse({"error": "Access denied."}, status_code=403)

    with _db_lock:
        conn = _get_db()
        try:
            # Update message status for messages sent by the other participant
            other_id = conv["user2_id"] if conv["user1_id"] == user_id else conv["user1_id"]
            _execute(
                conn,
                "UPDATE dm_messages SET status='read' WHERE conv_id=? AND sender_id=? AND status!='read'"
                if not USE_POSTGRES else
                "UPDATE dm_messages SET status='read' WHERE conv_id=%s AND sender_id=%s AND status!='read'",
                (conv_id, other_id),
            )
            # Reset unread counter for the current user
            _dm_reset_unread(conn, conv, user_id)
            conn.commit()
        finally:
            conn.close()

    # Notify the other side that messages are read
    room = f"dm_{conv_id}"
    await sio.emit("dm_read", {"conv_id": conv_id, "reader_id": user_id}, room=room)
    return JSONResponse({"ok": True})


@fastapi_app.get("/api/users/list")
async def api_list_users(request: Request):
    """Return a list of all registered users (name + user_id) for starting new conversations."""
    user_id = request.session.get("app_user_id")
    if not user_id:
        return JSONResponse({"error": "Login required."}, status_code=401)
    with _db_lock:
        conn = _get_db()
        try:
            cur = _execute(
                conn,
                "SELECT user_id, name FROM app_users WHERE user_id!=? ORDER BY name ASC"
                if not USE_POSTGRES else
                "SELECT user_id, name FROM app_users WHERE user_id!=%s ORDER BY name ASC",
                (user_id,),
            )
            rows = cur.fetchall()
        finally:
            conn.close()
    users = [{"user_id": r[0], "name": r[1]} for r in rows]
    return JSONResponse({"users": users})


# =========================================================
# SOCKET.IO EVENTS
# =========================================================

@sio.event
async def connect(sid, environ):
    """Handle client connection"""
    logger.info(f"Client connected: {sid}")

@sio.event
async def disconnect(sid):
    """Handle client disconnection"""
    logger.info(f"Client disconnected: {sid}")
    with _socket_user_lock:
        user_id = _sid_to_user.pop(sid, None)
        if user_id:
            _user_to_sid.pop(user_id, None)

@sio.on("subscribe")
async def on_subscribe(sid, data):
    """Subscribe to download updates"""
    download_id = data.get("download_id") if isinstance(data, dict) else None
    if download_id:
        sio.enter_room(sid, download_id)
        await sio.emit("subscribed", {"id": download_id}, room=sid)
        logger.info(f"Client {sid} subscribed to {download_id}")

@sio.on("identify")
async def on_identify(sid, data):
    """Associate a socket connection with a logged-in user for targeted notifications."""
    user_id = data.get("user_id") if isinstance(data, dict) else None
    if user_id:
        with _socket_user_lock:
            _sid_to_user[sid] = user_id
            _user_to_sid[user_id] = sid
        await sio.emit("identified", {"user_id": user_id}, room=sid)
        logger.info(f"Socket {sid} identified as user {user_id}")


# ── Ride live-chat ──────────────────────────────────────────────────────────

@sio.on("join_ride_chat")
async def on_join_ride_chat(sid, data):
    """Subscribe the caller to a ride's chat room and send recent message history."""
    ride_id = data.get("ride_id") if isinstance(data, dict) else None
    if not ride_id:
        return
    room = f"ride_chat_{ride_id}"
    sio.enter_room(sid, room)
    sender_name = data.get("name", "Someone")

    # Load recent persisted messages
    history = []
    try:
        with _db_lock:
            conn = _get_db()
            try:
                cols = ["msg_id", "ride_id", "sender_name", "sender_role",
                        "text", "media_type", "media_data", "lat", "lng", "ts"]
                if USE_POSTGRES:
                    cur = conn.cursor()
                    cur.execute(
                        "SELECT msg_id,ride_id,sender_name,sender_role,text,media_type,media_data,lat,lng,ts"
                        " FROM ride_chat_messages WHERE ride_id=%s ORDER BY ts ASC LIMIT 50",
                        (ride_id,),
                    )
                    rows = cur.fetchall()
                else:
                    cur = conn.execute(
                        "SELECT msg_id,ride_id,sender_name,sender_role,text,media_type,media_data,lat,lng,ts"
                        " FROM ride_chat_messages WHERE ride_id=? ORDER BY ts ASC LIMIT 50",
                        (ride_id,),
                    )
                    rows = cur.fetchall()
                for row in rows:
                    d = dict(zip(cols, row))
                    d["name"] = d.pop("sender_name")
                    d["role"] = d.pop("sender_role")
                    d["id"]   = d.pop("msg_id")
                    history.append(d)
                _resolve_chat_media(history)
            finally:
                conn.close()
    except Exception as e:
        logger.warning(f"Failed to load ride chat history: {e}")

    await sio.emit("ride_chat_joined", {
        "ride_id": ride_id, "name": sender_name, "history": history,
    }, room=sid)
    logger.info(f"Socket {sid} joined ride chat room {room} ({len(history)} history messages)")


@sio.on("leave_ride_chat")
async def on_leave_ride_chat(sid, data):
    """Unsubscribe the caller from a ride's chat room."""
    ride_id = data.get("ride_id") if isinstance(data, dict) else None
    if not ride_id:
        return
    room = f"ride_chat_{ride_id}"
    sio.leave_room(sid, room)
    logger.info(f"Socket {sid} left ride chat room {room}")


@sio.on("ride_chat_message")
async def on_ride_chat_message(sid, data):
    """Broadcast a chat message to all members of a ride's chat room.

    Supports text, image (media_type='image'), audio (media_type='audio'),
    and location (media_type='location') payloads.
    Persists the message and notifies the ride poster when the sender is
    a different user.
    """
    if not isinstance(data, dict):
        return
    ride_id    = data.get("ride_id")
    text       = str(data.get("text", "")).strip()
    name       = str(data.get("name", "Anonymous")).strip()
    role       = str(data.get("role", "passenger")).strip()
    msg_id     = data.get("id") or f"{time.time()}-{sid}"
    media_type = data.get("media_type")  # 'image' | 'audio' | 'location' | None
    media_data = data.get("media_data")  # base64 string for image/audio
    msg_lat    = data.get("lat")         # float, for location messages
    msg_lng    = data.get("lng")         # float, for location messages

    # Require either text or media
    if not ride_id or (not text and not media_type):
        return

    # Validate media_type
    _ALLOWED_MEDIA = {"image", "audio", "location"}
    if media_type and media_type not in _ALLOWED_MEDIA:
        media_type = None

    # Cap text length; limit base64 payload to ~1 MB
    text = text[:500]
    if media_data and len(media_data) > 1_400_000:
        media_data = None

    # When S3 is configured, upload image/audio media to the bucket and store
    # the S3 object key in place of the raw base64 payload.  This keeps large
    # binary blobs out of the database and ensures persistence across restarts.
    _MEDIA_TYPE_INFO = {
        "image": ("jpg",  "image/jpeg"),
        "audio": ("webm", "audio/webm"),
    }
    db_media_data = media_data  # value written to the database
    if _S3_ENABLED and media_data and media_type in _MEDIA_TYPE_INFO:
        try:
            raw_bytes = base64.b64decode(media_data)
            ext, mime = _MEDIA_TYPE_INFO[media_type]
            s3_key = f"chat_media/{msg_id}.{ext}"
            if _s3_upload_bytes(raw_bytes, s3_key, mime):
                db_media_data = s3_key  # store only the S3 key in the DB
        except Exception as _exc:
            logger.warning("Failed to upload ride chat media to S3: %s", _exc)

    ts = time.time()
    msg = {
        "ride_id":    ride_id,
        "name":       name,
        "text":       text,
        "ts":         ts,
        "role":       role,
        "id":         msg_id,
        "media_type": media_type,
        "media_data": media_data,
        "lat":        msg_lat,
        "lng":        msg_lng,
    }

    # Persist the message to the database
    created = datetime.now(timezone.utc).isoformat()
    try:
        with _db_lock:
            conn = _get_db()
            try:
                if USE_POSTGRES:
                    cur = conn.cursor()
                    cur.execute(
                        "INSERT INTO ride_chat_messages"
                        " (msg_id,ride_id,sender_name,sender_role,text,media_type,media_data,lat,lng,ts,created_at)"
                        " VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
                        " ON CONFLICT (msg_id) DO NOTHING",
                        (msg_id, ride_id, name, role, text or None, media_type, db_media_data,
                         msg_lat, msg_lng, ts, created),
                    )
                else:
                    conn.execute(
                        "INSERT OR IGNORE INTO ride_chat_messages"
                        " (msg_id,ride_id,sender_name,sender_role,text,media_type,media_data,lat,lng,ts,created_at)"
                        " VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                        (msg_id, ride_id, name, role, text or None, media_type, db_media_data,
                         msg_lat, msg_lng, ts, created),
                    )
                conn.commit()
            finally:
                conn.close()
    except Exception as e:
        logger.warning(f"Failed to persist ride chat message: {e}")

    # Notify the ride poster if the sender is a different user
    try:
        with _db_lock:
            conn = _get_db()
            try:
                if USE_POSTGRES:
                    cur = conn.cursor()
                    cur.execute("SELECT user_id FROM rides WHERE ride_id=%s", (ride_id,))
                    row = cur.fetchone()
                else:
                    cur = conn.execute("SELECT user_id FROM rides WHERE ride_id=?", (ride_id,))
                    row = cur.fetchone()
                poster_user_id = row[0] if row else None
            finally:
                conn.close()

        if poster_user_id:
            # Check if sender is the poster
            poster = _get_app_user(poster_user_id)
            if poster and poster.get("name") != name:
                # Notify poster about new message
                preview = text if text else f"[{media_type}]" if media_type else ""
                notif_id = _create_notification(
                    poster_user_id,
                    "chat_message",
                    f"New message from {name}",
                    f"Ride: {ride_id[:8]}… — {preview[:80]}",
                )
                # Push real-time notification to the poster's socket if online
                with _socket_user_lock:
                    poster_sid = _user_to_sid.get(poster_user_id)
                if poster_sid:
                    await sio.emit(
                        "ride_chat_notification",
                        {
                            "notif_id": notif_id,
                            "ride_id":  ride_id,
                            "from":     name,
                            "preview":  preview[:80],
                        },
                        room=poster_sid,
                    )
    except Exception as e:
        logger.warning(f"Failed to notify ride poster of chat message: {e}")

    room = f"ride_chat_{ride_id}"
    await sio.emit("ride_chat_message", msg, room=room)
    log_preview = text[:80] if text else f"[{media_type}]"
    logger.info(f"Ride chat [{ride_id}] from {name}: {log_preview}")

    # Auto-response: when a passenger sends their first message in the chat,
    # immediately reply with a structured booking prompt so the driver (or
    # system) collects all required details without back-and-forth friction.
    if role != "driver" and text:
        try:
            with _db_lock:
                conn = _get_db()
                try:
                    if USE_POSTGRES:
                        cur = conn.cursor()
                        cur.execute(
                            "SELECT COUNT(*) FROM ride_chat_messages"
                            " WHERE ride_id=%s AND sender_role != 'driver'",
                            (ride_id,),
                        )
                        passenger_msg_count = cur.fetchone()[0]
                    else:
                        cur = conn.execute(
                            "SELECT COUNT(*) FROM ride_chat_messages"
                            " WHERE ride_id=? AND sender_role != 'driver'",
                            (ride_id,),
                        )
                        passenger_msg_count = cur.fetchone()[0]
                finally:
                    conn.close()

            if passenger_msg_count == 1:
                # This is the passenger's first message — send the auto-response.
                auto_id = f"auto-{ride_id}-{time.time()}"
                auto_msg = {
                    "ride_id":    ride_id,
                    "name":       "System",
                    "text":       (
                        "Please share your current location, full name, and "
                        "contact number to complete your booking."
                    ),
                    "ts":         time.time(),
                    "role":       "system",
                    "id":         auto_id,
                    "media_type": None,
                    "media_data": None,
                    "lat":        None,
                    "lng":        None,
                }
                auto_created = datetime.now(timezone.utc).isoformat()
                with _db_lock:
                    conn = _get_db()
                    try:
                        if USE_POSTGRES:
                            cur = conn.cursor()
                            cur.execute(
                                "INSERT INTO ride_chat_messages"
                                " (msg_id,ride_id,sender_name,sender_role,text,"
                                "  media_type,media_data,lat,lng,ts,created_at)"
                                " VALUES (%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s)"
                                " ON CONFLICT (msg_id) DO NOTHING",
                                (auto_id, ride_id, "System", "system",
                                 auto_msg["text"], None, None,
                                 None, None, auto_msg["ts"], auto_created),
                            )
                        else:
                            conn.execute(
                                "INSERT OR IGNORE INTO ride_chat_messages"
                                " (msg_id,ride_id,sender_name,sender_role,text,"
                                "  media_type,media_data,lat,lng,ts,created_at)"
                                " VALUES (?,?,?,?,?,?,?,?,?,?,?)",
                                (auto_id, ride_id, "System", "system",
                                 auto_msg["text"], None, None,
                                 None, None, auto_msg["ts"], auto_created),
                            )
                        conn.commit()
                    finally:
                        conn.close()
                await sio.emit("ride_chat_message", auto_msg, room=room)
                logger.info(f"Ride chat [{ride_id}] auto-response sent to {name}")
        except Exception as e:
            logger.warning(f"Failed to send auto-response for ride {ride_id}: {e}")


@sio.on("ride_chat_typing")
async def on_ride_chat_typing(sid, data):
    """Broadcast typing indicator to the ride's chat room (exclude sender)."""
    if not isinstance(data, dict):
        return
    ride_id = data.get("ride_id")
    name    = str(data.get("name", "")).strip()
    if not ride_id or not name:
        return
    room = f"ride_chat_{ride_id}"
    await sio.emit("ride_chat_typing", {"ride_id": ride_id, "name": name}, room=room, skip_sid=sid)


@sio.on("ride_chat_stop_typing")
async def on_ride_chat_stop_typing(sid, data):
    """Broadcast stop-typing event to the ride's chat room."""
    if not isinstance(data, dict):
        return
    ride_id = data.get("ride_id")
    name    = str(data.get("name", "")).strip()
    if not ride_id or not name:
        return
    room = f"ride_chat_{ride_id}"
    await sio.emit("ride_chat_stop_typing", {"ride_id": ride_id, "name": name}, room=room, skip_sid=sid)


@sio.on("ride_chat_read")
async def on_ride_chat_read(sid, data):
    """Broadcast read receipt to the ride's chat room."""
    if not isinstance(data, dict):
        return
    ride_id = data.get("ride_id")
    msg_id  = data.get("msg_id")
    reader  = str(data.get("reader", "")).strip()
    if not ride_id or not msg_id or not reader:
        return
    room = f"ride_chat_{ride_id}"
    await sio.emit("ride_chat_read", {"ride_id": ride_id, "msg_id": msg_id, "reader": reader}, room=room)


# ── Agent live-chat ─────────────────────────────────────────────────────────

@sio.on("agent_chat_join")
async def on_agent_chat_join(sid, data):
    """Subscribe caller to an agent's chat room."""
    if not isinstance(data, dict):
        return
    agent_id = data.get("agent_id")
    user_id  = data.get("user_id")
    if not agent_id:
        return
    room = f"agent_chat_{agent_id}_{user_id}" if user_id else f"agent_chat_{agent_id}"
    sio.enter_room(sid, room)
    await sio.emit("agent_chat_joined", {"agent_id": agent_id, "room": room}, room=sid)


@sio.on("agent_chat_leave")
async def on_agent_chat_leave(sid, data):
    """Unsubscribe caller from an agent's chat room."""
    if not isinstance(data, dict):
        return
    agent_id = data.get("agent_id")
    user_id  = data.get("user_id")
    if not agent_id:
        return
    room = f"agent_chat_{agent_id}_{user_id}" if user_id else f"agent_chat_{agent_id}"
    sio.leave_room(sid, room)


@sio.on("agent_chat_message")
async def on_agent_chat_message(sid, data):
    """Persist and broadcast a chat message between a user and an agent."""
    if not isinstance(data, dict):
        return
    agent_id    = data.get("agent_id")
    user_id     = data.get("user_id")
    sender_role = str(data.get("sender_role", "user"))
    text        = str(data.get("text", "")).strip()[:500]
    if not agent_id or not user_id or not text:
        return
    if sender_role not in ("user", "agent"):
        sender_role = "user"

    msg_id = str(uuid.uuid4())
    ts     = time.time()
    now    = datetime.utcnow().isoformat()

    # Persist message
    try:
        with _db_lock:
            conn = _get_db()
            try:
                if USE_POSTGRES:
                    conn.cursor().execute(
                        "INSERT INTO agent_chat_messages (msg_id,agent_id,user_id,sender_role,text,ts,created_at) VALUES (%s,%s,%s,%s,%s,%s,%s)",
                        (msg_id, agent_id, user_id, sender_role, text, ts, now),
                    )
                else:
                    conn.execute(
                        "INSERT OR IGNORE INTO agent_chat_messages (msg_id,agent_id,user_id,sender_role,text,ts,created_at) VALUES (?,?,?,?,?,?,?)",
                        (msg_id, agent_id, user_id, sender_role, text, ts, now),
                    )
                conn.commit()
            finally:
                conn.close()
    except Exception as e:
        logger.warning(f"Failed to persist agent chat message: {e}")

    msg = {"msg_id": msg_id, "agent_id": agent_id, "user_id": user_id,
           "sender_role": sender_role, "text": text, "ts": ts}
    room = f"agent_chat_{agent_id}_{user_id}"
    await sio.emit("agent_chat_message", msg, room=room)

    # Also notify the agent's linked socket if message is from user
    if sender_role == "user":
        agent = _get_agent_row(agent_id)
        if agent and agent.get("user_id"):
            agent_user_id = agent["user_id"]
            with _socket_user_lock:
                agent_sid = _user_to_sid.get(agent_user_id)
            if agent_sid:
                try:
                    await sio.emit("agent_chat_message", msg, room=agent_sid)
                except Exception as e:
                    logger.warning(f"Failed to notify agent of chat message: {e}")

    logger.info(f"Agent chat [{agent_id}] user {user_id}: {text[:80]}")


@sio.on("agent_chat_typing")
async def on_agent_chat_typing(sid, data):
    """Broadcast typing indicator in an agent chat room."""
    if not isinstance(data, dict):
        return
    agent_id = data.get("agent_id")
    user_id  = data.get("user_id")
    name     = str(data.get("name", "")).strip()
    if not agent_id or not user_id:
        return
    room = f"agent_chat_{agent_id}_{user_id}"
    await sio.emit("agent_chat_typing", {"agent_id": agent_id, "user_id": user_id, "name": name}, room=room, skip_sid=sid)


@sio.on("agent_chat_stop_typing")
async def on_agent_chat_stop_typing(sid, data):
    """Broadcast stop-typing in an agent chat room."""
    if not isinstance(data, dict):
        return
    agent_id = data.get("agent_id")
    user_id  = data.get("user_id")
    name     = str(data.get("name", "")).strip()
    if not agent_id or not user_id:
        return
    room = f"agent_chat_{agent_id}_{user_id}"
    await sio.emit("agent_chat_stop_typing", {"agent_id": agent_id, "user_id": user_id, "name": name}, room=room, skip_sid=sid)


@sio.on("agent_profile_visit")
async def on_agent_profile_visit(sid, data):
    """Notify an agent when a client visits their profile page."""
    if not isinstance(data, dict):
        return
    agent_id  = str(data.get("agent_id", "")).strip()
    visitor_id = str(data.get("user_id", "anonymous")).strip()
    if not agent_id:
        return

    # Notify the agent's linked socket if they are online
    agent = _get_agent_row(agent_id)
    if agent and agent.get("user_id"):
        agent_user_id = agent["user_id"]
        with _socket_user_lock:
            agent_sid = _user_to_sid.get(agent_user_id)
        if agent_sid:
            try:
                await sio.emit(
                    "agent_profile_visit_notify",
                    {"agent_id": agent_id, "visitor_id": visitor_id, "ts": time.time()},
                    room=agent_sid,
                )
            except Exception as e:
                logger.warning(f"Failed to notify agent of profile visit: {e}")

    logger.info(f"Agent profile visited: agent {agent_id} by visitor {visitor_id}")


# ── Direct Messaging live-chat ───────────────────────────────────────────────

@sio.on("dm_join")
async def on_dm_join(sid, data):
    """Subscribe caller to a DM conversation room and send recent history."""
    if not isinstance(data, dict):
        return
    conv_id = data.get("conv_id")
    if not conv_id:
        return
    room = f"dm_{conv_id}"
    sio.enter_room(sid, room)

    # Load recent persisted messages
    history = []
    try:
        with _db_lock:
            conn = _get_db()
            try:
                cur = _execute(
                    conn,
                    "SELECT msg_id,conv_id,sender_id,content,status,reply_to_id,ts FROM dm_messages WHERE conv_id=? ORDER BY ts ASC LIMIT 100"
                    if not USE_POSTGRES else
                    "SELECT msg_id,conv_id,sender_id,content,status,reply_to_id,ts FROM dm_messages WHERE conv_id=%s ORDER BY ts ASC LIMIT 100",
                    (conv_id,),
                )
                cols = ["msg_id","conv_id","sender_id","content","status","reply_to_id","ts"]
                history = [dict(zip(cols, r)) for r in cur.fetchall()]
            finally:
                conn.close()
    except Exception as e:
        logger.warning(f"Failed to load DM history: {e}")

    await sio.emit("dm_joined", {"conv_id": conv_id, "history": history}, room=sid)
    logger.info(f"Socket {sid} joined DM room {room} ({len(history)} history messages)")


@sio.on("dm_leave")
async def on_dm_leave(sid, data):
    """Unsubscribe caller from a DM conversation room."""
    if not isinstance(data, dict):
        return
    conv_id = data.get("conv_id")
    if not conv_id:
        return
    sio.leave_room(sid, f"dm_{conv_id}")


@sio.on("dm_message")
async def on_dm_message(sid, data):
    """Persist and broadcast a direct message."""
    if not isinstance(data, dict):
        return
    conv_id    = data.get("conv_id")
    sender_id  = data.get("sender_id")
    content    = str(data.get("content", "")).strip()[:1000]
    reply_to   = data.get("reply_to_id")
    client_id  = data.get("id")  # client-generated optimistic id

    if not conv_id or not sender_id or not content:
        return

    # Verify conversation exists
    conv = _get_dm_conversation(conv_id)
    if not conv or sender_id not in (conv["user1_id"], conv["user2_id"]):
        return

    msg_id = client_id or str(uuid.uuid4())
    ts     = time.time()
    now    = datetime.now(timezone.utc).isoformat()
    other_id = conv["user2_id"] if conv["user1_id"] == sender_id else conv["user1_id"]

    try:
        with _db_lock:
            conn = _get_db()
            try:
                _execute(
                    conn,
                    "INSERT OR IGNORE INTO dm_messages (msg_id,conv_id,sender_id,content,status,reply_to_id,ts,created_at) VALUES (?,?,?,?,'sent',?,?,?)"
                    if not USE_POSTGRES else
                    "INSERT INTO dm_messages (msg_id,conv_id,sender_id,content,status,reply_to_id,ts,created_at) VALUES (%s,%s,%s,%s,'sent',%s,%s,%s) ON CONFLICT (msg_id) DO NOTHING",
                    (msg_id, conv_id, sender_id, content, reply_to, ts, now),
                )
                _dm_increment_unread(conn, conv, other_id)
                conn.commit()
            finally:
                conn.close()
    except Exception as e:
        logger.warning(f"Failed to persist DM message: {e}")

    msg = {
        "msg_id":      msg_id,
        "conv_id":     conv_id,
        "sender_id":   sender_id,
        "content":     content,
        "status":      "sent",
        "reply_to_id": reply_to,
        "ts":          ts,
    }

    room = f"dm_{conv_id}"
    await sio.emit("dm_message", msg, room=room)

    # Notify other user if online
    with _socket_user_lock:
        other_sid = _user_to_sid.get(other_id)
    if other_sid:
        me = _get_app_user(sender_id)
        sender_name = me["name"] if me else "Someone"
        await sio.emit(
            "dm_notification",
            {"conv_id": conv_id, "from": sender_name, "preview": content[:80]},
            room=other_sid,
        )

    logger.info(f"DM [{conv_id}] from {sender_id}: {content[:80]}")


@sio.on("dm_typing")
async def on_dm_typing(sid, data):
    """Broadcast typing indicator in a DM conversation room (exclude sender)."""
    if not isinstance(data, dict):
        return
    conv_id   = data.get("conv_id")
    sender_id = data.get("sender_id")
    if not conv_id or not sender_id:
        return
    room = f"dm_{conv_id}"
    await sio.emit("dm_typing", {"conv_id": conv_id, "sender_id": sender_id}, room=room, skip_sid=sid)


@sio.on("dm_stop_typing")
async def on_dm_stop_typing(sid, data):
    """Broadcast stop-typing indicator in a DM conversation room."""
    if not isinstance(data, dict):
        return
    conv_id   = data.get("conv_id")
    sender_id = data.get("sender_id")
    if not conv_id or not sender_id:
        return
    room = f"dm_{conv_id}"
    await sio.emit("dm_stop_typing", {"conv_id": conv_id, "sender_id": sender_id}, room=room, skip_sid=sid)


@sio.on("dm_read")
async def on_dm_read(sid, data):
    """Mark messages in a DM conversation as read and broadcast the event."""
    if not isinstance(data, dict):
        return
    conv_id   = data.get("conv_id")
    reader_id = data.get("reader_id")
    if not conv_id or not reader_id:
        return

    conv = _get_dm_conversation(conv_id)
    if not conv or reader_id not in (conv["user1_id"], conv["user2_id"]):
        return

    other_id = conv["user2_id"] if conv["user1_id"] == reader_id else conv["user1_id"]

    try:
        with _db_lock:
            conn = _get_db()
            try:
                _execute(
                    conn,
                    "UPDATE dm_messages SET status='read' WHERE conv_id=? AND sender_id=? AND status!='read'"
                    if not USE_POSTGRES else
                    "UPDATE dm_messages SET status='read' WHERE conv_id=%s AND sender_id=%s AND status!='read'",
                    (conv_id, other_id),
                )
                _dm_reset_unread(conn, conv, reader_id)
                conn.commit()
            finally:
                conn.close()
    except Exception as e:
        logger.warning(f"Failed to mark DM messages read: {e}")

    room = f"dm_{conv_id}"
    await sio.emit("dm_read", {"conv_id": conv_id, "reader_id": reader_id}, room=room)


# ── Property Conversation live-chat ─────────────────────────────────────────

@sio.on("prop_conv_join")
async def on_prop_conv_join(sid, data):
    """Subscribe caller to a property conversation room and send recent history.

    Only the two participants of the conversation (buyer and agent) are allowed
    to join.  Any other socket is silently rejected to preserve message privacy.
    """
    if not isinstance(data, dict):
        return
    conv_id = data.get("conv_id")
    if not conv_id:
        return

    # Resolve the authenticated user for this socket connection
    with _socket_user_lock:
        user_id = _sid_to_user.get(sid)

    # Load the conversation record to verify participation
    conv = _get_property_conversation(conv_id)
    if not conv:
        await sio.emit("prop_conv_error", {"conv_id": conv_id, "error": "Conversation not found."}, room=sid)
        return

    # Enforce privacy: only the buyer (user_id) or the agent (agent_id) may join
    if not user_id or user_id not in (conv["user_id"], conv["agent_id"]):
        await sio.emit("prop_conv_error", {"conv_id": conv_id, "error": "Access denied."}, room=sid)
        return

    room = f"prop_conv_{conv_id}"
    sio.enter_room(sid, room)

    history = []
    try:
        with _db_lock:
            conn = _get_db()
            try:
                cur = _execute(
                    conn,
                    "SELECT msg_id,conv_id,sender_id,sender_role,content,ts FROM property_messages WHERE conv_id=? ORDER BY ts ASC LIMIT 100"
                    if not USE_POSTGRES else
                    "SELECT msg_id,conv_id,sender_id,sender_role,content,ts FROM property_messages WHERE conv_id=%s ORDER BY ts ASC LIMIT 100",
                    (conv_id,),
                )
                cols = ["msg_id","conv_id","sender_id","sender_role","content","ts"]
                history = [dict(zip(cols, r)) for r in cur.fetchall()]
            finally:
                conn.close()
    except Exception as e:
        logger.warning(f"Failed to load property conversation history: {e}")

    await sio.emit("prop_conv_joined", {"conv_id": conv_id, "history": history}, room=sid)


@sio.on("prop_conv_leave")
async def on_prop_conv_leave(sid, data):
    """Unsubscribe caller from a property conversation room."""
    if not isinstance(data, dict):
        return
    conv_id = data.get("conv_id")
    if not conv_id:
        return
    sio.leave_room(sid, f"prop_conv_{conv_id}")


@sio.on("prop_conv_typing")
async def on_prop_conv_typing(sid, data):
    """Broadcast typing indicator to property conversation room."""
    if not isinstance(data, dict):
        return
    conv_id   = data.get("conv_id")
    sender_id = data.get("sender_id")
    if not conv_id or not sender_id:
        return
    room = f"prop_conv_{conv_id}"
    await sio.emit("prop_conv_typing", {"conv_id": conv_id, "sender_id": sender_id}, room=room, skip_sid=sid)


@sio.on("prop_conv_stop_typing")
async def on_prop_conv_stop_typing(sid, data):
    """Broadcast stop-typing indicator to property conversation room."""
    if not isinstance(data, dict):
        return
    conv_id   = data.get("conv_id")
    sender_id = data.get("sender_id")
    if not conv_id or not sender_id:
        return
    room = f"prop_conv_{conv_id}"
    await sio.emit("prop_conv_stop_typing", {"conv_id": conv_id, "sender_id": sender_id}, room=room, skip_sid=sid)


def cleanup_old_files():
    """Delete files that have been on disk for longer than FILE_RETENTION_MINUTES."""
    _ensure_download_folder()
    try:
        current_time = time.time()
        cutoff = current_time - (Config.FILE_RETENTION_MINUTES * 60)
        files_deleted = False

        try:
            listing = os.listdir(DOWNLOAD_FOLDER)
        except FileNotFoundError:
            # Directory was removed between the ensure call and the listdir;
            # nothing to clean up this cycle.
            return

        for filename in listing:
            filepath = os.path.join(DOWNLOAD_FOLDER, filename)
            if os.path.isfile(filepath):
                if os.path.getmtime(filepath) < cutoff:
                    os.remove(filepath)
                    if _S3_ENABLED:
                        _s3_delete_file(filename)
                    logger.info(f"Auto-cleaned file (>{Config.FILE_RETENTION_MINUTES} min): {filename}")
                    files_deleted = True

        if files_deleted:
            # Notify all connected clients from this background thread so they
            # refresh their "Downloaded Videos" list without waiting for the next
            # manual refresh or Socket.IO reconnect.
            emit_from_thread("files_updated")
    except Exception as e:
        logger.error(f"Cleanup error: {e}")

def cleanup_thread():
    """Background thread for periodic cleanup"""
    while True:
        time.sleep(Config.CLEANUP_INTERVAL)
        cleanup_old_files()

        # Flush in-memory records to DB before evicting stale ones so that
        # the admin dashboard (which reads from the DB) retains a complete
        # download log even after the video files and in-memory entries are
        # cleaned up.
        save_downloads_to_disk()

        # Remove stale in-memory download records (keep for 1 h after completion)
        with downloads_lock:
            current_time = time.time()
            to_delete = []
            for did, d in downloads.items():
                if d["status"] in ("completed", "failed", "cancelled"):
                    if current_time - d.get("end_time", current_time) > 3600:
                        to_delete.append(did)
            for did in to_delete:
                del downloads[did]

# =========================================================
# INITIALIZATION
# =========================================================

logger.info("=" * 50)
logger.info("Starting Video Downloader (Production)")
logger.info("=" * 50)

# Warn if ADMIN_PASSWORD was not explicitly set
if not os.environ.get("ADMIN_PASSWORD"):
    logger.warning(
        "WARNING: ADMIN_PASSWORD env var is not set. "
        "A random password has been generated for this session: %s  "
        "Set the ADMIN_PASSWORD environment variable to a persistent strong password before deploying.",
        Config.ADMIN_PASSWORD,
    )

# Check dependencies
check_yt_dlp()
check_ffmpeg()

# Log paths
logger.info(f"Root directory: {ROOT_DIR}")
logger.info(f"Templates directory: {TEMPLATES_DIR}")
logger.info(f"Downloads directory: {DOWNLOAD_FOLDER}")
if DOWNLOAD_FOLDER != _configured_download_folder:
    logger.warning(
        f"Could not create downloads directory at {_configured_download_folder!r}; "
        f"falling back to {DOWNLOAD_FOLDER!r}. "
        "Set the DOWNLOAD_DIR environment variable to a writable path to persist downloads."
    )
logger.info(f"Template exists: {os.path.exists(os.path.join(TEMPLATES_DIR, 'index.html'))}")

# Initialise database schema
init_db()
if USE_POSTGRES:
    logger.info("Using PostgreSQL database via DATABASE_URL")
else:
    logger.info(f"Using SQLite database at {DB_PATH}")

if _S3_ENABLED:
    logger.info(
        "S3 object storage enabled: bucket=%s region=%s endpoint=%s",
        BUCKET_NAME,
        BUCKET_REGION or "(default)",
        BUCKET_ENDPOINT or "(AWS)",
    )
else:
    logger.info("S3 object storage disabled (BUCKET_* variables not set)")

# Load persisted data
load_persistence()

# Start cleanup thread
cleanup_thread = threading.Thread(target=cleanup_thread, daemon=True)
cleanup_thread.start()

logger.info("=" * 50)

# =========================================================
# ERROR HANDLERS
# =========================================================

@fastapi_app.exception_handler(StarletteHTTPException)
async def http_exception_handler(request: Request, exc: StarletteHTTPException):
    # For 404s on non-API paths, serve the React SPA so client-side routing works
    if exc.status_code == 404:
        path = request.url.path
        # Only serve SPA for non-API / non-static paths
        api_prefixes = (
            "/video_info", "/start_download", "/status/", "/files", "/downloads/",
            "/stream/", "/delete/", "/cancel", "/cancel_all", "/stats",
            "/active_downloads", "/download_zip", "/reviews",
            "/admin/downloads", "/admin/visitors", "/admin/analytics",
            "/admin/cancel_download/", "/admin/delete_record/", "/admin/clear_visitors",
            "/admin/db/", "/admin/cookies", "/admin/auth_status", "/admin/has_admin",
            "/admin/api/", "/health", "/ads.txt", "/static/", "/assets/",
            "/api/", "/yotweek.png",
        )
        if not any(path.startswith(p) for p in api_prefixes):
            return _react_index()
        return JSONResponse({"error": "Not found"}, status_code=404)
    if exc.status_code == 500:
        logger.error(f"Internal error: {exc.detail}")
        return JSONResponse({"error": "Internal server error"}, status_code=500)
    return JSONResponse({"error": exc.detail or "Error"}, status_code=exc.status_code)

@fastapi_app.exception_handler(500)
async def internal_error(request: Request, exc):
    logger.error(f"Internal error: {exc}")
    return JSONResponse({"error": "Internal server error"}, status_code=500)

# =========================================================
# ENTRY POINT
# =========================================================

if __name__ == "__main__":
    import uvicorn
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("FLASK_DEBUG", "False").lower() == "true"

    logger.info(f"Starting server on port {port}")
    logger.info(f"Debug mode: {debug}")

    uvicorn.run(
        "api.app:app",
        host="0.0.0.0",
        port=port,
        reload=debug,
        workers=1,
    )